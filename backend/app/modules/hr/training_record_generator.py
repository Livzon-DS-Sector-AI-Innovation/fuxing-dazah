"""Generate 新员工培训记录 (new employee training record) from DOCX template."""

from datetime import date
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document

from app.modules.hr.models import Employee


def _find_template() -> Path:
    candidates = [
        Path("assets/hr/新员工培训记录-模板.docx"),
        Path("../assets/hr/新员工培训记录-模板.docx"),
        Path(__file__).resolve().parent.parent.parent.parent
        / "assets/hr"
        / "新员工培训记录-模板.docx",
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError("模板文件未找到: 新员工培训记录-模板.docx")


def _replace_all(doc: Document, placeholders: dict[str, str]) -> None:
    # 按 key 长度倒序，避免短 key 先替换破坏长 key（如 {入职日期} 破坏 {{入职日期}+3M-1day}）
    ordered = sorted(placeholders.items(), key=lambda kv: -len(kv[0]))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, val in ordered:
                        if key in para.text:
                            para.text = para.text.replace(key, val)


def _calc_age(birth_year: int | None) -> str:
    if not birth_year:
        return ""
    try:
        return str(date.today().year - int(birth_year))
    except (ValueError, TypeError):
        return ""


def _calc_probation_end(date_val: Any) -> str:
    """入职日期 + 3 个月 - 1 天"""
    from datetime import timedelta
    from dateutil.relativedelta import relativedelta

    if not date_val:
        return ""
    d = date_val if isinstance(date_val, date) else date.fromisoformat(str(date_val))
    end = d + relativedelta(months=3) - timedelta(days=1)
    return str(end)


def generate_training_record(
    employee: Employee,
    training_items: list[dict[str, Any]] | None = None,
) -> BytesIO:
    """生成新员工培训记录 Word 文档。

    training_items 每项包含: file_name, trainer, method, plan_date
    """
    doc = Document(str(_find_template()))

    # ── Part I: 员工概况占位符 ──
    probation_end = _calc_probation_end(employee.hire_date)
    age_str = _calc_age(employee.birth_year)

    _replace_all(doc, {
        "{姓名}": employee.name or "",
        "{性别}": employee.gender or "",
        "{年龄}": age_str,
        "{体现部门}": employee.department or "",
        "{体现岗位}": employee.position or "",
        "{入职日期}": str(employee.hire_date) if employee.hire_date else "",
        "{{入职日期}+3M-1day}": probation_end,
    })

    # ── Part II: 填写其他培训内容 (R17-R26, 10个空行) ──
    if training_items and doc.tables:
        t = doc.tables[0]

        # 从 R17 开始填入动态培训内容（skip 固定内容 R6-R16）
        start_row = 17
        max_rows = 10  # 最多填充10行
        row_offset = 0

        for item in training_items[:max_rows]:
            ri = start_row + row_offset
            if ri >= len(t.rows):
                break
            row = t.rows[ri]

            content = item.get("content") or item.get("file_name") or ""
            sop_number = item.get("sop_number", "")
            title = f"{sop_number} {content}" if sop_number else content

            # 第一列：培训内容 + 考核成绩
            if row.cells[0].paragraphs:
                row.cells[0].paragraphs[0].text = f"{title}                             考核成绩/Score："

            row_offset += 1

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
