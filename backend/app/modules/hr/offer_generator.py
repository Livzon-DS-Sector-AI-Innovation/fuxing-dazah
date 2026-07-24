"""入职 Offer 生成器：DOCX模板填占位符 + HTML预览 + Weasyprint转PDF。"""

from datetime import date
from io import BytesIO

from docx import Document

from app.modules.hr.template_utils import find_hr_template

# 模板占位符 → 参数名映射
PLACEHOLDER_MAP = {
    "{姓名}": "name",
    "{岗位}": "position",
    "{底薪}": "base_salary",
    "{综合月薪}": "salary_range",
    "{体检日期}": "medical_date",
    "{报到日期}": "report_date",
    "{岗位保留时间}": "offer_expire_date",
    "{发送日期}": "send_date",
}


def _replace_in_paragraph(p, replacements: dict[str, str]) -> None:
    """替换段落中跨 run 的占位符。"""
    # 合并所有 run 的文本
    full_text = "".join(r.text for r in p.runs)
    for placeholder, new_val in replacements.items():
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, new_val or "")
    # 全部写入第一个 run，清空其余
    if p.runs:
        p.runs[0].text = full_text
        for r in p.runs[1:]:
            r.text = ""


def generate_offer_docx(**kwargs) -> BytesIO:
    """填充 DOCX 模板，返回 BytesIO。"""
    template_path = find_hr_template("offer_template.docx")
    doc = Document(str(template_path))

    today = date.today()
    send_date = kwargs.get("send_date") or f"{today.year}年{today.month:02d}月{today.day:02d}日"

    replacements: dict[str, str] = {}
    for placeholder, key in PLACEHOLDER_MAP.items():
        if key == "send_date":
            replacements[placeholder] = send_date
        else:
            replacements[placeholder] = str(kwargs.get(key, "") or "")

    # 替换段落中的占位符
    for p in doc.paragraphs:
        _replace_in_paragraph(p, replacements)

    # 公章压在落款公司名上：同一段落内，图片先，文字后，图片自动居左上，文字在下
    from docx.enum.text import WD_ALIGN_PARAGRAPH as WD_AP
    from docx.shared import Cm as Cm_
    from docx.shared import Pt as Pt_
    stamp_path = find_hr_template("company_stamp.png")
    for i, p in enumerate(doc.paragraphs):
        if "丽珠集团福州福兴医药有限公司" in p.text and i > 0:
            para = doc.paragraphs[i]
            company_text = para.text
            para.clear()
            para.alignment = WD_AP.RIGHT
            # 公章 run
            stamp_run = para.add_run()
            stamp_run.add_picture(str(stamp_path), width=Cm_(5))
            # 换行 + 公司名
            name_run = para.add_run()
            name_run.text = "\n" + company_text
            name_run.font.size = Pt_(12)
            break

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf



def _find_cjk_font() -> str:
    import os
    for p in ["/System/Library/Fonts/Supplemental/Songti.ttc", "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc"]:
        if os.path.exists(p): return p
    return ""


def generate_offer_pdf(**kwargs) -> BytesIO:
    """fpdf2 纯 Python 生成 PDF，跨平台无系统依赖。"""
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    font_path = _find_cjk_font()
    if font_path:
        pdf.add_font("CJK", "", font_path, uni=True)
        pdf.add_font("CJK", "B", font_path, uni=True)
        fn = "CJK"
    else:
        fn = "Helvetica"
    today = date.today()
    send_date = kwargs.get("send_date") or f"{today.year}年{today.month:02d}月{today.day:02d}日"
    vals: dict[str, str] = {}
    for placeholder, key in PLACEHOLDER_MAP.items():
        if key == "send_date": vals[key] = send_date
        else: vals[key] = str(kwargs.get(key, "") or "")
    pdf.set_font(fn, "B", 18); pdf.cell(0, 14, "丽珠集团福州福兴医药有限公司", ln=True, align="C")
    pdf.set_font(fn, "", 12); pdf.cell(0, 10, "录用通知书", ln=True, align="C")
    pdf.ln(6)
    pdf.set_font(fn, "", 11)
    lines = [f"{vals["name"]} 先生/女士：", "", f"我们很高兴通知您，经公司研究决定，拟录用您担任我司 {vals["position"]} 岗位。", "", "现将入职相关事宜通知如下：", "", f"一、薪资待遇", f"  底薪：{vals["base_salary"] or "面议"}", f"  综合月薪范围：{vals["salary_range"] or "面议"}", "", f"二、报到安排", f"  体检日期：{vals["medical_date"] or "另行通知"}", f"  报到日期：{vals["report_date"]}", f"  岗位保留至：{vals["offer_expire_date"]}", "", "三、其他说明", "  1. 报到时请携带本人身份证、学历证书、资格证书原件及复印件。", "  2. 逾期未报到且未说明原因者，本录用通知自动失效。", "  3. 本通知仅限本人使用，不得转让。", "", "如您接受以上条件，请在3个工作日内回复确认。", "", "", f"丽珠集团福州福兴医药有限公司", f"发送日期：{send_date}"]
    for line in lines:
        pdf.set_font(fn, "B" if line.startswith("一、") or line.startswith("二、") or line.startswith("三、") or line.startswith("丽珠") else "", 11)
        pdf.cell(0, 8, line, ln=True, align="R" if line.startswith("丽珠") or line.startswith("发送") else "L")
    buf = BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf


def generate_offer_html(**kwargs) -> str:
    """生成 Offer HTML（用于前端预览和 Weasyprint 转 PDF）。"""
    today = date.today()
    send_date = kwargs.get("send_date") or f"{today.year}年{today.month:02d}月{today.day:02d}日"

    vals: dict[str, str] = {}
    for placeholder, key in PLACEHOLDER_MAP.items():
        if key == "send_date":
            vals[key] = send_date
        else:
            vals[key] = str(kwargs.get(key, "") or "")

    stamp_path = find_hr_template("company_stamp.png")
    logo_path = find_hr_template("company_logo.png")
    import base64
    stamp_b64 = base64.b64encode(stamp_path.read_bytes()).decode()
    logo_b64 = base64.b64encode(logo_path.read_bytes()).decode()

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<style>
  @page {{ size: A4; margin: 2cm 2.5cm 2cm 2.5cm; }}
  @media screen {{ body {{ max-width: 700px; margin: 30px auto; padding: 20px; }} }}
  body {{ font-family: "SimSun", "宋体", serif; font-size: 12pt; line-height: 1.8; color: #000; }}
  .center {{ text-align: center; }}
  .right {{ text-align: right; }}
  .indent {{ text-indent: 2em; }}
  .title {{ font-size: 18pt; font-weight: bold; margin: 12px 0 8px 0; }}
  .letterhead {{ margin: 0 0 16px 0; }}
  .letterhead-row {{ display: flex; justify-content: space-between; align-items: center; }}
  .letterhead-name {{ font-size: 12pt; font-weight: bold; margin: 0; text-align: left; }}
  .letterhead-logo {{ height: 50px; }}
  .letterhead-line {{ border: none; border-top: 1px solid #000; margin: 8px 0 0 0; }}
  .signature {{ text-align: right; margin-top: 24px; }}
  .signature-stamp {{ display: block; width: 6cm; margin-left: auto; }}
  .signature-text {{ margin-top: -3.5cm; font-size: 12pt; }}
</style>
</head>
<body>
<div class="letterhead">
  <div class="letterhead-row">
    <p class="letterhead-name">丽珠集团福州福兴医药有限公司</p>
    <img class="letterhead-logo" src="data:image/png;base64,{logo_b64}" alt="logo">
  </div>
  <hr class="letterhead-line">
</div>
<p class="right" style="font-size:9pt;color:#666;">HR-RE-006</p>
<p class="center title">录用通知函</p>
<p class="indent">{vals["name"]}同学：</p>
<p class="indent">恭喜您从众多应聘者中脱颖而出，成为丽珠大家庭的一员。根据您的经验、技能、学历和各方综合素质，并与您本人协商后，我公司对您的入职相关事宜做以下安排：</p>
<p>一、职&emsp;&emsp;位：{vals["position"]}</p>
<p>二、转正薪酬：底薪{vals["base_salary"]}元+其他津贴+绩效奖金，综合税前月薪{vals["salary_range"]}元</p>
<p>三、年收入=综合月薪*12+年终双薪+公司年度奖金+项目攻关奖励（根据参与的项目完成情况以及负责的职责由项目负责人进行分配）</p>
<p>四、福利待遇：</p>
<p>1、公司提供食宿，热水器、空调均有配套；</p>
<p>2、假期：法定节假日+公司规定的其他假期；</p>
<p>3、保险及公积金：按政府相关政策及公司的有关规定执行，公积金按照最高比例12%缴纳</p>
<p>4、试用期：根据《劳动合同法》的相关规定，您的试用期为&emsp;叁&emsp;个月（根据工作表现和工作业绩可适当缩短试用期）。首次期间签订&emsp;叁年&emsp;的劳动合同及签订关键岗位保密协议。</p>
<p>5、请您{vals["medical_date"]}前至我公司指定体检中心参加职业健康体检，体检合格后上岗：</p>
<p>6、报到时间：请于&emsp;{vals["report_date"]}前到公司人力资源部报到。</p>
<p>五、请您在报到时携带以下资料的原件及复印件：</p>
<p>身份证；学历证、学位证等其他相关证书；</p>
<p>个人一寸彩照2张；</p>
<p>中国银行卡（工资卡，可入职后办理）；</p>
<p>征信报告（任意银行APP可下载）</p>
<p>竭诚欢迎您的加入，相信以您的能力，必能在公司一展所长！</p>
<p>说明：1、公司将为你保留职位至{vals["offer_expire_date"]}，如您不能在此前在贵校的线上就业协议上应约，公司将视您为自动放弃本工作机会，本录用通知失效</p>
<p>2、公司对薪酬福利要求保密，请勿同其它任何第三方讨论薪酬福利事宜，谢谢！</p>
<div class="signature">
  <img class="signature-stamp" src="data:image/png;base64,{stamp_b64}" alt="公章">
  <p class="signature-text">丽珠集团福州福兴医药有限公司</p>
</div>
<p class="right">{send_date}</p>
<p style="margin-top:24px;">公司地址：福建省福州市福清市江阴工业集中区&emsp;联系人：王琳18650755207</p>
<p>邮箱：wanglin03@livzon.cn</p>
</body>
</html>"""
