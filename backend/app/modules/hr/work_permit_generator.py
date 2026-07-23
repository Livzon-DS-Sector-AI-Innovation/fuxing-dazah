"""Generate 上岗证 (work permit) from DOCX template."""

from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document

from app.modules.hr.models import Employee


def _find_template() -> Path:
    candidates = [
        Path("assets/hr/上岗证模板.docx"),
        Path("../assets/hr/上岗证模板.docx"),
        Path(__file__).resolve().parent.parent.parent.parent
        / "assets/hr"
        / "上岗证模板.docx",
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError("模板文件未找到: 上岗证模板.docx")


def _replace_all(doc: Document, placeholders: dict[str, str]) -> None:
    ordered = sorted(placeholders.items(), key=lambda kv: -len(kv[0]))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, val in ordered:
                        if key in para.text:
                            para.text = para.text.replace(key, val)


def generate_work_permit(
    employee: Employee,
    training_items: list[dict[str, Any]] | None = None,
) -> BytesIO:
    """生成上岗证 Word 文档。

    training_items 每项包含: file_name, trainer, plan_date
    - 培训起止日期: 从 plan_date 中取最早和最晚
    - 培训清单内容: 所有 file_name 逗号拼接
    - 培训师: 所有 trainer 去重后逗号拼接
    """
    doc = Document(str(_find_template()))

    # ── 计算培训计划的日期和内容 ──
    items = training_items or []
    dates = [it.get("plan_date") for it in items if it.get("plan_date")]
    content_names = [it.get("content") or it.get("file_name") or "" for it in items]
    trainer_names: list[str] = []
    for it in items:
        t = (it.get("trainer") or "").strip()
        if t and t not in trainer_names:
            trainer_names.append(t)

    date_range = ""
    if dates:
        sorted_dates = sorted(dates)
        date_range = f"{sorted_dates[0]} ~ {sorted_dates[-1]}"

    content_list = "、".join(filter(None, content_names))
    trainer_list = "、".join(trainer_names)

    _replace_all(doc, {
        "{姓名}": employee.name or "",
        "{性别}": employee.gender or "",
        "{学历}": employee.education or "",
        "{体现部门}": employee.department or "",
        "{培训起止日期}": date_range,
        "{培训清单内容}": content_list,
        "{培训师}": trainer_list,
    })

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
