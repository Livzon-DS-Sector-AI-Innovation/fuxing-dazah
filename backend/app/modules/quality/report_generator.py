"""报告单模板填充引擎。

读取 .docx 模板中的 {{占位符}}，根据数据源映射表填充值，
生成最终检验报告单。

格式说明（来自爬数软件模板约定）：
  {{字段名}}                        → 直接替换为文本
  {{字段名| dec(N)}}                → 保留 N 位小数
  {{字段名| dec(N, '单位')}}        → 小数 + 单位后缀
  {{字段名| dec(N, '单位', 阈值)}}   → 小于阈值时显示"未检出"
"""

from __future__ import annotations

import re
from copy import deepcopy
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document


# ─── 占位符解析 ───


@dataclass
class PlaceholderSpec:
    """占位符规格说明"""

    raw: str  # 原始占位符文本，如 "水分 | dec(1, '%')"
    name: str  # 字段名
    decimals: int = 2  # 小数位数
    suffix: str = ""  # 单位后缀
    threshold: float | None = None  # 低于此值显示"未检出"


def parse_placeholder(raw: str) -> PlaceholderSpec:
    """解析占位符字符串为结构化规格。

    Examples:
        "水分 | dec(1, '%')" → decimals=1, suffix='%'
        "杂质A| dec(2, '%',0.1)" → decimals=2, suffix='%', threshold=0.1
        "批号" → decimals=0, suffix=''
    """
    parts = [p.strip() for p in raw.split("|")]
    name = parts[0].strip()

    spec = PlaceholderSpec(raw=raw, name=name)

    if len(parts) < 2:
        return spec

    fmt = parts[1]
    # 解析 dec(N, 'unit', threshold)
    # 匹配 dec(数字)
    dec_match = re.search(r"dec\s*\(\s*(\d+)\s*", fmt)
    if dec_match:
        spec.decimals = int(dec_match.group(1))

    # 匹配单引号/中文引号中的单位
    unit_match = re.search(r"['‘’]([^'‘’]*)['‘’]", fmt)
    if unit_match:
        spec.suffix = unit_match.group(1)

    # 匹配阈值（最后一个数字参数）
    nums = re.findall(r"(\d+\.?\d*)", fmt)
    if len(nums) >= 2:
        spec.threshold = float(nums[-1])

    return spec


def format_value(value: Any, spec: PlaceholderSpec) -> str:
    """根据规格格式化值。"""
    if value is None or value == "":
        return ""

    try:
        num = float(value)
    except (ValueError, TypeError):
        return str(value)

    # 百分比后缀处理：若值为小数形式（<1.0），自动转换为百分比显示
    # 例如 0.956 → 95.6%，但 3.8（已是百分比）→ 3.8%
    is_percent = spec.suffix in ("%", "％")
    if is_percent and 0 < num < 1.0:
        num = num * 100
        if spec.threshold is not None:
            spec.threshold = spec.threshold * 100

    # 阈值检查：低于阈值显示"未检出"
    if spec.threshold is not None and spec.threshold > 0 and num < spec.threshold:
        return "未检出"

    # 小数格式化
    formatted = f"{num:.{spec.decimals}f}"

    if spec.suffix:
        formatted += spec.suffix

    return formatted


# ─── 模板填充 ───


@dataclass
class FillResult:
    """模板填充结果"""

    filled_count: int = 0
    unfilled: list[str] = field(default_factory=list)  # 未能填充的占位符
    file_bytes: BytesIO | None = None


class ReportFiller:
    """报告单模板填充器。

    支持段落和表格中的占位符替换。
    """

    def __init__(self, template_path: str | Path):
        self.template_path = Path(template_path)
        self.doc = Document(str(template_path))
        self._placeholder_pattern = re.compile(r"\{\{(.+?)\}\}")

    def extract_placeholders(self) -> list[PlaceholderSpec]:
        """提取模板中所有占位符。"""
        all_text = self._collect_all_text()
        raw_placeholders = self._placeholder_pattern.findall(all_text)
        # 去重保序
        seen = set()
        unique = []
        for p in raw_placeholders:
            if p not in seen:
                seen.add(p)
                unique.append(p)
        return [parse_placeholder(p) for p in unique]

    def fill(self, data: dict[str, Any]) -> FillResult:
        """用数据字典填充模板。

        Args:
            data: 字段名 → 值的映射字典

        Returns:
            FillResult: 填充结果（含未填充列表）
        """
        result = FillResult()
        filled_names = set()

        # 填充表格
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    filled = self._replace_in_cell(cell, data)
                    filled_names.update(filled)
                    result.filled_count += len(filled)

        # 填充段落
        for para in self.doc.paragraphs:
            filled = self._replace_in_paragraph(para, data)
            filled_names.update(filled)
            result.filled_count += len(filled)

        # 检查未填充的占位符
        all_specs = self.extract_placeholders()
        result.unfilled = [
            s.raw for s in all_specs if s.name not in filled_names and s.name not in data
        ]

        # 保存到 BytesIO
        output = BytesIO()
        self.doc.save(output)
        output.seek(0)
        result.file_bytes = output

        return result

    def _collect_all_text(self) -> str:
        """收集文档所有文本（用于提取占位符）。"""
        parts = []
        for para in self.doc.paragraphs:
            parts.append(para.text or "")
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text or "")
        return "\n".join(parts)

    def _replace_in_cell(self, cell, data: dict[str, Any]) -> set[str]:
        """在表格单元格中替换占位符。"""
        filled = set()
        for para in cell.paragraphs:
            filled.update(self._replace_in_paragraph(para, data))
        return filled

    def _replace_in_paragraph(self, para, data: dict[str, Any]) -> set[str]:
        """在段落中替换占位符。处理跨 run 的情况。"""
        # 收集段落完整文本
        full_text = para.text
        matches = list(self._placeholder_pattern.finditer(full_text))
        if not matches:
            return set()

        filled = set()

        # 逐个替换
        new_text = full_text
        for match in reversed(matches):  # 从后往前替换，保持位置
            raw = match.group(1)
            spec = parse_placeholder(raw)
            if spec.name in data:
                val = format_value(data[spec.name], spec)
                new_text = new_text[: match.start()] + val + new_text[match.end() :]
                filled.add(spec.name)

        # 将替换后的文本写回段落
        if filled:
            self._set_paragraph_text(para, new_text)

        return filled

    def _set_paragraph_text(self, para, text: str) -> None:
        """设置段落文本，保留第一个 run 的格式。"""
        if para.runs:
            # 保留第一个 run 的格式
            first_run = para.runs[0]
            # 清除其他 runs
            for run in para.runs[1:]:
                run.text = ""
            first_run.text = text
        else:
            para.add_run(text)


# ─── 便捷函数 ───


def fill_report(
    template_path: str | Path,
    data: dict[str, Any],
) -> FillResult:
    """填充报告单模板。"""
    filler = ReportFiller(template_path)
    return filler.fill(data)


def extract_template_placeholders(template_path: str | Path) -> list[PlaceholderSpec]:
    """提取模板占位符列表。"""
    filler = ReportFiller(template_path)
    return filler.extract_placeholders()
