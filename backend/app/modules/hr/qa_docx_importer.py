"""从完整培训记录 .docx 中提取问答记录表题目，供一键导入共享题库。"""

import re
from pathlib import Path
from typing import Any

from docx import Document


def extract_subject(doc: Document) -> str:
    """从第一个含"培训内容"字样的表格中提取主题。"""
    for table in doc.tables:
        for row in table.rows:
            texts = [c.text.strip() for c in row.cells]
            line = " ".join(texts)
            if "培训内容" in line:
                # 取后面非空的单元格内容作为主题
                for t in texts[1:]:
                    if t:
                        return re.sub(r"^Training\s*content", "", t).strip()
    return ""


def _is_qa_header(row_cells: list[Any]) -> bool:
    """判断是否为问答记录表的表头行（含 文件编号+考题+答案）。"""
    texts = [c.text.strip() for c in row_cells]
    joined = " ".join(texts)
    has_file = "文件编号" in joined or "Document No" in joined
    has_question = "考题" in joined or "Question" in joined
    has_answer = "答案" in joined or "Answer" in joined
    return has_file and (has_question or has_answer)


def _is_data_terminator(row_cells: list[Any]) -> bool:
    """判断是否为题目数据区的结束行（出现"姓名"或空白行）。"""
    texts = [c.text.strip() for c in row_cells]
    joined = " ".join(texts)
    return "姓名" in joined or ("Name" in joined and "考核" in joined) or all(t == "" for t in texts)


def extract_questions(doc: Document) -> list[dict[str, Any]]:
    """从 docx 的所有表格中提取问答题目。"""
    results: list[dict[str, Any]] = []

    for table in doc.tables:
        rows = list(table.rows)
        for ri, row in enumerate(rows):
            if not _is_qa_header(row.cells):
                continue
            # 找到表头后的数据行
            for data_row in rows[ri + 1:]:
                if _is_data_terminator(data_row.cells):
                    break
                texts = [c.text.strip() for c in data_row.cells]
                # 去重合并单元格导致的重复
                seen, cells = set(), []
                for t in texts:
                    if t not in seen:
                        seen.add(t)
                        cells.append(t)
                # 格式：序号, 文件编号, (空合并列), (空), 考题, (空), (空), 答案, 分数
                # 过滤掉序号（纯数字）和仅含空格/符号的残留
                meaningful = [t for t in cells if t and not t.isspace() and not re.fullmatch(r"\d+", t)]
                if len(meaningful) < 3:
                    continue
                # 如果 meaningful 的第 0 个是文件编号，取 file_no=meaningful[0]
                # 否则取 question, answer, score 尽量匹配
                file_no = meaningful[0] if len(meaningful) > 0 else ""
                question = meaningful[1] if len(meaningful) > 1 else ""
                answer = meaningful[2] if len(meaningful) > 2 else ""
                score_str = meaningful[3] if len(meaningful) > 3 else ""
                try:
                    score = int(score_str)
                except (ValueError, TypeError):
                    score = 10
                if question:
                    results.append({
                        "file_no": file_no,
                        "question": question,
                        "answer": answer,
                        "score": score,
                    })
            # 一张表只取一块问答区
            break

    return results


def parse_training_record(file_bytes: bytes, filename: str = "") -> tuple[str, list[dict[str, Any]]]:
    """解析完整培训记录 docx，返回 (培训主题, 题目列表)。"""
    from io import BytesIO

    doc = Document(BytesIO(file_bytes))
    subject = extract_subject(doc) or Path(filename).stem
    questions = extract_questions(doc)
    return subject, questions
