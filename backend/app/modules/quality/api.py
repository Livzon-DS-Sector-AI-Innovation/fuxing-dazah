"""Quality 模块 API 路由。"""

import re
import uuid
from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi import (
    APIRouter,
    Body,
    Depends,
    File,
    Form,
    HTTPException,
    Query,
    UploadFile,
)
from fastapi.responses import FileResponse, JSONResponse, Response, StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.database import get_db
from app.core.response import paginated_response, success_response
from app.modules.quality.report_generator import extract_template_placeholders
from app.modules.quality.repository import (
    create_report_record,
    create_standard_document,
    create_standard_item,
    delete_standard_document,
    delete_standard_item,
    get_standard_document,
    list_standard_documents,
    list_standard_items,
    update_standard_document,
    update_standard_item,
    delete_inspection_record,
    get_product_names,
    get_report_record,
    get_summary_by_product,
    list_inspection_records,
    list_report_records,
)
from app.modules.quality.schemas import (
    GenerateReportRequest,
    InspectionRecordListItem,
    StandardDocumentCreate,
    StandardDocumentUpdate,
    StandardItemCreate,
    StandardItemUpdate,
    UploadLcResponse,
)
from app.modules.quality.service import lc_report_service
from app.shared.module_registry import MODULES_BY_CODE

router = APIRouter()
_module = MODULES_BY_CODE["quality"]
REPORT_TEMPLATE_DIR = Path(__file__).resolve().parent.parent.parent.parent / "报告模板"
PLACEHOLDER_RE = re.compile(r"\{\{(.+?)\}\}")


@router.get("/", summary=f"{_module.name}模块信息")
async def read_module() -> dict[str, str]:
    return _module.as_dict()


# ─── 液相解析上传 ───


@router.post("/lc/upload", response_model=UploadLcResponse, summary="上传液相计算表并解析")
async def upload_lc_excel(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
):
    filename = file.filename or "unknown.xlsx"
    if not filename.lower().endswith((".xlsx", ".xls")):
        raise HTTPException(status_code=400, detail="仅支持 .xlsx 或 .xls")
    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(status_code=400, detail="文件为空")
    if len(file_bytes) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="文件超过 10MB")
    try:
        return await lc_report_service.parse_and_validate(file_bytes, filename, db=db)
    except ValueError as e:
        raise HTTPException(status_code=422, detail=f"无法解析：{e}") from e


# ─── 检验记录查询 ───


@router.get("/lc/records", summary="分页查询液相解析历史记录")
async def list_lc_records(
    product_name: str | None = Query(default=None, description="产品名称（模糊搜索）"),
    batch_number: str | None = Query(default=None, description="批号（模糊搜索）"),
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=20, ge=1, le=200),
    db: AsyncSession = Depends(get_db),
) -> JSONResponse:
    items, total = await list_inspection_records(
        db, product_name=product_name, batch_number=batch_number,
        page=page, page_size=page_size,
    )
    return paginated_response(
        data=[
            InspectionRecordListItem(
                id=it.id,
                product_name=it.product_name,
                batch_number=it.batch_number,
                form_id=it.form_id,
                standard_type=it.standard_type,
                all_pass=it.all_pass,
                excel_filename=it.excel_filename,
                created_at=it.created_at,
            ).model_dump(mode="json")
            for it in items
        ],
        page=page,
        page_size=page_size,
        total=total,
    )


@router.get("/lc/records/{record_id}", summary="查询单条检验记录详情（含杂质明细）")
async def get_lc_record(
    record_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
) -> JSONResponse:
    detail = await lc_report_service.get_record_detail(db, record_id)
    if not detail:
        raise HTTPException(status_code=404, detail="检验记录不存在")
    return success_response(data=detail.model_dump(mode="json"))


@router.delete("/lc/records/{record_id}", summary="删除检验记录（软删除）")
async def delete_lc_record(
    record_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
) -> JSONResponse:
    result = await delete_inspection_record(db, record_id)
    if not result:
        raise HTTPException(status_code=404, detail="检验记录不存在")
    return success_response(message="已删除")


# ─── 模板管理 ───


def _scan_templates(root: Path) -> list[dict]:
    items = []
    for p in sorted(root.iterdir()):
        if p.is_dir():
            items.append({"type": "folder", "name": p.name, "children": _scan_templates(p)})
        elif p.suffix == ".docx":
            phs = extract_template_placeholders(str(p))
            st = p.stat()
            items.append({
                "type": "template",
                "filename": p.name,
                "size_kb": round(st.st_size / 1024, 1),
                "modified": st.st_mtime,
                "placeholder_count": len(phs),
                "placeholders": [
                    {"name": ph.name, "decimals": ph.decimals, "suffix": ph.suffix}
                    for ph in phs
                ],
            })
    return items


@router.get("/templates", summary="列出模板")
async def list_templates():
    return _scan_templates(REPORT_TEMPLATE_DIR) if REPORT_TEMPLATE_DIR.exists() else []


@router.get("/templates/all-placeholders", summary="获取所有模板的所有占位符名称")
async def all_placeholders():
    names: set[str] = set()

    def walk(items: list[dict]):
        for item in items:
            if item["type"] == "folder":
                walk(item.get("children", []))
            else:
                for ph in item.get("placeholders", []):
                    names.add(ph["name"])

    walk(_scan_templates(REPORT_TEMPLATE_DIR) if REPORT_TEMPLATE_DIR.exists() else [])
    base = ["流水号", "批号", "规格", "生产日期", "批量_kg", "有效期_年"]
    rest = sorted(n for n in names if n not in base)
    return base + rest


@router.post("/templates/folders", summary="创建文件夹")
async def create_folder(name: str = Body(..., embed=True)):
    (REPORT_TEMPLATE_DIR / name).mkdir(parents=True, exist_ok=True)
    return {"name": name}


@router.delete("/templates/folders", summary="删除空文件夹")
async def delete_folder(name: str = Body(..., embed=True)):
    p = REPORT_TEMPLATE_DIR / name
    if not p.exists():
        raise HTTPException(status_code=404, detail="不存在")
    try:
        p.rmdir()
    except OSError:
        raise HTTPException(status_code=400, detail="文件夹不为空") from None
    return {"message": f"已删除 {name}"}


@router.post("/templates/upload", summary="上传模板")
async def upload_template(folder: str = Form(""), file: UploadFile = File(...)):
    if not file.filename or not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="仅支持 .docx")
    dest = REPORT_TEMPLATE_DIR / folder
    dest.mkdir(parents=True, exist_ok=True)
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="文件为空")
    if len(content) > 5 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="超过5MB")
    (dest / file.filename).write_bytes(content)
    return {"filename": file.filename, "folder": folder}


@router.get("/templates/{path:path}/download", summary="下载模板")
async def download_template(path: str):
    full = REPORT_TEMPLATE_DIR / path
    if not full.exists():
        raise HTTPException(status_code=404, detail="不存在")
    return FileResponse(
        str(full),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=full.name,
    )


@router.delete("/templates/{path:path}", summary="删除模板")
async def delete_template(path: str):
    full = REPORT_TEMPLATE_DIR / path
    if not full.exists():
        raise HTTPException(status_code=404, detail="不存在")
    full.unlink()
    return {"message": "已删除"}


# ─── 产品代码管理 ───

PRODUCTS_FILE = REPORT_TEMPLATE_DIR.parent / "products.json"


def _load_products() -> list[dict]:
    import json

    if PRODUCTS_FILE.exists():
        try:
            return json.loads(PRODUCTS_FILE.read_text())
        except Exception:
            pass
    return []


def _save_products(data: list[dict]):
    import json

    PRODUCTS_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2))


@router.get("/products", summary="列出产品代码映射")
async def list_products():
    return _load_products()


@router.post("/products", summary="保存产品代码映射")
async def save_products(data: list[dict] = Body(...)):
    _save_products(data)
    return {"message": "已保存", "count": len(data)}


# ─── 报告生成 ───


def _replace_placeholders(para, data: dict):
    full = para.text
    matches = list(PLACEHOLDER_RE.finditer(full))
    if not matches:
        return
    new_text = full
    for m in reversed(matches):
        raw = m.group(1).strip()
        key = raw.split("|")[0].strip()
        val = str(data.get(key, "-"))
        new_text = new_text[: m.start()] + val + new_text[m.end() :]
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""


@router.post("/report/generate", summary="生成报告单")
async def generate_report(
    payload: GenerateReportRequest = Body(...),
    db: AsyncSession = Depends(get_db),
):
    tp = REPORT_TEMPLATE_DIR / payload.template
    if not tp.exists():
        raise HTTPException(status_code=404, detail=f"模板不存在：{payload.template}")

    # 获取填充数据
    fill_data: dict = {}
    product_name = ""
    batch_number = ""
    record_id: uuid.UUID | None = None

    if payload.inspection_record_id:
        # 从数据库加载检验记录自动填充
        detail = await lc_report_service.get_record_detail(db, payload.inspection_record_id)
        if not detail:
            raise HTTPException(status_code=404, detail="检验记录不存在")
        fill_data = lc_report_service.build_report_data(detail.report)
        product_name = detail.product_name
        batch_number = detail.batch_number
        record_id = detail.id
    elif payload.data:
        fill_data = payload.data
        product_name = fill_data.get("产品名称", "")
        batch_number = fill_data.get("批号", "")
    else:
        raise HTTPException(status_code=400, detail="请提供 inspection_record_id 或 data")

    # 填充模板
    doc = Document(str(tp))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_placeholders(para, fill_data)
    for para in doc.paragraphs:
        _replace_placeholders(para, fill_data)

    # 保存到临时文件
    output_dir = REPORT_TEMPLATE_DIR / "_generated"
    output_dir.mkdir(parents=True, exist_ok=True)
    safe_batch = batch_number.replace("/", "_").replace("\\", "_")
    output_filename = f"COA-{product_name}-{safe_batch}.docx"
    output_path = output_dir / output_filename
    doc.save(str(output_path))
    file_size = output_path.stat().st_size

    # 保存报告单记录到数据库
    if record_id:
        await create_report_record(
            db=db,
            inspection_record_id=record_id,
            template_path=payload.template,
            product_name=product_name,
            batch_number=batch_number,
            file_path=str(output_path),
            file_size=file_size,
        )

    # 返回文件
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return StreamingResponse(
        out,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename={output_filename}"
        },
    )


# ─── 报告单记录 ───


@router.get("/report/records", summary="报告单历史列表")
async def list_reports(
    product_name: str | None = Query(default=None),
    batch_number: str | None = Query(default=None),
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=20, ge=1, le=200),
    db: AsyncSession = Depends(get_db),
) -> JSONResponse:
    items, total = await list_report_records(
        db, product_name=product_name, batch_number=batch_number,
        page=page, page_size=page_size,
    )
    return paginated_response(
        data=[
            {
                "id": str(it.id),
                "inspection_record_id": str(it.inspection_record_id),
                "template_path": it.template_path,
                "product_name": it.product_name,
                "batch_number": it.batch_number,
                "file_path": it.file_path,
                "file_size": it.file_size,
                "created_at": it.created_at.isoformat() if it.created_at else None,
            }
            for it in items
        ],
        page=page,
        page_size=page_size,
        total=total,
    )


@router.get("/report/records/{report_id}/download", summary="下载已生成的报告单")
async def download_report(report_id: uuid.UUID, db: AsyncSession = Depends(get_db)):
    report = await get_report_record(db, report_id)
    if not report or not report.file_path:
        raise HTTPException(status_code=404, detail="报告文件不存在")
    fp = Path(report.file_path)
    if not fp.exists():
        raise HTTPException(status_code=404, detail="报告文件已被清理")
    return FileResponse(
        str(fp),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=fp.name,
    )


# ─── 汇总表 ───


@router.get("/summary/batch/{record_id}", summary="单批次汇总")
async def batch_summary(
    record_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
) -> JSONResponse:
    detail = await lc_report_service.get_record_detail(db, record_id)
    if not detail:
        raise HTTPException(status_code=404, detail="检验记录不存在")

    # 生成文字化判定摘要
    lines = [
        f"产品：{detail.product_name}",
        f"批号：{detail.batch_number}",
        f"标准：{detail.standard_type or '-'}",
        f"总体判定：{'合格' if detail.all_pass else '不合格'}",
        f"超趋势：{'是' if detail.has_oot else '否'}",
        f"杂质项目数：{len(detail.impurities)}",
    ]
    if detail.report.vancomycin_b:
        vb = detail.report.vancomycin_b
        lines.append(
            f"万古霉素B：{vb.rounded_first} (判定：{'合格' if vb.is_pass else '不合格'})"
        )
    if detail.report.total_impurities:
        ti = detail.report.total_impurities
        lines.append(
            f"总杂质：{ti.rounded_first} (判定：{'合格' if ti.is_pass else '不合格'})"
        )

    return success_response(
        data={
            "record": detail.model_dump(mode="json"),
            "summary_text": "\n".join(lines),
        }
    )


@router.get("/summary/history", summary="多批次历史汇总")
async def history_summary(
    product_name: str | None = Query(default=None, description="产品名称"),
    date_from: str | None = Query(default=None, description="起始日期 ISO 格式"),
    date_to: str | None = Query(default=None, description="结束日期 ISO 格式"),
    db: AsyncSession = Depends(get_db),
) -> JSONResponse:
    dt_from = datetime.fromisoformat(date_from) if date_from else None
    dt_to = datetime.fromisoformat(date_to) if date_to else None

    summary = await get_summary_by_product(
        db, product_name=product_name, date_from=dt_from, date_to=dt_to
    )
    return success_response(data=summary)


@router.get("/summary/products", summary="已检验产品列表")
async def list_summary_products(
    db: AsyncSession = Depends(get_db),
) -> JSONResponse:
    names = await get_product_names(db)
    return success_response(data=names)


# ─── 质量标准文档 / 项目行（SOP 号为匹配键）───


@router.post("/standards/import-doc", summary="上传标准文档解析导入")
async def import_standard_doc(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
) -> JSONResponse:
    """上传 .doc/.docx 质量标准文件：解析文件头+项目行，纯文字标准自动过滤。

    返回解析草稿与落库结果，前端可继续在明细表格中校正。
    """
    from app.modules.quality.standard_doc_parser import extract_text_async, parse_standard_doc

    filename = file.filename or "standard.doc"
    if not filename.lower().endswith((".doc", ".docx")):
        raise HTTPException(status_code=400, detail="仅支持 .doc / .docx 格式")
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="文件为空")
    try:
        text = await extract_text_async(content, filename)
    except Exception:
        raise HTTPException(status_code=500, detail="文档文本提取失败")
    if not text.strip():
        raise HTTPException(status_code=400, detail="未能从文档中提取文本")

    parsed = parse_standard_doc(text)
    if not parsed.product_name and not parsed.file_no:
        raise HTTPException(status_code=400, detail="解析失败：未识别出标准文件头信息")

    doc = await create_standard_document(db, {
        "file_no": parsed.file_no or filename,
        "product_name": parsed.product_name or filename,
        "product_code": parsed.product_code,
        "product_internal_code": parsed.product_internal_code,
        "specification": parsed.specification,
        "valid_years": parsed.valid_years,
        "effective_date": parsed.effective_date,
        "version": parsed.version,
    })
    created = 0
    for it in parsed.items:
        try:
            await create_standard_item(db, doc.id, {
                "seq": it.seq,
                "category": it.category,
                "item_name": it.item_name,
                "sop_no": it.sop_no,
                "standard_text": it.standard_text,
                "operator": it.operator,
                "limit_min": it.limit_min,
                "limit_max": it.limit_max,
                "method_source": it.method_source,
                "remark": it.remark,
            })
            created += 1
        except Exception:
            continue
    return success_response(
        data={"id": str(doc.id), "created_items": created, "parsed_items": len(parsed.items)},
        message=f"解析导入完成：{created}/{len(parsed.items)} 条标准行（纯文字标准已过滤）",
        status_code=201,
    )


@router.get("/standards/documents", summary="质量标准文档列表")
async def list_standard_docs(
    product_name: str | None = Query(default=None),
    db: AsyncSession = Depends(get_db),
) -> JSONResponse:
    docs = await list_standard_documents(db, product_name=product_name)
    return success_response(data=[
        {
            "id": str(d.id),
            "file_no": d.file_no,
            "product_name": d.product_name,
            "product_code": d.product_code,
            "product_internal_code": d.product_internal_code,
            "specification": d.specification,
            "valid_years": d.valid_years,
            "effective_date": d.effective_date,
            "version": d.version,
            "created_at": d.created_at.isoformat() if d.created_at else None,
        }
        for d in docs
    ])


@router.post("/standards/documents", summary="创建质量标准文档")
async def create_standard_doc(
    payload: StandardDocumentCreate = Body(...),
    db: AsyncSession = Depends(get_db),
) -> JSONResponse:
    doc = await create_standard_document(db, payload.model_dump())
    return success_response(data={"id": str(doc.id)}, message="标准文档创建成功", status_code=201)


@router.put("/standards/documents/{doc_id}", summary="更新质量标准文档")
async def update_standard_doc(
    doc_id: uuid.UUID,
    payload: StandardDocumentUpdate = Body(...),
    db: AsyncSession = Depends(get_db),
) -> JSONResponse:
    updates = {k: v for k, v in payload.model_dump().items() if v is not None}
    doc = await update_standard_document(db, doc_id, **updates)
    if not doc:
        raise HTTPException(status_code=404, detail="标准文档不存在")
    return success_response(message="已更新")


@router.delete("/standards/documents/{doc_id}", summary="删除质量标准文档")
async def delete_standard_doc(
    doc_id: uuid.UUID, db: AsyncSession = Depends(get_db),
) -> JSONResponse:
    doc = await delete_standard_document(db, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="标准文档不存在")
    return success_response(message="已删除")


@router.get("/standards/documents/{doc_id}/items", summary="标准项目行列表")
async def list_standard_doc_items(
    doc_id: uuid.UUID, db: AsyncSession = Depends(get_db),
) -> JSONResponse:
    items = await list_standard_items(db, doc_id)
    return success_response(data=[
        {
            "id": str(it.id),
            "seq": it.seq,
            "category": it.category,
            "item_name": it.item_name,
            "sop_no": it.sop_no,
            "standard_text": it.standard_text,
            "operator": it.operator,
            "limit_min": it.limit_min,
            "limit_max": it.limit_max,
            "method_source": it.method_source,
            "remark": it.remark,
        }
        for it in items
    ])


@router.post("/standards/documents/{doc_id}/items", summary="新增标准项目行")
async def create_standard_doc_item(
    doc_id: uuid.UUID,
    payload: StandardItemCreate = Body(...),
    db: AsyncSession = Depends(get_db),
) -> JSONResponse:
    it = await create_standard_item(db, doc_id, payload.model_dump())
    return success_response(data={"id": str(it.id)}, message="已添加", status_code=201)


@router.put("/standards/items/{item_id}", summary="更新标准项目行")
async def update_standard_doc_item(
    item_id: uuid.UUID,
    payload: StandardItemUpdate = Body(...),
    db: AsyncSession = Depends(get_db),
) -> JSONResponse:
    updates = {k: v for k, v in payload.model_dump().items() if v is not None}
    it = await update_standard_item(db, item_id, **updates)
    if not it:
        raise HTTPException(status_code=404, detail="标准行不存在")
    return success_response(message="已更新")


@router.delete("/standards/items/{item_id}", summary="删除标准项目行")
async def delete_standard_doc_item(
    item_id: uuid.UUID, db: AsyncSession = Depends(get_db),
) -> JSONResponse:
    it = await delete_standard_item(db, item_id)
    if not it:
        raise HTTPException(status_code=404, detail="标准行不存在")
    return success_response(message="已删除")
