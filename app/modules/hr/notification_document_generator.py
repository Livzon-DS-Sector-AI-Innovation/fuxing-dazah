"""培训通知 Word 文档生成器 — 基于表格模板."""

from datetime import date, datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from pydantic import BaseModel


class TrainingNotificationInput(BaseModel):
    department: str
    training_date: date
    subject: str
    training_time_start: str | None = None
    training_time_end: str | None = None
    location: str | None = None
    trainer: str | None = None
    training_method: str | None = None
    assessment_method: str | None = None
    content: str | None = None
    trainee_names: list[str] = []
    issuer_department: str | None = None
    issue_date: date | None = None


def _find_template() -> Path:
    candidates = [
        Path("员工培训教育管理规程/7.4培训通知书.docx"),
        Path("../员工培训教育管理规程/7.4培训通知书.docx"),
        Path(__file__).resolve().parent.parent.parent.parent
        / "员工培训教育管理规程"
        / "7.4培训通知书.docx",
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError("模板文件未找到: 7.4培训通知书.docx")


def _set_cell(cell, text: str) -> None:
    """Set cell text, preserving first run's formatting."""
    first = None
    for p in cell.paragraphs:
        for r in p.runs:
            if first is None:
                first = r
            r.text = ""
    if first is not None:
        first.text = str(text or "")
    elif cell.paragraphs:
        cell.paragraphs[0].add_run(str(text or ""))


def _compute_hours(start: str | None, end: str | None) -> str:
    if not start or not end:
        return ""
    try:
        s = datetime.strptime(start, "%H:%M")
        e = datetime.strptime(end, "%H:%M")
        diff = (e - s).total_seconds() / 3600
        if diff <= 0:
            return ""
        rounded = round(diff * 2) / 2
        if rounded == int(rounded):
            return f"{int(rounded)}小时"
        return f"{rounded}小时"
    except ValueError:
        return ""


def generate_training_notification(data: TrainingNotificationInput) -> BytesIO:
    """Generate training notification docx from template."""
    template_path = _find_template()
    doc = Document(str(template_path))
    table = doc.tables[0]

    # ── Row 0: 培训内容 (label col0, value cols 1-3 merged) ──
    topic_parts = [data.subject]
    if data.content:
        topic_parts.append(data.content)
    _set_cell(table.rows[0].cells[1], " — ".join(topic_parts))

    # ── Row 1: 培训日期 | value | 课时 | value ──
    _set_cell(table.rows[1].cells[1], str(data.training_date) if data.training_date else "")
    _set_cell(table.rows[1].cells[3], _compute_hours(data.training_time_start, data.training_time_end))

    # ── Row 2: 培训方式 | value | 授课人 | value ──
    _set_cell(table.rows[2].cells[1], data.training_method or "")
    _set_cell(table.rows[2].cells[3], data.trainer or "")

    # ── Row 3: 培训对象 (merged cols 1-3) ──
    people = "、".join(data.trainee_names) if data.trainee_names else ""
    _set_cell(table.rows[3].cells[1], people)

    # ── Row 4: 培训地点 (merged cols 1-3) ──
    _set_cell(table.rows[4].cells[1], data.location or "")

    # ── Row 5: 考核方式 (merged cols 1-3) ──
    _set_cell(table.rows[5].cells[1], data.assessment_method or "")

    # ── Row 6: 注意事项 (merged cols 1-3) ──
    _set_cell(table.rows[6].cells[1],
        "1. 请培训人员自带笔记本、笔，做好笔记。\n"
        "2. 请部门安排好参训人员的工作时间，做到培训工作两不误。\n"
        "3. 不得无故缺席、迟到，到场签到，有特殊情况须提前请假。"
    )

    # ── Para 0: department line ──
    for p_idx in (0, 1):
        if p_idx < len(doc.paragraphs):
            p = doc.paragraphs[p_idx]
            # Clear all runs first
            for r in p.runs:
                r.text = ""
    if len(doc.paragraphs) > 0:
        dept_name = data.issuer_department or data.department or ""
        doc.paragraphs[0].runs[0].text = f"部门/Dept：{dept_name}          签发人/ Issued by："
    if len(doc.paragraphs) > 1:
        d = data.training_date
        doc.paragraphs[1].runs[0].text = f"{d.year}年{d.month:02d}月{d.day:02d}日"

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
