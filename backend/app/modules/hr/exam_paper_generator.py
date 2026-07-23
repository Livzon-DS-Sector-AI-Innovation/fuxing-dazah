"""笔试试卷生成器 — 基于 assets/hr/试卷模板.docx。"""

from datetime import date
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Pt

TEMPLATE = Path(__file__).parent.parent.parent.parent / "assets/hr/试卷模板.docx"


def _find_template() -> Path:
    for p in [TEMPLATE, Path("assets/hr/试卷模板.docx"), Path("../assets/hr/试卷模板.docx")]:
        if p.exists(): return p
    raise FileNotFoundError("试卷模板.docx 未找到")


def _replace_para_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for r in paragraph.runs[1:]: r.text = ""


def _add_hr(doc) -> None:
    """添加水平横线。"""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_question(doc, no: int, q: dict) -> None:
    q_text = q.get("question") or q.get("content") or ""
    score = q.get("score", "")
    p = doc.add_paragraph()
    run = p.add_run(f"{no}. {q_text}  （{score}分）")
    run.font.size = Pt(12)

    opts = q.get("options") or []
    for opt in opts:
        if isinstance(opt, dict):
            doc.add_paragraph(f"    {opt.get('label', '')}. {opt.get('text', opt)}")
        else:
            doc.add_paragraph(f"    {opt}")

    if q.get("type") == "true_false":
        doc.add_paragraph("    对 □    错 □")
    if q.get("type") == "fill_blank":
        doc.add_paragraph("    ____________________")
    doc.add_paragraph("")


def generate_exam_paper_docx(
    *, subject: str, department: str | None = None,
    training_date: date | None = None, training_method: str | None = None,
    questions: list[dict] | None = None, full_score: int = 100, pass_line: int = 60,
) -> BytesIO:
    doc = Document(str(_find_template()))

    # 替换标题占位符
    for p in doc.paragraphs:
        if "{{培训内容}}" in (p.text or ""):
            _replace_para_text(p, f"{subject or ''}试题")
            break

    # 添加题目
    questions = questions or []
    choice_qs = [q for q in questions if q.get("type") == "choice"]
    tf_qs = [q for q in questions if q.get("type") == "true_false"]
    multi_qs = [q for q in questions if q.get("type") == "multi_choice"]
    fill_qs = [q for q in questions if q.get("type") == "fill_blank"]

    no = 1
    for heading, qs in [("一、单选题", choice_qs), ("二、判断题", tf_qs), ("三、多选题", multi_qs), ("四、填空题", fill_qs)]:
        if not qs: continue
        p = doc.add_paragraph(); run = p.add_run(f"{heading}（共{len(qs)}题）"); run.bold = True; run.font.size = Pt(14)
        doc.add_paragraph("")
        for q in qs:
            _add_question(doc, no, q)
            no += 1

    # 参考答案
    doc.add_page_break()
    p = doc.add_paragraph(); run = p.add_run("参考答案"); run.bold = True; run.font.size = Pt(16)
    doc.add_paragraph("")
    no = 1
    for heading, qs in [("一、单选题", choice_qs), ("二、判断题", tf_qs), ("三、多选题", multi_qs), ("四、填空题", fill_qs)]:
        if not qs: continue
        p = doc.add_paragraph(); run = p.add_run(heading); run.bold = True; run.font.size = Pt(12)
        for q in qs:
            ans = q.get("answer", "")
            if isinstance(ans, list): ans = "、".join(str(a) for a in ans)
            elif ans is True: ans = "对"
            elif ans is False: ans = "错"
            doc.add_paragraph(f"{no}. {ans}")
            no += 1
        doc.add_paragraph("")

    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf
