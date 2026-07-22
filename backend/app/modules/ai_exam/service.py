"""AI exam generation service. Supports both API and local mode."""

import json
import logging
import re
from io import BytesIO

import httpx
from docx import Document as DocxDocument

from app.core.config import get_settings

logger = logging.getLogger(__name__)


def _extract_docx_content(file_bytes: bytes) -> dict:
    doc = DocxDocument(BytesIO(file_bytes))
    all_paragraphs: list[str] = []
    bold_texts: list[str] = []
    for para in doc.paragraphs:
        parts: list[str] = []
        for run in para.runs:
            text = run.text.strip()
            if not text:
                continue
            parts.append(text)
            if run.bold:
                bold_texts.append(text)
        if parts:
            all_paragraphs.append(" ".join(parts))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        text = run.text.strip()
                        if text and run.bold:
                            bold_texts.append(text)
    return {"full_text": "\n".join(all_paragraphs), "bold_texts": bold_texts}


def _extract_text_content(file_bytes: bytes) -> dict:
    text = file_bytes.decode("utf-8", errors="ignore")
    return {"full_text": text, "bold_texts": []}


def _parse_file(file_bytes: bytes, filename: str) -> dict:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext in ("docx", "doc"):
        return _extract_docx_content(file_bytes)
    return _extract_text_content(file_bytes)


def _generate_local(content: dict, config: dict | None = None) -> dict:
    """离线模式：基于规则从文本生成多题型试卷。加粗文字优先挖空出题。"""
    choice_count = config.get("choice_count", 5) if config else 5
    tf_count = config.get("true_false_count", 5) if config else 5
    multi_count = config.get("multi_choice_count", 0) if config else 0
    fill_count = config.get("fill_blank_count", 0) if config else 0

    sentences = [s.strip() for s in content["full_text"].replace("\n", "。").split("。") if len(s.strip()) > 8]
    bold_texts = content.get("bold_texts", [])

    choice_qs = []
    tf_qs = []
    multi_qs = []
    fill_qs = []

    # ── 先对句子打分 ──
    keywords = ["必须", "应当", "禁止", "要求", "包括", "负责", "管理", "检查", "培训", "安全", "操作", "设备"]
    scored_sentences = []
    for s in sentences:
        score = sum(1 for kw in keywords if kw in s)
        if score > 0:
            scored_sentences.append((score, s))
    scored_sentences.sort(key=lambda x: -x[0])

    # ── 填空题：先从加粗文字挖空，不够再从关键句中提取 ──
    for i, bold in enumerate(bold_texts):
        fill_qs.append({"number": len(fill_qs) + 1, "question": f"请填写关键知识点：{bold}", "answer": bold})
    for _, s in scored_sentences:
        if len(fill_qs) >= fill_count:
            break
        for kw in keywords:
            if kw in s and not any(q["question"].find(kw) >= 0 for q in fill_qs):
                q_text = s.replace(kw, "______", 1)
                fill_qs.append({"number": len(fill_qs) + 1, "question": q_text, "answer": kw})
                break

    # ── 判断题：加粗文字生成正误判断 ──
    for i, bold in enumerate(bold_texts[:tf_count]):
        is_correct = i % 2 == 0  # 一半正确一半错误
        if is_correct:
            tf_qs.append({"number": i + 1, "question": f"「{bold}」这个说法是否正确？", "answer": "正确"})
        else:
            # 故意改错
            wrong = bold.replace("必须", "不必").replace("禁止", "允许").replace("应", "不应")
            if wrong == bold:
                wrong = "不" + bold
            tf_qs.append({"number": i + 1, "question": f"「{wrong}」这个说法是否正确？", "answer": "错误"})

    # 补判断题
    for s in sentences:
        if len(tf_qs) >= tf_count:
            break
        if s not in bold_texts and len(s) > 12:
            tf_qs.append({"number": len(tf_qs) + 1, "question": s + "？", "answer": "正确"})

    # ── 单选题：从高分句子中提取核心词挖空 ──
    for i, (_, s) in enumerate(scored_sentences[:choice_count]):
        # 找句子中的核心关键词挖空
        target_kw = None
        for kw in keywords:
            if kw in s:
                target_kw = kw
                break
        if target_kw:
            q = s.replace(target_kw, "____", 1)
            choice_qs.append({
                "number": i + 1,
                "question": q,
                "options": [
                    {"label": "A", "text": target_kw},
                    {"label": "B", "text": "不需要" + target_kw},
                    {"label": "C", "text": "视情况" + target_kw},
                    {"label": "D", "text": "由领导决定"},
                ],
                "answer": "A",
            })

    # ── 多选题：选包含多个关键点的句子 ──
    multi_candidates = [(sum(1 for kw in keywords if kw in s), s) for s in sentences if sum(1 for kw in keywords if kw in s) >= 2]
    multi_candidates.sort(key=lambda x: -x[0])
    for i, (_, s) in enumerate(multi_candidates[:multi_count]):
        # 提取2-3个关键词作为正确答案
        found_kws = [kw for kw in keywords if kw in s][:3]
        options = [{"label": chr(65 + j), "text": kw} for j, kw in enumerate(found_kws)]
        answer = ",".join(chr(65 + j) for j in range(len(found_kws)))
        # 加一个干扰项
        options.append({"label": chr(65 + len(found_kws)), "text": "以上都不对"})
        multi_qs.append({
            "number": i + 1,
            "question": s,
            "options": options,
            "answer": answer,
        })

    result = {
        "choice_questions": choice_qs[:choice_count],
        "true_false_questions": tf_qs[:tf_count],
        "multi_choice_questions": multi_qs[:multi_count],
        "fill_blank_questions": fill_qs[:fill_count],
    }
    logger.info(
        "_generate_local: bold=%d sentences=%d -> choice=%d tf=%d multi=%d fill=%d",
        len(bold_texts), len(sentences),
        len(result["choice_questions"]), len(result["true_false_questions"]),
        len(result["multi_choice_questions"]), len(result["fill_blank_questions"]),
    )
    return result


def _build_prompt(full_text: str, bold_texts: list[str], config: dict | None = None) -> str:
    if config:
        choice_count = config.get("choice_count", 5)
        tf_count = config.get("true_false_count", 5)
        multi_count = config.get("multi_choice_count", 0)
        fill_count = config.get("fill_blank_count", 0)
    else:
        choice_count, tf_count, multi_count, fill_count = 5, 5, 0, 0

    bold_hint = ""
    if bold_texts:
        bold_hint = "\n\n【重点关注】以下内容在原文件中被加粗标记，请优先作为考点出题：\n" + "\n".join(f"  - {t}" for t in bold_texts[:20])

    reqs = []
    if choice_count > 0:
        reqs.append(f"出 {choice_count} 道单选题（4 选项 A/B/C/D，只有一个正确答案）")
    if tf_count > 0:
        reqs.append(f"出 {tf_count} 道判断题（正确/错误）")
    if multi_count > 0:
        reqs.append(f"出 {multi_count} 道多选题（4 选项，有 2-3 个正确答案）")
    if fill_count > 0:
        reqs.append(f"出 {fill_count} 道填空题（将关键知识点挖空）")
    reqs.append("题目应覆盖材料关键知识点，加粗内容优先")
    reqs.append("严格按 JSON 格式输出")

    json_parts = []
    if choice_count > 0:
        json_parts.append('"choice_questions": [{"question":"...","options":[{"label":"A","text":"..."},...],"answer":"A"}]')
    if tf_count > 0:
        json_parts.append('"true_false_questions": [{"question":"...","answer":"正确"}]')
    if multi_count > 0:
        json_parts.append('"multi_choice_questions": [{"question":"...","options":[{"label":"A","text":"..."},...],"answer":"A,B"}]')
    if fill_count > 0:
        json_parts.append('"fill_blank_questions": [{"question":"...","answer":"关键词"}]')

    return f"""你是培训考核出题老师。根据以下材料生成试卷。

要求：{'; '.join(reqs)}

输出格式：{{ {', '.join(json_parts)} }}

材料：
{full_text[:8000]}{bold_hint}"""


async def generate_exam(file_bytes: bytes, filename: str, config: dict | None = None) -> dict:
    """Generate exam questions using local offline mode (no external API dependency)."""
    content = _parse_file(file_bytes, filename)
    if not content["full_text"].strip():
        raise ValueError("文件中未检测到文本内容")

    result = _generate_local(content, config)

    for key in ("choice_questions", "true_false_questions", "multi_choice_questions", "fill_blank_questions"):
        for i, q in enumerate(result.get(key, [])):
            q["number"] = i + 1

    return result


def export_exam(data: dict) -> BytesIO:
    """Export exam questions as a Word document."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)
    doc.add_heading(data.get("title", "培训考试试卷"), level=1)
    doc.add_paragraph("")

    info = []
    if data.get("examiner"):
        info.append(f"出卷人：{data['examiner']}")
    if data.get("exam_date"):
        info.append(f"考试日期：{data['exam_date']}")
    if data.get("assessment_date"):
        info.append(f"评估日期：{data['assessment_date']}")
    if info:
        doc.add_paragraph("  |  ".join(info))
        doc.add_paragraph("")

    choice_qs = data.get("choice_questions", [])
    if choice_qs:
        doc.add_heading("一、单选题", level=2)
        for q in choice_qs:
            doc.add_paragraph(f"{q['number']}. {q['question']}")
            for opt in q.get("options", []):
                doc.add_paragraph(f"    {opt['label']}. {opt['text']}")
            doc.add_paragraph("")

    tf_qs = data.get("true_false_questions", [])
    if tf_qs:
        doc.add_heading("二、判断题", level=2)
        for q in tf_qs:
            doc.add_paragraph(f"{q['number']}. {q['question']} （  ）")
            doc.add_paragraph("")

    multi_qs = data.get("multi_choice_questions", [])
    if multi_qs:
        doc.add_heading("三、多选题", level=2)
        for q in multi_qs:
            doc.add_paragraph(f"{q['number']}. {q['question']}")
            for opt in q.get("options", []):
                doc.add_paragraph(f"    {opt['label']}. {opt['text']}")
            doc.add_paragraph("")

    fill_qs = data.get("fill_blank_questions", [])
    if fill_qs:
        doc.add_heading("四、填空题", level=2)
        for q in fill_qs:
            doc.add_paragraph(f"{q['number']}. {q['question']}")
            doc.add_paragraph("")

    doc.add_heading("参考答案", level=2)
    if choice_qs:
        doc.add_paragraph("单选题答案：")
        for q in choice_qs:
            doc.add_paragraph(f"  {q['number']}. {q.get('answer', '?')}")
    if tf_qs:
        doc.add_paragraph("判断题答案：")
        for q in tf_qs:
            doc.add_paragraph(f"  {q['number']}. {q.get('answer', '?')}")
    if multi_qs:
        doc.add_paragraph("多选题答案：")
        for q in multi_qs:
            doc.add_paragraph(f"  {q['number']}. {q.get('answer', '?')}")
    if fill_qs:
        doc.add_paragraph("填空题答案：")
        for q in fill_qs:
            doc.add_paragraph(f"  {q['number']}. {q.get('answer', '?')}")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
