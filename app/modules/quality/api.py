"""Quality 模块 API 路由。"""

import re
from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi import APIRouter, Body, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse, StreamingResponse

from app.modules.quality.report_generator import extract_template_placeholders
from app.modules.quality.schemas import UploadLcResponse
from app.modules.quality.service import lc_report_service
from app.shared.module_registry import MODULES_BY_CODE

router = APIRouter()
_module = MODULES_BY_CODE["quality"]
REPORT_TEMPLATE_DIR = Path(__file__).resolve().parent.parent.parent.parent / "报告模板"
PLACEHOLDER_RE = re.compile(r"\{\{(.+?)\}\}")


@router.get("/", summary=f"{_module.name}模块信息")
async def read_module() -> dict[str, str]:
    return _module.as_dict()


@router.post("/lc/upload", response_model=UploadLcResponse, summary="上传液相计算表并解析")
async def upload_lc_excel(file: UploadFile = File(...)):
    filename = file.filename or "unknown.xlsx"
    if not filename.lower().endswith((".xlsx", ".xls")):
        raise HTTPException(status_code=400, detail="仅支持 .xlsx 或 .xls")
    file_bytes = await file.read()
    if not file_bytes: raise HTTPException(status_code=400, detail="文件为空")
    if len(file_bytes) > 10 * 1024 * 1024: raise HTTPException(status_code=400, detail="文件超过 10MB")
    try: return lc_report_service.parse_and_validate(file_bytes, filename)
    except ValueError as e: raise HTTPException(status_code=422, detail=f"无法解析：{e}") from e


# ─── 模板管理 ───

def _scan_templates(root: Path) -> list[dict]:
    items = []
    for p in sorted(root.iterdir()):
        if p.is_dir():
            items.append({"type": "folder", "name": p.name, "children": _scan_templates(p)})
        elif p.suffix == ".docx":
            phs = extract_template_placeholders(str(p))
            st = p.stat()
            items.append({"type": "template", "filename": p.name, "size_kb": round(st.st_size / 1024, 1), "modified": st.st_mtime, "placeholder_count": len(phs),
                "placeholders": [{"name": ph.name, "decimals": ph.decimals, "suffix": ph.suffix} for ph in phs]})
    return items


@router.get("/templates", summary="列出模板")
async def list_templates():
    return _scan_templates(REPORT_TEMPLATE_DIR) if REPORT_TEMPLATE_DIR.exists() else []


@router.get("/templates/all-placeholders", summary="获取所有模板的所有占位符名称")
async def all_placeholders():
    """收集所有模板中出现的所有占位符名称（去重），用于动态生成表单。"""
    names: set[str] = set()
    def walk(items: list[dict]):
        for item in items:
            if item["type"] == "folder": walk(item.get("children", []))
            else:
                for ph in item.get("placeholders", []): names.add(ph["name"])
    walk(_scan_templates(REPORT_TEMPLATE_DIR) if REPORT_TEMPLATE_DIR.exists() else [])
    # 基础字段始终在前面
    base = ["流水号", "批号", "规格", "生产日期", "批量_kg", "有效期_年"]
    rest = sorted(n for n in names if n not in base)
    return base + rest


@router.post("/templates/folders", summary="创建文件夹")
async def create_folder(name: str = Body(..., embed=True)):
    (REPORT_TEMPLATE_DIR / name).mkdir(parents=True, exist_ok=True)
    return {"name": name}


@router.delete("/templates/folders", summary="删除空文件夹")
async def delete_folder(name: str = Body(..., embed=True)):
    p = REPORT_TEMPLATE_DIR / name
    if not p.exists(): raise HTTPException(status_code=404, detail="不存在")
    try: p.rmdir()
    except OSError: raise HTTPException(status_code=400, detail="文件夹不为空") from None
    return {"message": f"已删除 {name}"}


@router.post("/templates/upload", summary="上传模板")
async def upload_template(folder: str = Form(""), file: UploadFile = File(...)):
    if not file.filename or not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="仅支持 .docx")
    dest = REPORT_TEMPLATE_DIR / folder
    dest.mkdir(parents=True, exist_ok=True)
    content = await file.read()
    if not content: raise HTTPException(status_code=400, detail="文件为空")
    if len(content) > 5 * 1024 * 1024: raise HTTPException(status_code=400, detail="超过5MB")
    (dest / file.filename).write_bytes(content)
    return {"filename": file.filename, "folder": folder}


@router.get("/templates/{path:path}/download", summary="下载模板")
async def download_template(path: str):
    full = REPORT_TEMPLATE_DIR / path
    if not full.exists(): raise HTTPException(status_code=404, detail="不存在")
    return FileResponse(str(full), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=full.name)


@router.delete("/templates/{path:path}", summary="删除模板")
async def delete_template(path: str):
    full = REPORT_TEMPLATE_DIR / path
    if not full.exists(): raise HTTPException(status_code=404, detail="不存在")
    full.unlink()
    return {"message": "已删除"}


# ─── 产品代码管理 ───

PRODUCTS_FILE = REPORT_TEMPLATE_DIR.parent / "products.json"


def _load_products() -> list[dict]:
    import json
    if PRODUCTS_FILE.exists():
        try: return json.loads(PRODUCTS_FILE.read_text())
        except: pass
    return []


def _save_products(data: list[dict]):
    import json
    PRODUCTS_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2))


@router.get("/products", summary="列出产品代码映射")
async def list_products():
    return _load_products()


@router.post("/products", summary="保存产品代码映射")
async def save_products(data: list[dict] = Body(...)):
    _save_products(data)
    return {"message": "已保存", "count": len(data)}


# ─── 报告生成 ───

def _replace_placeholders(para, data: dict):
    full = para.text
    matches = list(PLACEHOLDER_RE.finditer(full))
    if not matches: return
    new_text = full
    for m in reversed(matches):
        raw = m.group(1).strip(); key = raw.split("|")[0].strip()
        val = str(data.get(key, "-"))
        new_text = new_text[:m.start()] + val + new_text[m.end():]
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]: r.text = ""


@router.post("/report/generate", summary="生成报告单")
async def generate_report(data: dict = Body(...), template: str = Body("万古霉素/3205.docx", embed=True)):
    tp = REPORT_TEMPLATE_DIR / template
    if not tp.exists(): raise HTTPException(status_code=404, detail=f"模板不存在：{template}")
    doc = Document(str(tp))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs: _replace_placeholders(para, data)
    for para in doc.paragraphs: _replace_placeholders(para, data)
    out = BytesIO(); doc.save(out); out.seek(0)
    return StreamingResponse(out, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=COA-{data.get('批号','report')}.docx"})
