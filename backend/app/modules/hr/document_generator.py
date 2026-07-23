"""Generate onboarding training record documents from templates."""

from datetime import date
from io import BytesIO
from pathlib import Path
from typing import Any
import copy

from dateutil.relativedelta import relativedelta
from docx import Document

from app.modules.hr.models import Employee


def _resolve_template(filename: str) -> Path:
    candidates = [
        Path(f"assets/hr/{filename}"),
        Path(f"../assets/hr/{filename}"),
        Path(__file__).resolve().parent.parent.parent.parent / "assets/hr" / filename,
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError(f"模板文件未找到: {filename}")


def _replace_all(doc, placeholders: dict[str, str]) -> None:
    """Replace all {key} placeholders across the document."""
    for para in doc.paragraphs:
        for key, val in placeholders.items():
            if key in para.text:
                para.text = para.text.replace(key, val)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, val in placeholders.items():
                        if key in para.text:
                            para.text = para.text.replace(key, val)


def _calc_probation_date(hire_date: date) -> str:
    return str(hire_date + relativedelta(months=3, days=-1))


def generate_onboarding_training_record(
    employee: Employee,
    training_items: list[dict[str, Any]] | None = None,
) -> BytesIO:
    """新员工培训记录 — 使用新模板（含年龄、转正日期等占位符）。"""
    doc = Document(str(_resolve_template("新员工培训记录-模板.docx")))

    hire_d = employee.hire_date
    hire_str = str(hire_d) if hire_d else ""
    # 先替换双括号公式（避免 {入职日期} 被先替换掉）
    _replace_all(doc, {"{{入职日期}+3M-1day}": _calc_probation_date(hire_d) if hire_d else ""})
    _replace_all(doc, {
        "{姓名}": employee.name or "",
        "{性别}": employee.gender or "",
        "{年龄}": str(employee.age) if employee.age else "",
        "{体现部门}": employee.department or "",
        "{体现岗位}": employee.position or "",
        "{入职日期}": hire_str,
    })

    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf


async def _get_position_training(employee) -> tuple[list[str], list[str]]:
    """获取员工岗位关联的培训科目及培训师。返回 (培训内容列表, 培训师列表)。"""
    from sqlalchemy import select
    from app.core.database import async_session_factory
    from app.modules.hr.models import PositionTraining

    contents = []
    trainers = []
    async with async_session_factory() as s:
        r = await s.execute(select(PositionTraining).where(
            PositionTraining.department == employee.department,
            PositionTraining.position_name == employee.position,
            PositionTraining.is_deleted == False,
        ))
        for pt in r.scalars().all():
            label = f"{pt.training_category}" + (f" （{pt.training_method}）" if pt.training_method else "")
            contents.append(label)
            if pt.trainer and pt.trainer not in trainers:
                trainers.append(pt.trainer)
    return contents, trainers


async def generate_onboarding_certificate(employee: Employee) -> BytesIO:
    """上岗证 — 培训内容从岗位培训表自动填充。"""
    doc = Document(str(_resolve_template("上岗证模板.docx")))
    contents, trainer_list = await _get_position_training(employee)

    _replace_all(doc, {
        "{姓名}": employee.name or "",
        "{性别}": employee.gender or "",
        "{学历}": employee.education or "",
        "{体现部门}": employee.department or "",
        "{培训清单内容}": "\n".join(contents) if contents else "",
        "{培训师}": "、".join(trainer_list) if trainer_list else "",
    })

    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf
