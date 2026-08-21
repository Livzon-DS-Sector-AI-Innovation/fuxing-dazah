"""AI exam generation service. Supports both API and local mode."""

import logging
from io import BytesIO
from typing import Any

from docx import Document as DocxDocument

from app.modules.hr import config as hr_config

logger = logging.getLogger(__name__)

EXAM_QUESTION_KEYS = (
    "choice_questions",
    "true_false_questions",
    "multi_choice_questions",
    "fill_blank_questions",
)

ExamQuestions = dict[str, list[dict[str, Any]]]


def _extract_docx_content(file_bytes: bytes) -> dict[str, Any]:
    doc = DocxDocument(BytesIO(file_bytes))
    all_paragraphs: list[str] = []
    bold_texts: list[str] = []
    for para in doc.paragraphs:
        parts: list[str] = []
        for run in para.runs:
            text = run.text.strip()
            if not text:
                continue
            parts.append(text)
            if run.bold:
                bold_texts.append(text)
        if parts:
            all_paragraphs.append(" ".join(parts))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        text = run.text.strip()
                        if text and run.bold:
                            bold_texts.append(text)
    return {"full_text": "\n".join(all_paragraphs), "bold_texts": bold_texts}


def _extract_text_content(file_bytes: bytes) -> dict[str, Any]:
    text = file_bytes.decode("utf-8", errors="ignore")
    return {"full_text": text, "bold_texts": []}


def _parse_file(file_bytes: bytes, filename: str) -> dict[str, Any]:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext in ("docx", "doc"):
        return _extract_docx_content(file_bytes)
    return _extract_text_content(file_bytes)


def _generate_local(content: dict[str, Any], config: dict[str, Any] | None = None) -> ExamQuestions:
    """离线模式：基于规则从文本生成多题型试卷。加粗文字优先挖空出题。"""
    choice_count = config.get("choice_count", 5) if config else 5
    tf_count = config.get("true_false_count", 5) if config else 5
    multi_count = config.get("multi_choice_count", 0) if config else 0
    fill_count = config.get("fill_blank_count", 0) if config else 0

    sentences = [s.strip() for s in content["full_text"].replace("\n", "。").split("。") if len(s.strip()) > 8]
    bold_texts = content.get("bold_texts", [])

    choice_qs: list[dict[str, Any]] = []
    tf_qs: list[dict[str, Any]] = []
    multi_qs: list[dict[str, Any]] = []
    fill_qs: list[dict[str, Any]] = []

    # ── 先对句子打分 ──
    keywords = ["必须", "应当", "禁止", "要求", "包括", "负责", "管理", "检查", "培训", "安全", "操作", "设备"]
    scored_sentences = []
    for s in sentences:
        score = sum(1 for kw in keywords if kw in s)
        if score > 0:
            scored_sentences.append((score, s))
    scored_sentences.sort(key=lambda x: -x[0])

    # ── 填空题：先从加粗文字挖空，不够再从关键句中提取 ──
    for i, bold in enumerate(bold_texts):
        fill_qs.append({"number": len(fill_qs) + 1, "question": f"请填写关键知识点：{bold}", "answer": bold})
    for _, s in scored_sentences:
        if len(fill_qs) >= fill_count:
            break
        for kw in keywords:
            if kw in s and not any(q["question"].find(kw) >= 0 for q in fill_qs):
                q_text = s.replace(kw, "______", 1)
                fill_qs.append({"number": len(fill_qs) + 1, "question": q_text, "answer": kw})
                break

    # ── 判断题：加粗文字生成正误判断 ──
    for i, bold in enumerate(bold_texts[:tf_count]):
        is_correct = i % 2 == 0  # 一半正确一半错误
        if is_correct:
            tf_qs.append({"number": i + 1, "question": f"「{bold}」这个说法是否正确？", "answer": "正确"})
        else:
            # 故意改错
            wrong = bold.replace("必须", "不必").replace("禁止", "允许").replace("应", "不应")
            if wrong == bold:
                wrong = "不" + bold
            tf_qs.append({"number": i + 1, "question": f"「{wrong}」这个说法是否正确？", "answer": "错误"})

    # 补判断题
    for s in sentences:
        if len(tf_qs) >= tf_count:
            break
        if s not in bold_texts and len(s) > 12:
            tf_qs.append({"number": len(tf_qs) + 1, "question": s + "？", "answer": "正确"})

    # ── 单选题：从高分句子中提取核心词挖空 ──
    for i, (_, s) in enumerate(scored_sentences[:choice_count]):
        # 找句子中的核心关键词挖空
        target_kw = None
        for kw in keywords:
            if kw in s:
                target_kw = kw
                break
        if target_kw:
            q = s.replace(target_kw, "____", 1)
            choice_qs.append({
                "number": i + 1,
                "question": q,
                "options": [
                    {"label": "A", "text": target_kw},
                    {"label": "B", "text": "不需要" + target_kw},
                    {"label": "C", "text": "视情况" + target_kw},
                    {"label": "D", "text": "由领导决定"},
                ],
                "answer": "A",
            })

    # ── 多选题：选包含多个关键点的句子 ──
    multi_candidates = [(sum(1 for kw in keywords if kw in s), s) for s in sentences if sum(1 for kw in keywords if kw in s) >= 2]
    multi_candidates.sort(key=lambda x: -x[0])
    for i, (_, s) in enumerate(multi_candidates[:multi_count]):
        # 提取2-3个关键词作为正确答案
        found_kws = [kw for kw in keywords if kw in s][:3]
        options = [{"label": chr(65 + j), "text": kw} for j, kw in enumerate(found_kws)]
        answer = ",".join(chr(65 + j) for j in range(len(found_kws)))
        # 加一个干扰项
        options.append({"label": chr(65 + len(found_kws)), "text": "以上都不对"})
        multi_qs.append({
            "number": i + 1,
            "question": s,
            "options": options,
            "answer": answer,
        })

    result = {
        "choice_questions": choice_qs[:choice_count],
        "true_false_questions": tf_qs[:tf_count],
        "multi_choice_questions": multi_qs[:multi_count],
        "fill_blank_questions": fill_qs[:fill_count],
    }
    logger.info(
        "_generate_local: bold=%d sentences=%d -> choice=%d tf=%d multi=%d fill=%d",
        len(bold_texts), len(sentences),
        len(result["choice_questions"]), len(result["true_false_questions"]),
        len(result["multi_choice_questions"]), len(result["fill_blank_questions"]),
    )
    return result


def _build_prompt(full_text: str, bold_texts: list[str], config: dict[str, Any] | None = None) -> str:
    if config:
        choice_count = config.get("choice_count", 5)
        tf_count = config.get("true_false_count", 5)
        multi_count = config.get("multi_choice_count", 0)
        fill_count = config.get("fill_blank_count", 0)
    else:
        choice_count, tf_count, multi_count, fill_count = 5, 5, 0, 0

    bold_hint = ""
    if bold_texts:
        bold_hint = "\n\n【重点关注】以下内容在原文件中被加粗标记，请优先作为考点出题：\n" + "\n".join(f"  - {t}" for t in bold_texts[:20])

    reqs = []
    if choice_count > 0:
        reqs.append(f"出 {choice_count} 道单选题（4 选项 A/B/C/D，只有一个正确答案）")
    if tf_count > 0:
        reqs.append(f"出 {tf_count} 道判断题（answer 只能是\"正确\"或\"错误\"，正确与错误大致各半）")
    if multi_count > 0:
        reqs.append(f"出 {multi_count} 道多选题（4 选项，有 2-3 个正确答案）")
    if fill_count > 0:
        reqs.append(f"出 {fill_count} 道填空题（将关键知识点挖空，question 中用______表示空位，question 里不能出现答案）")
    reqs.append("题目必须严格基于材料内容，禁止编造材料中不存在的知识点")
    reqs.append("题目应覆盖材料关键知识点，加粗内容优先作为考点")
    reqs.append("干扰选项要合理、有区分度，正确答案在各选项间均匀分布，不要固定为 A")
    reqs.append("严格按 JSON 格式输出，不要输出 JSON 以外的内容")

    json_parts = []
    if choice_count > 0:
        json_parts.append('"choice_questions": [{"question":"...","options":[{"label":"A","text":"..."},...],"answer":"A"}]')
    if tf_count > 0:
        json_parts.append('"true_false_questions": [{"question":"...","answer":"正确"}]')
    if multi_count > 0:
        json_parts.append('"multi_choice_questions": [{"question":"...","options":[{"label":"A","text":"..."},...],"answer":"A,B"}]')
    if fill_count > 0:
        json_parts.append('"fill_blank_questions": [{"question":"...","answer":"关键词"}]')

    return f"""你是培训考核出题老师。根据以下材料生成试卷。

要求：{'; '.join(reqs)}

输出格式：{{ {', '.join(json_parts)} }}

材料：
{full_text[:8000]}{bold_hint}"""


def _expected_counts(config: dict[str, Any] | None) -> dict[str, int]:
    return {
        "choice_questions": config.get("choice_count", 5) if config else 5,
        "true_false_questions": config.get("true_false_count", 5) if config else 5,
        "multi_choice_questions": config.get("multi_choice_count", 0) if config else 0,
        "fill_blank_questions": config.get("fill_blank_count", 0) if config else 0,
    }


def _validate_exam(result: dict[str, Any], config: dict[str, Any] | None) -> ExamQuestions:
    """校验 LLM 返回的试卷结构，不合格直接抛错由上层降级。"""
    if not isinstance(result, dict):
        raise ValueError("AI 返回结果不是合法的试卷结构")

    expected = _expected_counts(config)
    validated: ExamQuestions = {}
    for key in EXAM_QUESTION_KEYS:
        qs = result.get(key, [])
        want = expected[key]
        if want <= 0:
            validated[key] = []
            continue
        if not isinstance(qs, list) or not qs:
            raise ValueError(f"AI 未生成任何{key}")
        clean = []
        for q in qs[:want]:
            if not isinstance(q, dict) or not str(q.get("question", "")).strip():
                raise ValueError(f"AI 生成的{key}缺少题干")
            item: dict[str, Any] = {"question": str(q["question"]).strip()}
            answer = str(q.get("answer", "")).strip()
            if not answer:
                raise ValueError(f"AI 生成的{key}缺少答案")
            if key in ("choice_questions", "multi_choice_questions"):
                options = q.get("options")
                if not isinstance(options, list) or len(options) < 2:
                    raise ValueError(f"AI 生成的{key}选项不足")
                labels: list[str] = []
                clean_opts: list[dict[str, str]] = []
                for opt in options:
                    if not isinstance(opt, dict) or not str(opt.get("text", "")).strip():
                        raise ValueError(f"AI 生成的{key}选项不合法")
                    label = str(opt.get("label", "")).strip().upper() or chr(65 + len(labels))
                    labels.append(label)
                    clean_opts.append({"label": label, "text": str(opt["text"]).strip()})
                item["options"] = clean_opts
                answers = [a.strip().upper() for a in answer.replace("，", ",").split(",") if a.strip()]
                if not answers or any(a not in labels for a in answers):
                    raise ValueError(f"AI 生成的{key}答案与选项不匹配: {answer}")
                item["answer"] = ",".join(answers)
            elif key == "true_false_questions":
                if answer in ("对", "是", "√", "T", "true", "True"):
                    answer = "正确"
                elif answer in ("错", "否", "×", "F", "false", "False"):
                    answer = "错误"
                if answer not in ("正确", "错误"):
                    raise ValueError(f"AI 生成的判断题答案不合法: {answer}")
                item["answer"] = answer
            else:
                item["answer"] = answer
            clean.append(item)
        validated[key] = clean
    return validated


async def _generate_via_ai(content: dict[str, Any], config: dict[str, Any] | None = None) -> ExamQuestions:
    """在线模式：调用大模型基于培训材料生成试卷。"""
    from app.modules.hr.ai_service import AiChatService

    prompt = _build_prompt(content["full_text"], content["bold_texts"], config)
    result = await AiChatService.call_json(
        prompt,
        system_prompt="你是培训考核出题老师。严格基于给定材料出题，只输出合法 JSON。",
        api_key=hr_config.HR_AI_API_KEY,
    )
    validated = _validate_exam(result, config)
    logger.info(
        "_generate_via_ai: choice=%d tf=%d multi=%d fill=%d",
        len(validated["choice_questions"]), len(validated["true_false_questions"]),
        len(validated["multi_choice_questions"]), len(validated["fill_blank_questions"]),
    )
    return validated


async def generate_exam(file_bytes: bytes, filename: str, config: dict[str, Any] | None = None) -> ExamQuestions:
    """Generate exam questions. 配置了 HR_AI_API_KEY 时优先走大模型出题，失败自动降级到本地规则模式。"""
    content = _parse_file(file_bytes, filename)
    if not content["full_text"].strip():
        raise ValueError("文件中未检测到文本内容")

    result: ExamQuestions | None = None
    if hr_config.HR_AI_API_KEY:
        try:
            result = await _generate_via_ai(content, config)
        except Exception:
            logger.exception("AI 出题失败，降级为本地规则出题")
    if result is None:
        result = _generate_local(content, config)

    for key in EXAM_QUESTION_KEYS:
        for i, q in enumerate(result.get(key, [])):
            q["number"] = i + 1

    return result


def _find_exam_template() -> str:
    """查找试卷模板文件，复用 find_hr_template 的多路径搜索逻辑。"""
    from app.modules.hr.template_utils import find_hr_template
    try:
        return str(find_hr_template("试卷模板.docx"))
    except FileNotFoundError:
        return ""


def export_exam(data: dict[str, Any]) -> BytesIO:
    """导出考试试卷为 Word 文档，优先使用试卷模板。"""
    from docx import Document
    from docx.shared import Pt

    template_path = _find_exam_template()
    if template_path:
        doc = Document(template_path)
    else:
        doc = Document()
    doc.styles["Normal"].font.size = Pt(11)

    title = data.get("title", "培训考试试卷")
    header_para = None
    for p in doc.paragraphs:
        if "{{培训内容}}" in (p.text or ""):
            header_para = p
            break

    if header_para:
        if header_para.runs:
            header_para.runs[0].text = f"{title}试题"
            for r in header_para.runs[1:]:
                r.text = ""
    else:
        doc.add_heading(title, level=1)
        doc.add_paragraph("")

    # 出卷人信息行
    if data.get("examiner"):
        doc.add_paragraph(f"出卷人：{data['examiner']}")
        doc.add_paragraph("")

    choice_qs = data.get("choice_questions", [])
    tf_qs = data.get("true_false_questions", [])
    multi_qs = data.get("multi_choice_questions", [])
    fill_qs = data.get("fill_blank_questions", [])

    for heading, qs in [("一、单选题", choice_qs), ("二、判断题", tf_qs),
                          ("三、多选题", multi_qs), ("四、填空题", fill_qs)]:
        if not qs:
            continue
        p = doc.add_paragraph()
        run = p.add_run(f"{heading}（共{len(qs)}题）")
        run.bold = True
        run.font.size = Pt(14)
        doc.add_paragraph("")

        if heading == "一、单选题":
            for q in qs:
                doc.add_paragraph(f"{q['number']}. {q['question']}")
                for opt in q.get("options", []):
                    doc.add_paragraph(f"    {opt['label']}. {opt['text']}")
                doc.add_paragraph("")
        elif heading == "二、判断题":
            for q in qs:
                doc.add_paragraph(f"{q['number']}. {q['question']}  对 □    错 □")
                doc.add_paragraph("")
        elif heading == "三、多选题":
            for q in qs:
                doc.add_paragraph(f"{q['number']}. {q['question']}")
                for opt in q.get("options", []):
                    doc.add_paragraph(f"    {opt['label']}. {opt['text']}")
                doc.add_paragraph("")
        elif heading == "四、填空题":
            for q in qs:
                doc.add_paragraph(f"{q['number']}. {q['question']}")
                doc.add_paragraph("    ____________________")
                doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("参考答案")
    run.bold = True
    run.font.size = Pt(14)
    if choice_qs:
        doc.add_paragraph("单选题答案：")
        for q in choice_qs:
            doc.add_paragraph(f"  {q['number']}. {q.get('answer', '?')}")
    if tf_qs:
        doc.add_paragraph("判断题答案：")
        for q in tf_qs:
            doc.add_paragraph(f"  {q['number']}. {q.get('answer', '?')}")
    if multi_qs:
        doc.add_paragraph("多选题答案：")
        for q in multi_qs:
            doc.add_paragraph(f"  {q['number']}. {q.get('answer', '?')}")
    if fill_qs:
        doc.add_paragraph("填空题答案：")
        for q in fill_qs:
            doc.add_paragraph(f"  {q['number']}. {q.get('answer', '?')}")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
