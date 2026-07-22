"""解除劳动关系证明 DOCX / PDF 生成器。"""

from datetime import date
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from app.modules.hr.template_utils import find_hr_template


def _fmt_date(d: date | str) -> str:
    if isinstance(d, date):
        return f"{d.year} 年 {d.month} 月 {d.day} 日"
    return str(d or "")


def generate_termination_certificate_docx(
    *,
    name: str,
    id_number: str,
    department: str,
    position: str,
    entry_date: date | str,
    leave_date: date | str,
    leave_reason: str = "个人原因",
) -> BytesIO:
    """生成解除劳动关系证明 DOCX，返回 BytesIO。"""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.size = Pt(12)
    style.font.name = "宋体"
    style.paragraph_format.line_spacing = 1.8

    # 页边距
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

    # 文件编号（右上角）
    p = doc.add_paragraph("HR-RE-013")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(24)

    # 标题
    title = doc.add_paragraph("解除劳动关系证明")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(18)
        run.font.bold = True
    title.paragraph_format.space_after = Pt(24)

    entry_str = _fmt_date(entry_date)
    leave_str = _fmt_date(leave_date)

    # 正文
    body_text = (
        f"兹有我司原职工姓名：{name}，身份证号：{id_number}，"
        f"入职时间 {entry_str} 到我公司工作，{department}部门，"
        f"{position}岗位 工作。现因{leave_reason}，"
        f"于 {leave_str} 正式解除劳动关系。"
    )
    p = doc.add_paragraph(body_text)
    p.paragraph_format.first_line_indent = Cm(0.7)
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph("特此证明。")
    p.paragraph_format.space_after = Pt(18)

    # 注意事项
    p = doc.add_paragraph(
        "1、员工离职后仍需履行保密义务，未经我公司书面许可，不得向任何单位和个人透露"
        "我公司商业秘密和其他经营秘密（造成影响公司保留追责权力）"
    )
    p.paragraph_format.space_after = Pt(4)
    p = doc.add_paragraph("2、本证明仅开具一次，请妥善保管，如遗失不补开。")
    p.paragraph_format.space_after = Pt(36)

    # 落款 + 公章 + 日期（右对齐）
    today = date.today()
    today_str = f"{today.year} 年 {today.month} 月 {today.day} 日"

    stamp_path = find_hr_template("company_stamp.png")
    # 公章压在落款公司名上
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    stamp_run = p.add_run()
    stamp_run.add_picture(str(stamp_path), width=Cm(5))
    name_run = p.add_run()
    name_run.text = "\n丽珠集团福州福兴医药有限公司"
    name_run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph(today_str)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 清空页脚页码
    footer = section.footer
    footer.is_linked_to_previous = False
    for fp in list(footer.paragraphs):
        fp._element.getparent().remove(fp._element)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_termination_certificate_html(
    *,
    name: str,
    id_number: str,
    department: str,
    position: str,
    entry_date: date | str,
    leave_date: date | str,
    leave_reason: str = "个人原因",
) -> str:
    """生成解除劳动关系证明 HTML（用于预览和 PDF）。"""
    import base64

    entry_str = _fmt_date(entry_date)
    leave_str = _fmt_date(leave_date)
    today = date.today()
    today_str = f"{today.year} 年 {today.month} 月 {today.day} 日"

    stamp_path = find_hr_template("company_stamp.png")
    stamp_b64 = base64.b64encode(stamp_path.read_bytes()).decode()

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<style>
  @page {{ size: A4; margin: 2.5cm 3cm; }}
  body {{ font-family: "SimSun", "宋体", serif; font-size: 12pt; line-height: 2; color: #000; }}
  .center {{ text-align: center; }}
  .right {{ text-align: right; }}
  .indent {{ text-indent: 2em; }}
  .title {{ font-size: 18pt; font-weight: bold; letter-spacing: 0.3em; }}
  .signature {{ text-align: right; margin-top: 24px; }}
  .signature-stamp {{ display: block; width: 6cm; margin-left: auto; }}
  .signature-text {{ margin-top: -3.5cm; font-size: 12pt; }}
</style>
</head>
<body>
<p style="text-align:right;font-size:9pt;color:#666;">HR-RE-013</p>
<p class="center title">解除劳动关系证明</p>
<p>&nbsp;</p>
<p class="indent">兹有我司原职工姓名：{name}，身份证号：{id_number}，
入职时间 {entry_str} 到我公司工作，{department}部门，
{position}岗位 工作。现因{leave_reason}，
于 {leave_str} 正式解除劳动关系。</p>
<p>特此证明。</p>
<p>&nbsp;</p>
<p>1、员工离职后仍需履行保密义务，未经我公司书面许可，不得向任何单位和个人透露我公司商业秘密和其他经营秘密（造成影响公司保留追责权力）</p>
<p>2、本证明仅开具一次，请妥善保管，如遗失不补开。</p>
<p>&nbsp;</p>
<div class="signature">
  <img class="signature-stamp" src="data:image/png;base64,{stamp_b64}" alt="公章">
  <p class="signature-text">丽珠集团福州福兴医药有限公司</p>
</div>
<p class="right">{today_str}</p>
</body>
</html>"""
