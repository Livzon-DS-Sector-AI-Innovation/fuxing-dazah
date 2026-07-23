"""Generate training sign-in sheet (签到表) from Word template.

Uses the user-provided template for formatting.  Each page loads a fresh
copy of the template, fills metadata and employee names, then all pages are
combined into a single .docx with page breaks in between.
"""

import copy
import re
from io import BytesIO
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

from app.modules.hr.schemas import TrainingSignInSheetInput

EMPLOYEES_PER_PAGE = 15


def _find_template() -> Path:
    candidates = [
        Path("assets/hr/7.5培训签到表.docx"),
        Path("../assets/hr/7.5培训签到表.docx"),
        Path(__file__).resolve().parent.parent.parent.parent
        / "assets/hr"
        / "7.5培训签到表.docx",
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError("模板文件未找到: 7.5培训签到表.docx")


def _set_cell_text(cell, text: str) -> None:
    """Set the text of a cell, preserving the first run's formatting."""
    first_run = None
    for p in cell.paragraphs:
        for r in p.runs:
            if first_run is None:
                first_run = r
            r.text = ""
    if first_run is not None:
        first_run.text = str(text or "")
    elif cell.paragraphs:
        run = cell.paragraphs[0].add_run(str(text or ""))


def _set_xml_cell_text(tc, text: str) -> None:
    """Set text in a w:tc XML element."""
    wts = list(tc.iter(qn("w:t")))
    if wts:
        for t in wts:
            t.text = ""
        wts[0].text = str(text or "")
    else:
        for p in tc.iter(qn("w:p")):
            for r in p.iter(qn("w:r")):
                new_t = etree.SubElement(r, qn("w:t"))
                new_t.text = str(text or "")
                return


def _compute_duration_hours(start: str | None, end: str | None) -> str:
    if not start or not end:
        return ""
    try:
        s = datetime.strptime(start, "%H:%M")
        e = datetime.strptime(end, "%H:%M")
        diff_hours = (e - s).total_seconds() / 3600
        if diff_hours <= 0:
            return ""
        rounded = round(diff_hours * 2) / 2
        if rounded == int(rounded):
            return f"{int(rounded)}小时"
        return f"{rounded}小时"
    except ValueError:
        return ""


def _fill_metadata(table, data: TrainingSignInSheetInput) -> None:
    """Fill metadata into template table by cell position."""
    # Row 0, Col 1: 培训内容 (merged cols 1-3)
    _set_cell_text(table.rows[0].cells[1], data.topic)
    # Row 1, Col 1: 培训对象
    _set_cell_text(table.rows[1].cells[1], data.department)
    # Row 1, Col 3: 培训方式
    _set_cell_text(table.rows[1].cells[3], data.training_method or "")
    # Row 2, Col 1: 课时
    hours = _compute_duration_hours(data.training_time_start, data.training_time_end)
    _set_cell_text(table.rows[2].cells[1], hours)
    # Row 2, Col 3: 考核方式
    _set_cell_text(table.rows[2].cells[3], data.assessment_method or "")
    # Row 3, Col 1: 培训日期 (merged cols 1-3)
    _set_cell_text(table.rows[3].cells[1], str(data.training_date) if data.training_date else "")


def _fill_employee_rows(table, employee_names: list[str], employee_departments: dict[str, str], default_department: str) -> None:
    """Clone Row 5 for each employee and fill name + department."""
    if len(table.rows) < 7:
        return

    template_row = table.rows[5]
    footer_row = table.rows[6]

    for name in employee_names:
        new_tr = copy.deepcopy(template_row._tr)
        footer_row._tr.addprevious(new_tr)
        new_row = table.rows[len(table.rows) - 2]
        _set_cell_text(new_row.cells[0], name)
        _set_cell_text(new_row.cells[1], employee_departments.get(name, default_department))
        # Cell 2 merge ghost, cell 3 signature — empty

    table._tbl.remove(template_row._tr)


def _update_section_header_page_numbers(section, page: int, total_pages: int) -> None:
    """Update page numbers in a single section's header."""
    header = section.header
    if header is None:
        return
    for paragraph in header.paragraphs:
        runs = paragraph.runs
        page_run_idx = None
        total_run_idx = None
        for i in range(len(runs) - 2):
            if runs[i].text.strip() == "第" and runs[i + 2].text.strip() == "页":
                page_run_idx = i + 1
            if i + 6 < len(runs):
                if runs[i + 4].text.strip() == "，共" and runs[i + 6].text.strip() == "页":
                    total_run_idx = i + 5
        if page_run_idx is not None:
            runs[page_run_idx].text = str(page)
        if total_run_idx is not None:
            runs[total_run_idx].text = str(total_pages)


def _add_page_break(doc: Document) -> None:
    """Insert a page break paragraph at the end of the document body."""
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    doc.element.body.append(p)


# ─── public API ───

def generate_training_sign_in_sheet(data: TrainingSignInSheetInput) -> BytesIO:
    """Generate sign-in sheet from template, matching the preview layout.

    Page 1 is built directly in the template document (preserves header images).
    Pages 2+ are built in fresh template loads, then their tables are inserted
    into the output document with page breaks before each.
    """
    template_path = _find_template()
    total = len(data.employee_names)
    total_pages = max(1, (total + EMPLOYEES_PER_PAGE - 1) // EMPLOYEES_PER_PAGE)

    # Page 1: build directly in template (preserves header with logo)
    output_doc = Document(str(template_path))
    output_table = output_doc.tables[0]

    first_names = data.employee_names[:EMPLOYEES_PER_PAGE]
    _fill_metadata(output_table, data)
    _fill_employee_rows(output_table, first_names, data.employee_departments, data.department)

    # Remove trainer signature row for non-last pages (it only appears on the final page)
    # The footer is always the LAST row after _fill_employee_rows
    if total_pages > 1:
        last_row = output_table.rows[-1]
        output_table._tbl.remove(last_row._tr)

    _update_section_header_page_numbers(output_doc.sections[0], 1, total_pages)

    # Pages 2+: build new tables in the same document using the template's table as a starting point,
    # then append them with page breaks in between.
    for page_idx in range(1, total_pages):
        start = page_idx * EMPLOYEES_PER_PAGE
        end = min(start + EMPLOYEES_PER_PAGE, total)
        page_names = data.employee_names[start:end]
        is_last = page_idx == total_pages - 1

        # Add page break before this page
        output_doc.add_page_break()

        # Load fresh template to get a clean table with correct styling
        temp_doc = Document(str(template_path))
        temp_table = temp_doc.tables[0]
        _fill_metadata(temp_table, data)
        _fill_employee_rows(temp_table, page_names, data.employee_departments, data.department)
        if not is_last:
            # Footer is the last row after employee rows are added
            last_row = temp_table.rows[-1]
            temp_table._tbl.remove(last_row._tr)

        # Save temp doc to BytesIO, then re-open and extract the table's XML properly
        temp_buf = BytesIO()
        temp_doc.save(temp_buf)
        temp_buf.seek(0)

        # Now open the temp doc and copy its table into output
        # We use a fresh load to ensure all XML namespaces are properly handled
        reloaded = Document(temp_buf)
        reloaded_table = reloaded.tables[0]

        body = output_doc.element.body
        sect_pr = body.find(qn("w:sectPr"))
        imported = copy.deepcopy(reloaded_table._tbl)
        if sect_pr is not None:
            body.insert(list(body).index(sect_pr), imported)
        else:
            body.append(imported)

    buf = BytesIO()
    output_doc.save(buf)
    buf.seek(0)
    return buf
