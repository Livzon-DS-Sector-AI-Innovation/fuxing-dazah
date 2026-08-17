"""文档表格提取工具测试（LLM 调用全部 mock；转换链路为纯 Python，无系统依赖）。"""

import asyncio
from pathlib import Path

import pytest

from app.modules.toolbox.registry import ToolError, get_tool
from app.modules.toolbox.tools._qwen import parse_rows_response
from app.modules.toolbox.tools.doc_table_extract import (
    dedupe_rows,
    drop_header_rows,
    parse_field_lines,
    rows_to_csv,
)


def test_parse_field_lines_plain_and_with_desc() -> None:
    fields = parse_field_lines("批号\n含量%|色谱峰含量数值\n\n")
    assert fields == [
        {"name": "批号", "description": ""},
        {"name": "含量%", "description": "色谱峰含量数值"},
    ]


def test_parse_field_lines_duplicate_raises() -> None:
    with pytest.raises(ToolError):
        parse_field_lines("批号\n批号")


def test_dedupe_rows_removes_adjacent_duplicates() -> None:
    rows = [["a", "1"], ["a", "1"], ["b", "2"], ["a", "1"]]
    assert dedupe_rows(rows) == [["a", "1"], ["b", "2"], ["a", "1"]]


def test_drop_header_rows_filters_table_header() -> None:
    rows = [["批号", "含量%"], ["B-1", "98.5"], ["批号", "含量%"]]
    columns = ["批号", "含量%"]
    assert drop_header_rows(rows, columns) == [["B-1", "98.5"]]


def test_rows_to_csv_has_bom_and_quotes() -> None:
    csv_text = rows_to_csv([["批号", "含量%"], ["B-1", "98,5"]])
    assert csv_text.startswith("﻿批号,含量%")
    assert '"98,5"' in csv_text


def test_parse_rows_response_valid() -> None:
    rows = parse_rows_response('{"rows": [["B-1", 98.5], ["", 99]]}')
    assert rows == [["B-1", "98.5"], ["", "99"]]


def test_parse_rows_response_rows_is_string() -> None:
    with pytest.raises(ToolError):
        parse_rows_response('{"rows": "text"}')


def test_parse_rows_response_rows_is_dict() -> None:
    with pytest.raises(ToolError):
        parse_rows_response('{"rows": {"批号": "B-1"}}')


def test_parse_rows_response_body_not_json() -> None:
    with pytest.raises(ToolError):
        parse_rows_response("这不是 JSON")


def test_tool_registered_with_two_steps() -> None:
    t = get_tool("doc-table-extract")
    assert t is not None
    assert [s.id for s in t.steps] == ["configure", "extract"]
    files_input = t.steps[1].inputs[0]
    assert files_input.multiple is True
    assert files_input.accept == ".docx"


def test_docx_to_html_keeps_table_structure(tmp_path: Path) -> None:
    """mammoth 纯 Python 转换：docx → 含表格结构与单元格文本的 HTML。"""
    import docx

    from app.modules.toolbox.tools.doc_table_extract import _docx_to_html

    doc = docx.Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "批号"
    table.cell(1, 1).text = "99.2"
    src = tmp_path / "t.docx"
    doc.save(str(src))

    html = _docx_to_html(src)
    assert "<table>" in html
    assert "批号" in html
    assert "99.2" in html


def test_docx_to_html_rejects_corrupt_file(tmp_path: Path) -> None:
    """非 docx 内容（如改名后的 .doc）报友好 ToolError。"""
    from app.modules.toolbox.tools.doc_table_extract import _docx_to_html

    bad = tmp_path / "fake.docx"
    bad.write_bytes(b"\xd0\xcf\x11\xe0")  # OLE 头（.doc 老格式），不是 zip

    with pytest.raises(ToolError, match="无法解析文档"):
        _docx_to_html(bad)


async def test_extract_end_to_end_mocked_llm(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """用真 docx 走完整链路：mammoth → WeasyPrint → fitz 渲染 → mock LLM → 汇总。"""

    import docx

    from app.modules.toolbox import storage
    from app.modules.toolbox.registry import StepContext
    from app.modules.toolbox.tools import doc_table_extract as dte

    # 生成一个含表格的 docx fixture
    doc = docx.Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "批号"
    table.cell(0, 1).text = "含量%"
    table.cell(1, 0).text = "B-2026-001"
    table.cell(1, 1).text = "99.2"
    src = tmp_path / "fixture.docx"
    doc.save(str(src))

    fake_rows = [["批号", "含量%"], ["B-2026-001", "99.2"]]

    async def fake_extract_page(
        image_base64: str, fields: list[dict[str, str]], semaphore: asyncio.Semaphore
    ) -> list[list[str]]:
        return fake_rows

    monkeypatch.setattr(dte, "extract_page", fake_extract_page)
    # 产物写入 storage.EXEC_DIR_ROOT/{execution_id}（系统临时目录），
    # 测试中重定向到 tmp_path，避免污染真实临时目录
    monkeypatch.setattr(storage, "EXEC_DIR_ROOT", tmp_path / "exec-root")

    out_dir = tmp_path / "out"
    out_dir.mkdir()
    ctx = StepContext(
        execution_id="e2e",
        user_id="u1",
        prev_outputs={
            "configure": {
                "fields": [
                    {"name": "批号", "description": ""},
                    {"name": "含量%", "description": ""},
                ]
            }
        },
        file_paths={"files": [str(src)]},
        output_dir=out_dir,
    )
    try:
        result = await dte.doc_table_extract("extract", {"files": {"file_id": "x"}}, ctx)
    except OSError as e:
        # 本机（macOS）缺 WeasyPrint 系统库（libgobject 等）；Docker 镜像已随 hr 模块装齐，
        # 生产路径由 hr offer_generator 同款依赖保障，此处跳过渲染段。
        if "libgobject" in str(e):
            pytest.skip(f"本机缺少 WeasyPrint 系统库（Docker 镜像内已有）: {e}")
        raise
    assert result["columns"] == ["批号", "含量%"]
    assert result["rows"] == [["B-2026-001", "99.2"]]
    # 产物由 storage.save_upload 写入 exec 目录（而非 ctx.output_dir），
    # 通过 storage.resolve_file 校验，与下载端点同一解析路径
    csv_path = storage.resolve_file("e2e", result["csv_file"]["file_id"])
    assert csv_path is not None and csv_path.exists()
