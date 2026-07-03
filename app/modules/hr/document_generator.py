"""Generate onboarding training record documents from templates."""

from io import BytesIO
from pathlib import Path
from typing import Any
import copy

from docx import Document

from app.modules.hr.models import Employee


def _find_template() -> Path:
    candidates = [
        Path("assets/hr/7.3新员工入职培训记录.docx"),
        Path("../assets/hr/7.3新员工入职培训记录.docx"),
        Path(__file__).resolve().parent.parent.parent.parent
        / "assets/hr"
        / "7.3新员工入职培训记录.docx",
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError("模板文件未找到: 7.3新员工入职培训记录.docx")


def _replace_all(doc, placeholders: dict[str, str]) -> None:
    """Replace all {key} placeholders across the document."""
    for para in doc.paragraphs:
        for key, val in placeholders.items():
            if key in para.text:
                para.text = para.text.replace(key, val)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, val in placeholders.items():
                        if key in para.text:
                            para.text = para.text.replace(key, val)


def generate_onboarding_training_record(
    employee: Employee,
    training_items: list[dict[str, Any]] | None = None,
) -> BytesIO:
    """Fill the onboarding training record template with employee data."""
    template_path = _find_template()
    doc = Document(str(template_path))

    # ── Replace all employee info placeholders ──
    quals = employee.qualifications
    qual_str = ", ".join(quals) if isinstance(quals, list) else (str(quals) if quals else "")

    _replace_all(doc, {
        "{姓名}": employee.name or "",
        "{学历}": employee.education or "",
        "{毕业院校}": employee.school or "",
        "{毕业时间}": str(employee.graduation_date) if employee.graduation_date else "",
        "{体现部门}": employee.department or "",
        "{体现岗位}": employee.position or "",
        "{证书}": qual_str,
        "{入职日期}": str(employee.hire_date) if employee.hire_date else "",
        "{异动类别}": "新员工",
    })

    # ── Table 1: fill training items ──
    if training_items and len(doc.tables) > 1:
        t = doc.tables[1]

        # Find Part III row
        part3_idx = None
        for ri, row in enumerate(t.rows):
            if "第三部分" in row.cells[0].text:
                part3_idx = ri
                break
        if part3_idx is None:
            part3_idx = 3  # fallback

        # Template row is just above Part III (should be row 2)
        template_idx = part3_idx - 1

        # Clone and insert training items before Part III
        if template_idx < len(t.rows):
            template_tr = t.rows[template_idx]._tr
            for item in reversed(training_items):
                new_tr = copy.deepcopy(template_tr)
                template_tr.addnext(new_tr)

            # Remove template row
            t._tbl.remove(template_tr)
        else:
            # Fallback: add rows at end
            for item in training_items:
                t.add_row()

        # Fill in item data (now rows template_idx .. template_idx+N-1)
        for i, item in enumerate(training_items):
            row = t.rows[template_idx + i]
            file_name = item.get("file_name", "")
            sop_number = item.get("sop_number", "")
            trainer = item.get("trainer", "")
            method = item.get("method", "")
            title = f"{sop_number} {file_name}" if sop_number else file_name
            if len(row.cells) > 0:
                row.cells[0].paragraphs[0].text = str(i + 1)
            if len(row.cells) > 1:
                row.cells[1].paragraphs[0].text = title
            if len(row.cells) > 7:
                row.cells[7].paragraphs[0].text = trainer
            if len(row.cells) > 9:
                row.cells[9].paragraphs[0].text = method

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
