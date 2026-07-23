"""Generate pre-job training plan documents from templates."""

from copy import deepcopy
from datetime import date
from io import BytesIO
from pathlib import Path
from typing import Any
from uuid import UUID

from docx import Document
from docx.oxml.ns import qn

TEMPLATE_FILENAME = "新员工培训计划-模板.docx"


def _find_template() -> Path:
    """Locate the docx template, trying several path candidates."""
    candidates = [
        Path("assets/hr") / TEMPLATE_FILENAME,
        Path("../assets/hr") / TEMPLATE_FILENAME,
        Path(__file__).resolve().parent.parent.parent.parent
        / "assets/hr"
        / TEMPLATE_FILENAME,
    ]
    for p in candidates:
        if p.exists():
            return p.resolve()
    raise FileNotFoundError(f"模板文件未找到: {TEMPLATE_FILENAME}")


def _replace_in_paragraph(paragraph, mapping: dict[str, str]) -> None:
    """Replace placeholders in a paragraph, handling markers split across runs."""
    full_text = paragraph.text
    for ph, val in mapping.items():
        if ph in full_text:
            full_text = full_text.replace(ph, val)
    if full_text != paragraph.text:
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = full_text


def _replace_in_cell(cell, mapping: dict[str, str]) -> None:
    """Replace placeholders in all paragraphs of a cell."""
    for paragraph in cell.paragraphs:
        _replace_in_paragraph(paragraph, mapping)


def _cleanup_markers(doc: Document) -> None:
    """Remove template loop markers from all paragraphs."""
    markers = [
        "{#岗位异动培训}", "{/岗位异动培训}",
        "{#培训内容}", "{/培训内容}",
        "{$index+1}",
        "{培训类别}", "{培训师}", "{培训方式}",
    ]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    for marker in markers:
                        text = text.replace(marker, "")
                    if text != paragraph.text:
                        for run in paragraph.runs:
                            run.text = ""
                        if paragraph.runs:
                            paragraph.runs[0].text = text


def generate_prejob_training_plan(
    person: Any,
    training_items: list[dict[str, str]] | None = None,
) -> BytesIO:
    """根据人员数据和培训内容生成岗前培训计划 Word 文档。

    Args:
        person: 人员对象，需有 name, department, position, hire_date,
                education, school, graduation_date, classification, transfer_history
        training_items: 培训内容列表，每项包含 training_category, trainer, training_method
    """
    template_path = _find_template()
    doc = Document(str(template_path))

    def s(val) -> str:
        if val is None:
            return ""
        if isinstance(val, (date, UUID)):
            return str(val)
        return str(val)

    employee_map = {
        "{姓名}": s(getattr(person, "name", None)),
        "{学历}": s(getattr(person, "education", None)),
        "{毕业院校}": s(getattr(person, "school", None)),
        "{毕业时间}": s(getattr(person, "graduation_date", None)),
        "{体现部门}": s(getattr(person, "department", None)),
        "{体现岗位}": s(getattr(person, "position", None)),
        "{入职日期}": s(getattr(person, "hire_date", None)),
        "{证书}": s(getattr(person, "classification", None)),
        "{异动类别}": (
            "岗位变动" if getattr(person, "transfer_history", None)
            else "新员工"
        ),
    }

    # Fill employee placeholders
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_cell(cell, employee_map)

    # Fill training content loop
    if training_items and len(doc.tables) >= 2:
        _fill_training_table(doc, training_items)

    _cleanup_markers(doc)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _fill_training_table(doc: Document, items: list[dict[str, str]]) -> None:
    """Populate Table 1 with training content rows."""
    table = doc.tables[1]

    template_row = None
    for row in table.rows:
        row_text = "".join(c.text for c in row.cells)
        if "{#培训内容}" in row_text or "{培训类别}" in row_text:
            template_row = row
            break
    if template_row is None:
        return

    from docx.oxml import OxmlElement
    from lxml import etree

    # 扫描模板行，按原始文本内容定位每个字段的 XML cell 索引
    xml_cells = template_row._tr.findall(qn('w:tc'))
    idx_cell = None          # {$index+1}
    category_cell = None     # {培训类别}（合并单元格首格）
    date_cell = None         # 空单元格（计划完成期限）
    trainer_cell = None      # {培训师}（合并单元格首格）
    method_cell = None       # {培训方式}

    for ci, tc in enumerate(xml_cells):
        p_text = "".join(t.text or "" for t in tc.iter(qn('w:t')))

        if "{$index+1}" in p_text:
            idx_cell = ci
        elif "{培训类别}" in p_text and category_cell is None:
            category_cell = ci
        elif "{培训师}" in p_text and trainer_cell is None:
            trainer_cell = ci
        elif "{培训方式}" in p_text:
            method_cell = ci

    # 计划完成期限：培训类别和培训师之间的空单元格（取第一个）
    if category_cell is not None and trainer_cell is not None:
        for ci in range(category_cell + 1, trainer_cell):
            p_text = "".join(t.text or "" for t in xml_cells[ci].iter(qn('w:t')))
            if not p_text.strip():
                date_cell = ci
                break

    def _set_cell_text(xml_cell, value: str) -> None:
        for p_elem in xml_cell.findall(qn('w:p')):
            xml_cell.remove(p_elem)
        new_p = OxmlElement('w:p')
        new_r = OxmlElement('w:r')
        new_t = OxmlElement('w:t')
        new_t.text = value
        new_t.set(qn('xml:space'), 'preserve')
        new_r.append(new_t)
        new_p.append(new_r)
        xml_cell.append(new_p)

    for idx, item in enumerate(items):
        new_tr = deepcopy(template_row._tr)
        template_row._tr.addprevious(new_tr)
        new_cells = new_tr.findall(qn('w:tc'))

        if idx_cell is not None:
            _set_cell_text(new_cells[idx_cell], str(idx + 1))
        if category_cell is not None:
            _set_cell_text(new_cells[category_cell], item.get("training_category", ""))
        if date_cell is not None:
            _set_cell_text(new_cells[date_cell], item.get("plan_date", "") or "")
        if trainer_cell is not None:
            _set_cell_text(new_cells[trainer_cell], item.get("trainer", "") or "")
        if method_cell is not None:
            _set_cell_text(new_cells[method_cell], item.get("training_method", "") or "")

    table._tbl.remove(template_row._tr)
