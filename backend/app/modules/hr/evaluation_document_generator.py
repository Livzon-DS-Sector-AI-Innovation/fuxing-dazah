"""培训效果评估表 Word 文档生成器 — 基于模板."""

from datetime import date, datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from pydantic import BaseModel, Field


class TrainingEvaluationInput(BaseModel):
    subject: str = Field(..., max_length=256, description="培训主题")
    training_date: date | None = Field(None)
    training_time_start: str | None = Field(None, max_length=32)
    training_time_end: str | None = Field(None, max_length=32)
    duration_hours: float | None = Field(None)
    training_method: str | None = Field(None, max_length=32)
    trainer: str | None = Field(None, max_length=64)
    trainee_names: list[str] = Field(default_factory=list)
    assessment_method: str | None = Field(None, max_length=32)
    # Post-training fields
    expected_count: int | None = Field(None, description="应到人数")
    actual_count: int | None = Field(None, description="实到人数")
    absent_count: int | None = Field(None)
    sick_leave: int | None = Field(None)
    personal_leave: int | None = Field(None)
    maternity_leave: int | None = Field(None)
    exam_count: int | None = Field(None, description="参加考核人数")
    excellent_count: int | None = Field(None)
    qualified_count: int | None = Field(None)
    unqualified_count: int | None = Field(None)
    participation_rate: str | None = Field(None, max_length=16)
    pass_rate: str | None = Field(None, max_length=16)
    conclusion: str | None = Field(None, max_length=1024)
    organizer: str | None = Field(None, max_length=64)
    organizer_date: date | None = Field(None)


def _find_template() -> Path:
    candidates = [
        Path("assets/hr/7.11培训效果评估表.docx"),
        Path("../assets/hr/7.11培训效果评估表.docx"),
        Path(__file__).resolve().parent.parent.parent.parent
        / "assets/hr"
        / "7.11培训效果评估表.docx",
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError("模板文件未找到: 7.11培训效果评估表.docx")


def _set_cell(cell, text: str) -> None:
    """Set cell text using the first run's formatting."""
    first = None
    for p in cell.paragraphs:
        for r in p.runs:
            if first is None:
                first = r
            r.text = ""
    if first is not None:
        first.text = str(text or "")
    elif cell.paragraphs:
        cell.paragraphs[0].add_run(str(text or ""))


def _compute_hours(start: str | None, end: str | None) -> str:
    if not start or not end:
        return ""
    try:
        s = datetime.strptime(start, "%H:%M")
        e = datetime.strptime(end, "%H:%M")
        diff = (e - s).total_seconds() / 3600
        if diff <= 0:
            return ""
        rounded = round(diff * 2) / 2
        return f"{int(rounded)}h" if rounded == int(rounded) else f"{rounded}h"
    except ValueError:
        return ""


def generate_training_evaluation(data: TrainingEvaluationInput) -> BytesIO:
    """Generate evaluation form from template."""
    doc = Document(str(_find_template()))
    table = doc.tables[0]

    # Row 0: 培训内容 (col 0=label, cols 1-14 merged value)
    _set_cell(table.rows[0].cells[1], data.subject)

    # Row 1: date(cols 1-7), label(cols 8-10), hours(cols 11-14)
    date_str = data.training_date.strftime("%Y.%m.%d") if data.training_date else ""
    _set_cell(table.rows[1].cells[1], date_str)
    hours_str = _compute_hours(data.training_time_start, data.training_time_end)
    if not hours_str and data.duration_hours:
        h = data.duration_hours
        hours_str = f"{int(h)}h" if h == int(h) else f"{h}h"
    _set_cell(table.rows[1].cells[11], hours_str)

    # Row 2: method(cols 1-7), label(cols 8-10), trainer(cols 11-14)
    _set_cell(table.rows[2].cells[1], data.training_method or "")
    _set_cell(table.rows[2].cells[11], data.trainer or "")

    # Row 3: 培训教材 — skip

    # Row 4: 培训对象 (cols 1-14 merged)
    people = "、".join(data.trainee_names) if data.trainee_names else ""
    _set_cell(table.rows[4].cells[1], f"部门/班组/人员(Dept./group/personnel)：{people}")

    # Row 5-7: 应到/实到/缺席 (fill if data available)
    if data.expected_count is not None:
        _set_cell(table.rows[5].cells[3], str(data.expected_count))
        _set_cell(table.rows[6].cells[3], str(data.expected_count))
        _set_cell(table.rows[7].cells[3], str(data.expected_count))
    if data.actual_count is not None:
        _set_cell(table.rows[5].cells[8], str(data.actual_count))
        _set_cell(table.rows[6].cells[8], str(data.actual_count))
        _set_cell(table.rows[7].cells[8], str(data.actual_count))
    if data.sick_leave is not None:
        _set_cell(table.rows[6].cells[14], str(data.sick_leave))
    if data.personal_leave is not None:
        _set_cell(table.rows[5].cells[14], str(data.personal_leave))
    if data.maternity_leave is not None:
        _set_cell(table.rows[7].cells[14], str(data.maternity_leave))

    # Row 8: 考核方式 (cols 1-4) + 参加考核人数 (cols 10-14)
    _set_cell(table.rows[8].cells[1], data.assessment_method or "")
    if data.exam_count is not None:
        _set_cell(table.rows[8].cells[10], str(data.exam_count))

    # Row 11: 考核结果 优/合格/不合格
    if data.excellent_count is not None:
        _set_cell(table.rows[11].cells[3], str(data.excellent_count))
    if data.qualified_count is not None:
        _set_cell(table.rows[11].cells[8], str(data.qualified_count))
    if data.unqualified_count is not None:
        _set_cell(table.rows[11].cells[14], str(data.unqualified_count))

    # Row 15: 培训效果总结评估 (cols 2-14)
    if data.participation_rate or data.pass_rate:
        rate_text = ""
        if data.participation_rate:
            rate_text += f"参训人员参与率{data.participation_rate}"
        if data.pass_rate:
            rate_text += f"，合格率{data.pass_rate}"
        if rate_text:
            _set_cell(table.rows[15].cells[2], f"{rate_text}，达到培训效果。")

    # 最终清理所有残留模板占位符（段落级替换，避免 run 拆分问题）
    import re
    _marker_re = re.compile(r"\{[#/]?[^}]*\}")
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                full_text = "".join(r.text for r in p.runs)
                if _marker_re.search(full_text):
                    cleaned = _marker_re.sub("", full_text)
                    # 把清理后的文本写回第一个 run，清空其余
                    if p.runs:
                        p.runs[0].text = cleaned
                        for r in p.runs[1:]:
                            r.text = ""

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
