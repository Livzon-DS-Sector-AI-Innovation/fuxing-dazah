"""SOP 培训文件登记表 + 各部二级表路由。

登记表（对齐 010版 Excel 模板）：每年一张，每行一个培训文件，按年归档、颜色标注、可导出。
提交登记时按「培训涉及部门」自动生成各部二级表记录，并飞书卡片通知各部门培训管理员。
二级表：转培训自动带出部门当前培训师；培训管理员选择自定义分类人员；
支持多条 SOP 同场一起转训、合并生成一套培训材料。
"""
import json
import logging
import zipfile
from datetime import UTC, date, datetime
from io import BytesIO
from urllib.parse import quote
from uuid import UUID

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from sqlalchemy import select, text
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.database import get_db
from app.core.response import success_response
from app.modules.hr.deps import HrAccessContext, get_hr_scope

logger = logging.getLogger(__name__)

router = APIRouter(tags=["HR-SOP培训"])


# ─── 工具函数 ───


def _ensure_entry_in_scope(hr_scope: HrAccessContext, entry) -> None:
    """二级表记录数据范围校验：记录按部门归属，越界抛 403。"""
    if hr_scope.is_unrestricted:
        return
    if not hr_scope.scoped_departments:
        raise HTTPException(403, "数据范围限制：仅可访问本人相关数据")
    if entry.department not in hr_scope.scoped_departments:
        raise HTTPException(403, f"数据范围限制：仅可操作授权部门记录（{entry.department}）")


def _json_loads(raw) -> list:
    """解析 JSON 数组字段，兼容 list 与 str。"""
    if not raw:
        return []
    if isinstance(raw, list):
        return raw
    try:
        val = json.loads(raw)
        return val if isinstance(val, list) else []
    except Exception:
        return []


async def _lookup_level1_trainer(session: AsyncSession, department: str) -> str | None:
    """查找部门的一级培训师（当前培训师）。

    口径统一：优先取部门培训人员配置表 dept_training_personnel.level1_trainer；
    未配置时兜底取内训师台账 trainers 中 is_level1 的该部门培训师。
    """
    from app.modules.hr.models import DeptTrainingPersonnel, HrTrainer
    result = (await session.execute(
        select(DeptTrainingPersonnel).where(
            DeptTrainingPersonnel.is_deleted == False,  # noqa: E712
            DeptTrainingPersonnel.display_department == department,
        )
    )).scalars().first()
    if result and result.level1_trainer:
        return result.level1_trainer
    trainer = (await session.execute(
        select(HrTrainer).where(
            HrTrainer.is_deleted == False,  # noqa: E712
            HrTrainer.department == department,
            HrTrainer.is_level1.in_(["是", "1", "true", "True"]),
        ).order_by(HrTrainer.created_at.asc())
    )).scalars().first()
    return trainer.name if trainer else None


async def _notify_training_admins(session: AsyncSession, record, departments: list[str]) -> None:
    """发送飞书卡片通知各涉及部门的培训管理员。"""
    try:
        from app.modules.hr.models import DeptTrainingPersonnel
        if not departments:
            return
        result = (await session.execute(
            select(DeptTrainingPersonnel).where(
                DeptTrainingPersonnel.is_deleted == False,  # noqa: E712
                DeptTrainingPersonnel.display_department.in_(departments),
            )
        )).scalars().all()
        admins: set[str] = set()
        for r in result:
            if r.training_admin:
                for a in r.training_admin.split(","):
                    a = a.strip()
                    if a:
                        admins.add(a)
        if not admins:
            return

        from app.modules.hr.feishu_review_service import _lookup_open_id, _send_card
        dept_text = "、".join(departments)
        card = {
            "header": {"title": {"tag": "plain_text", "content": "📋 新的SOP培训文件登记"}, "template": "blue"},
            "elements": [
                {"tag": "div", "text": {"tag": "lark_md", "content": (
                    f"**文件**：{record.file_name}\n"
                    f"**涉及部门**：{dept_text}\n"
                    f"**发起部门**：{record.initiator_department or '-'}\n请在系统中查看并处理"
                )}},
            ],
        }
        for name in admins:
            oid = await _lookup_open_id(name)
            if oid:
                await _send_card(oid, card)
    except Exception:
        logger.exception("通知培训管理员失败: record_id=%s", getattr(record, "id", None))  # 记录失败，便于排查


async def _sync_entries(session: AsyncSession, record) -> None:
    """按登记记录的培训涉及部门同步二级表：新增缺失部门，移除已取消且未转训的部门。"""
    from app.modules.hr.models import SopTrainingEntry
    related = _json_loads(record.involved_departments)
    related = list(dict.fromkeys(str(d).strip() for d in related if str(d).strip()))
    existing_rows = (await session.execute(
        select(SopTrainingEntry).where(
            SopTrainingEntry.is_deleted == False,  # noqa: E712
            SopTrainingEntry.record_id == str(record.id),
        )
    )).scalars().all()
    existing_depts = {e.department for e in existing_rows}
    for dept in related:
        if dept not in existing_depts:
            session.add(SopTrainingEntry(record_id=str(record.id), department=dept))
    for e in existing_rows:
        if e.department not in related and e.status != "已转训":
            e.is_deleted = True


def _record_dict(r) -> dict:
    return {
        "id": str(r.id),
        "year": r.year,
        "training_date": r.training_date,
        "file_name": r.file_name,
        "file_no": r.file_no,
        "effective_date": r.effective_date,
        "method": r.method,
        "complete_time": r.complete_time,
        "trainer": r.trainer,
        "trainees": r.trainees,
        "involved_departments": _json_loads(r.involved_departments),
        "change_note": r.change_note,
        "color": r.color,
        "status": r.status,
        "initiator_department": r.initiator_department,
        "created_by": r.created_by,
        "created_at": str(r.created_at),
        "updated_at": str(r.updated_at),
    }


# ─── 登记表 ───


@router.get("/sop-training-records/years", summary="登记年份列表")
async def list_sop_record_years(
    session: AsyncSession = Depends(get_db),
):
    from app.modules.hr.models import SopTrainingRecord
    rows = (await session.execute(
        select(SopTrainingRecord.year)
        .where(SopTrainingRecord.is_deleted == False)  # noqa: E712
        .distinct()
        .order_by(SopTrainingRecord.year.desc())
    )).scalars().all()
    years = [r for r in rows if r]
    if str(date.today().year) not in years:
        years.insert(0, str(date.today().year))
    return success_response(data=years)


@router.get("/sop-training-records/dept-trainers", summary="按部门查一级培训师（被培训人员）")
async def sop_record_dept_trainers(
    departments: list[str] = Query(..., description="部门列表"),
    session: AsyncSession = Depends(get_db),
):
    """选择培训涉及部门后，自动关联各对应部门的一级培训师（即被培训人员）。"""
    result = []
    for dept in departments:
        trainer = await _lookup_level1_trainer(session, dept)
        result.append({"department": dept, "trainer": trainer})
    return success_response(data=result)


@router.get("/sop-training-records", summary="培训文件登记表列表")
async def list_sop_records(
    year: str | None = Query(None, description="年份筛选"),
    color: str | None = Query(None, description="颜色状态：新增/撤销/修改"),
    keyword: str | None = Query(None, description="文件名称/编号关键词"),
    session: AsyncSession = Depends(get_db),
):
    from app.modules.hr.models import SopTrainingRecord
    stmt = select(SopTrainingRecord).where(SopTrainingRecord.is_deleted == False)  # noqa: E712
    if year:
        stmt = stmt.where(SopTrainingRecord.year == year)
    if color:
        stmt = stmt.where(SopTrainingRecord.color == color)
    if keyword:
        stmt = stmt.where(
            SopTrainingRecord.file_name.ilike(f"%{keyword}%")
            | SopTrainingRecord.file_no.ilike(f"%{keyword}%")
        )
    rows = (await session.execute(
        stmt.order_by(SopTrainingRecord.training_date, SopTrainingRecord.created_at)
    )).scalars().all()
    return success_response(data=[_record_dict(r) for r in rows])


@router.post("/sop-training-records", summary="登记培训文件（保存为草稿）")
async def create_sop_record(
    payload: dict,
    session: AsyncSession = Depends(get_db),
    hr_scope: HrAccessContext = Depends(get_hr_scope),
):
    """登记培训文件，默认保存为草稿；点「提交/通知」后才生成二级表并通知各部门培训管理员。"""
    from app.modules.hr.models import SopTrainingRecord
    if not payload.get("file_name"):
        raise HTTPException(400, "请填写文件名称")
    # 列宽校验：非法值直接 400，而不是落库时 PG DataError 500
    year_val = str(payload.get("year") or date.today().year).strip()
    if len(year_val) != 4 or not year_val.isdigit():
        raise HTTPException(400, "年份格式不正确（应为 4 位数字年份）")
    if payload.get("method") and len(str(payload["method"]).strip()) > 4:
        raise HTTPException(400, "培训方式应为 R / T")
    if payload.get("color") and len(str(payload["color"]).strip()) > 8:
        raise HTTPException(400, "标记颜色值过长")
    involved = list(dict.fromkeys(
        str(d).strip() for d in _json_loads(payload.get("involved_departments")) if str(d).strip()
    ))
    initiator = payload.get("initiator_department") or ""
    r = SopTrainingRecord(
        year=year_val,
        training_date=payload.get("training_date"),
        file_name=payload["file_name"],
        file_no=payload.get("file_no"),
        effective_date=payload.get("effective_date"),
        method=payload.get("method"),
        complete_time=payload.get("complete_time"),
        trainer=payload.get("trainer"),
        trainees=payload.get("trainees") or (f"「{initiator}」全体员工及相关部门培训师" if initiator else None),
        involved_departments=json.dumps(involved, ensure_ascii=False),
        change_note=payload.get("change_note"),
        color=payload.get("color", "新增"),
        # 状态由提交流程控制：创建一律草稿，不允许客户端直接指定状态绕过状态机
        status="草稿",
        initiator_department=initiator or None,
        created_by=hr_scope.user.name or "",
    )
    session.add(r)
    await session.commit()
    return success_response(data={"id": str(r.id)}, message="已保存草稿", status_code=201)


@router.post("/sop-training-records/{record_id}/submit", summary="提交并通知")
async def submit_sop_record(
    record_id: UUID,
    session: AsyncSession = Depends(get_db),
):
    """提交/通知：自动关联到相关部门二级表，并飞书通知对应培训管理员。"""
    from app.modules.hr.models import SopTrainingRecord
    r = await session.get(SopTrainingRecord, record_id)
    if not r or r.is_deleted:
        raise HTTPException(404, "记录不存在")
    involved = _json_loads(r.involved_departments)
    if not involved:
        raise HTTPException(400, "请先填写培训涉及部门")
    was_submitted = r.status == "已提交"
    r.status = "已提交"
    await session.commit()
    await _sync_entries(session, r)
    await session.commit()
    if not was_submitted:
        await _notify_training_admins(session, r, involved)
        return success_response(message="已提交，二级表已生成并通知培训管理员")
    return success_response(message="已提交过，二级表已同步")


@router.put("/sop-training-records/{record_id}", summary="编辑登记记录")
async def update_sop_record(
    record_id: UUID,
    payload: dict,
    session: AsyncSession = Depends(get_db),
):
    from app.modules.hr.models import SopTrainingRecord
    r = await session.get(SopTrainingRecord, record_id)
    if not r or r.is_deleted:
        raise HTTPException(404, "记录不存在")
    # 列宽校验：非法值直接 400，而不是落库时 PG DataError 500
    if "year" in payload:
        year_val = str(payload["year"]).strip()
        if len(year_val) != 4 or not year_val.isdigit():
            raise HTTPException(400, "年份格式不正确（应为 4 位数字年份）")
        payload["year"] = year_val
    if payload.get("method") and len(str(payload["method"]).strip()) > 4:
        raise HTTPException(400, "培训方式应为 R / T")
    if payload.get("color") and len(str(payload["color"]).strip()) > 8:
        raise HTTPException(400, "标记颜色值过长")
    for key in ("year", "training_date", "file_name", "file_no", "effective_date",
                "method", "complete_time", "trainer", "trainees", "change_note",
                "color", "initiator_department"):
        if key in payload:
            setattr(r, key, payload[key])
    # status 由提交/转训流程控制，不允许客户端直接改写
    if "involved_departments" in payload:
        involved = list(dict.fromkeys(
            str(d).strip() for d in _json_loads(payload["involved_departments"]) if str(d).strip()
        ))
        r.involved_departments = json.dumps(involved, ensure_ascii=False)
    await session.commit()
    # 已提交的记录编辑后同步二级表；草稿不生成二级表
    if r.status == "已提交":
        await _sync_entries(session, r)
        await session.commit()
    return success_response(message="已更新")


@router.delete("/sop-training-records/{record_id}", summary="删除登记记录")
async def delete_sop_record(
    record_id: UUID,
    session: AsyncSession = Depends(get_db),
):
    from app.modules.hr.models import SopTrainingEntry, SopTrainingRecord
    r = await session.get(SopTrainingRecord, record_id)
    if not r or r.is_deleted:
        raise HTTPException(404, "记录不存在")
    r.is_deleted = True
    entries = (await session.execute(
        select(SopTrainingEntry).where(
            SopTrainingEntry.is_deleted == False,  # noqa: E712
            SopTrainingEntry.record_id == str(record_id),
        )
    )).scalars().all()
    for e in entries:
        if e.status != "已转训":
            e.is_deleted = True
    await session.commit()
    return success_response(message="已删除")


@router.get("/sop-training-records/export", summary="导出培训文件登记表（对齐模板）")
async def export_sop_records(
    year: str | None = Query(None, description="年份筛选"),
    session: AsyncSession = Depends(get_db),
):
    """导出对齐 010版 登记表版式的 Excel（含颜色标注：黄=新增/灰=撤销/红=修改）。"""
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill

    from app.modules.hr.models import SopTrainingRecord
    stmt = select(SopTrainingRecord).where(SopTrainingRecord.is_deleted == False)  # noqa: E712
    if year:
        stmt = stmt.where(SopTrainingRecord.year == year)
    rows = (await session.execute(
        stmt.order_by(SopTrainingRecord.year, SopTrainingRecord.training_date, SopTrainingRecord.created_at)
    )).scalars().all()

    wb = Workbook()
    ws = wb.active
    ws.title = f"{year}年" if year else "登记表"
    headers = ["培训日期", "文件名称", "文件编号", "生效日期", "培训方式\n（R/T）",
               "R：培训完成时间\nT：培训课时", "培训师", "培训对象", "培训涉及部门", "变更内容"]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    fills = {
        "新增": PatternFill("solid", fgColor="FFF2CC"),
        "撤销": PatternFill("solid", fgColor="D9D9D9"),
        "修改": PatternFill("solid", fgColor="FFC7CE"),
    }

    def _fmt_dot(v: str | None) -> str:
        if not v:
            return ""
        return v.replace("-", ".")

    for r in rows:
        depts = "/".join(_json_loads(r.involved_departments))
        train_date = _fmt_dot(r.training_date)
        # 培训日期显示为月.日（如 01.05）；生效日期显示为 年.月.日
        if len(train_date) >= 8:
            train_date = train_date[5:]
        ws.append([
            train_date, r.file_name or "", r.file_no or "",
            _fmt_dot(r.effective_date), r.method or "", r.complete_time or "",
            r.trainer or "", r.trainees or "", depts, r.change_note or "",
        ])
        fill = fills.get(r.color)
        if fill:
            for cell in ws[ws.max_row]:
                cell.fill = fill
    for col, width in zip("ABCDEFGHIJ", (12, 44, 20, 12, 10, 22, 10, 30, 32, 30)):
        ws.column_dimensions[col].width = width
    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    filename = f"培训文件登记表_{year}年.xlsx" if year else f"培训文件登记表_{date.today().year}年.xlsx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename*=utf-8''{quote(filename)}"},
    )


# ─── 各部二级表 ───


async def _records_by_ids(session: AsyncSession, ids: list[str]) -> dict[str, object]:
    from app.modules.hr.models import SopTrainingRecord
    try:
        uuids = [UUID(i) for i in ids]
    except ValueError:
        return {}
    rows = (await session.execute(
        select(SopTrainingRecord).where(SopTrainingRecord.id.in_(uuids))
    )).scalars().all()
    return {str(r.id): r for r in rows}


def _entry_dict(e, record_map: dict[str, object]) -> dict:
    rec = record_map.get(e.record_id or "")
    personnel = _json_loads(e.personnel)
    return {
        "id": str(e.id),
        "record_id": e.record_id,
        "department": e.department,
        "file_name": rec.file_name if rec else None,
        "file_no": rec.file_no if rec else None,
        "method": rec.method if rec else None,
        "trainer": e.trainer,
        "status": e.status,
        "complete_time": e.complete_time,
        "classification": e.classification,
        "personnel": personnel,
        "personnel_count": len(personnel),
        "transferred_by": e.transferred_by,
        "transferred_at": str(e.transferred_at) if e.transferred_at else None,
        "created_at": str(e.created_at),
        "updated_at": str(e.updated_at),
    }


@router.get("/sop-training-entries", summary="SOP培训二级表列表")
async def list_sop_entries(
    record_id: str | None = Query(None, description="登记记录筛选"),
    department: str | None = Query(None, description="部门筛选"),
    status: str | None = Query(None, description="状态筛选：待转训/已转训"),
    session: AsyncSession = Depends(get_db),
    hr_scope: HrAccessContext = Depends(get_hr_scope),
):
    from app.modules.hr.models import SopTrainingEntry
    stmt = select(SopTrainingEntry).where(SopTrainingEntry.is_deleted == False)  # noqa: E712
    if hr_scope.is_unrestricted:
        if department:
            stmt = stmt.where(SopTrainingEntry.department == department)
    else:
        scoped = hr_scope.scoped_departments
        if not scoped:
            return success_response(data=[])
        stmt = stmt.where(SopTrainingEntry.department.in_(scoped))
        if department:
            if department not in scoped:
                raise HTTPException(403, "无权查看该部门二级表")
            stmt = stmt.where(SopTrainingEntry.department == department)
    if record_id:
        stmt = stmt.where(SopTrainingEntry.record_id == record_id)
    if status:
        stmt = stmt.where(SopTrainingEntry.status == status)
    rows = (await session.execute(stmt.order_by(SopTrainingEntry.created_at.desc()))).scalars().all()
    record_ids = list({e.record_id for e in rows if e.record_id})
    record_map = await _records_by_ids(session, record_ids)
    return success_response(data=[_entry_dict(e, record_map) for e in rows])


@router.put("/sop-training-entries/{entry_id}", summary="编辑二级表（分类/人员/培训师/完成时间）")
async def update_sop_entry(
    entry_id: UUID,
    payload: dict,
    session: AsyncSession = Depends(get_db),
    hr_scope: HrAccessContext = Depends(get_hr_scope),
):
    from app.modules.hr.models import SopTrainingEntry
    e = await session.get(SopTrainingEntry, entry_id)
    if not e or e.is_deleted:
        raise HTTPException(404, "记录不存在")
    _ensure_entry_in_scope(hr_scope, e)
    for key in ("classification", "trainer", "complete_time"):
        if key in payload:
            setattr(e, key, payload[key])
    if "personnel" in payload:
        e.personnel = json.dumps(payload["personnel"], ensure_ascii=False)
    await session.commit()
    return success_response(message="已更新")


async def _do_transfer(session: AsyncSession, e, operator: str) -> None:
    """单条转培训：自动带出该部门当前培训师，并把涉及人员自动关联到该条分类标签。"""
    if e.status == "已转训":
        return
    trainer = await _lookup_level1_trainer(session, e.department)
    e.trainer = trainer or e.trainer
    e.status = "已转训"
    e.transferred_by = operator
    e.transferred_at = datetime.now(UTC)
    await _link_classification_tags(session, e, operator)


async def _link_classification_tags(session: AsyncSession, e, operator: str) -> None:
    """转训自动关联分类：二级表涉及人员的自定义分类写入员工标签（幂等）。

    之后「员工档案-分类」与「二级表按分类拉人」会自动包含这些人。
    """
    import json as _json

    from app.modules.hr.models import EmployeeTag

    if not e.classification:
        return
    personnel = e.personnel or []
    if isinstance(personnel, str):
        try:
            personnel = _json.loads(personnel)
        except ValueError:
            return
    if not isinstance(personnel, list):
        return
    numbers = {
        str(p.get("employee_number") or "").strip()
        for p in personnel
        if isinstance(p, dict) and (p.get("employee_number") or "").strip()
    }
    if not numbers:
        return
    # ON CONFLICT DO NOTHING：与活跃行部分唯一索引配合，并发转训/批量打标不 500
    # （部分索引推断必须显式给出 index_where）
    from sqlalchemy.dialects.postgresql import insert as pg_insert
    for num in numbers:
        await session.execute(
            pg_insert(EmployeeTag)
            .values(employee_number=num, tag_name=e.classification, created_by=operator)
            .on_conflict_do_nothing(
                index_elements=["employee_number", "tag_name"],
                index_where=text("is_deleted = false"),
            ),
        )
    await session.flush()


@router.post("/sop-training-entries/{entry_id}/transfer", summary="转培训")
async def transfer_sop_entry(
    entry_id: UUID,
    session: AsyncSession = Depends(get_db),
    hr_scope: HrAccessContext = Depends(get_hr_scope),
):
    """转培训：自动带出该部门一级培训师为培训师，状态置为已转训。"""
    from app.modules.hr.models import SopTrainingEntry
    e = await session.get(SopTrainingEntry, entry_id)
    if not e or e.is_deleted:
        raise HTTPException(404, "记录不存在")
    _ensure_entry_in_scope(hr_scope, e)
    await _do_transfer(session, e, hr_scope.user.name or "")
    await session.commit()
    return success_response(data={"trainer": e.trainer}, message="已转培训")


@router.post("/sop-training-entries/batch-transfer", summary="批量转培训（多条SOP一起转训）")
async def batch_transfer_sop_entries(
    payload: dict,
    session: AsyncSession = Depends(get_db),
    hr_scope: HrAccessContext = Depends(get_hr_scope),
):
    """勾选多条二级表记录一起转训，每条自动带出本部门当前培训师。"""
    from app.modules.hr.models import SopTrainingEntry
    ids = payload.get("ids") or []
    if not ids:
        raise HTTPException(400, "请选择要转训的记录")
    try:
        entry_ids = [UUID(i) for i in ids]
    except (TypeError, ValueError):
        raise HTTPException(400, "记录 ID 格式不正确")
    rows = (await session.execute(
        select(SopTrainingEntry).where(
            SopTrainingEntry.is_deleted == False,  # noqa: E712
            SopTrainingEntry.id.in_(entry_ids),
        )
    )).scalars().all()
    transferred = 0
    for e in rows:
        if e.status == "已转训":
            continue
        _ensure_entry_in_scope(hr_scope, e)
        await _do_transfer(session, e, hr_scope.user.name or "")
        transferred += 1
    await session.commit()
    return success_response(data={"transferred": transferred}, message=f"已转培训 {transferred} 条")



async def _query_leave_counts(session: AsyncSession, department: str) -> tuple[int, int]:
    """统计部门病假/产假人数（无则 0）；与员工查询口径一致：
    部门匹配 department 或（未分类员工的 actual_department）。"""
    sick = maternity = 0
    try:
        result = await session.execute(
            text(
                "SELECT status, count(*) FROM hr.employees "
                "WHERE (department = :dept OR (department = '未分类' AND actual_department = :dept)) "
                "AND is_deleted = false AND status IN ('病假', '产假') GROUP BY status"
            ),
            {"dept": department},
        )
        for row in result.fetchall():
            if row[0] == "病假":
                sick = row[1]
            elif row[0] == "产假":
                maternity = row[1]
    except Exception:
        pass
    return sick, maternity


async def _lookup_exam_questions(session, file_nos: list[str]) -> list[dict]:
    """按 SOP 编号查找题库题目（试卷打包用），返回 [{file_no, question, answer, score}]。"""
    if not file_nos:
        return []
    from app.modules.hr.models import QuestionBank

    rows = (
        await session.execute(
            select(QuestionBank).where(
                QuestionBank.is_deleted == False,  # noqa: E712
                QuestionBank.file_no.in_(file_nos),
            )
        )
    ).scalars().all()
    return [
        {"file_no": q.file_no or "", "question": q.question, "answer": q.answer, "score": q.score or 0}
        for q in rows
    ]


def _generate_exam_docx(subject: str, questions: list[dict]) -> BytesIO:
    """从零生成培训试卷 docx（题目+参考答案）。"""
    from docx import Document

    doc = Document()
    doc.add_heading(f"培训试卷 — {subject}", level=1)
    total = 0
    for i, q in enumerate(questions, start=1):
        doc.add_paragraph(f"{i}. {q['question']}（{q['score']}分）")
        total += int(q.get("score") or 0)
    if total:
        doc.add_paragraph(f"满分：{total} 分")
    doc.add_heading("参考答案", level=2)
    for i, q in enumerate(questions, start=1):
        doc.add_paragraph(f"{i}. {q.get('answer') or '（无参考答案）'}")
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _generate_register_docx(
    *, subject: str, training_date, department: str, trainer: str, method: str,
    trainee_names: list[str], sick: int, maternity: int,
) -> BytesIO:
    """从零生成培训登记表 docx（信息表+人员名单）。"""
    from docx import Document

    doc = Document()
    doc.add_heading("培训登记表", level=1)
    table = doc.add_table(rows=8, cols=2)
    table.style = "Table Grid"
    info = [
        ("培训主题", subject),
        ("培训日期", str(training_date or "")),
        ("培训部门", department),
        ("培训师", trainer or ""),
        ("培训方式", method or ""),
        ("应到人数", f"{len(trainee_names)} 人"),
        ("病假/产假", f"病假 {sick} 人、产假 {maternity} 人"),
        ("实到人数", "____ 人"),
    ]
    for i, (k, v) in enumerate(info):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    doc.add_heading("受训人员名单", level=2)
    if trainee_names:
        for name in trainee_names:
            doc.add_paragraph(name, style="List Number")
    else:
        doc.add_paragraph("（无）")
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf




@router.post("/sop-training-records/{record_id}/materials", summary="统筹总表一键生成全套培训材料")
async def generate_record_materials(
    record_id: UUID,
    session: AsyncSession = Depends(get_db),
    hr_scope: HrAccessContext = Depends(get_hr_scope),
):
    """按登记表涉及部门生成全套材料 zip：每部门 培训通知+签到表+试卷（有题则含）+培训登记表。"""
    from app.modules.hr.models import SopTrainingRecord

    record = (
        await session.execute(
            select(SopTrainingRecord).where(
                SopTrainingRecord.id == record_id,
                SopTrainingRecord.is_deleted == False,  # noqa: E712
            )
        )
    ).scalar_one_or_none()
    if not record:
        raise HTTPException(404, "登记记录不存在")
    departments = _json_loads(record.involved_departments) or []
    if not departments:
        raise HTTPException(400, "该记录未配置涉及部门")
    # 数据范围：涉及部门必须全部在授权范围内（防止导出其他部门员工名单）
    if not hr_scope.is_unrestricted:
        if not hr_scope.scoped_departments:
            raise HTTPException(403, "数据范围限制：仅可访问本人相关数据")
        out_of_scope = [d for d in departments if d not in hr_scope.scoped_departments]
        if out_of_scope:
            raise HTTPException(403, f"数据范围限制：涉及部门超出授权范围（{', '.join(map(str, out_of_scope))}）")

    from app.modules.hr.notification_document_generator import (
        generate_training_notification,
    )
    from app.modules.hr.schemas import (
        TrainingNotificationInput,
        TrainingSignInSheetInput,
    )
    from app.modules.hr.signin_document_generator import generate_training_sign_in_sheet

    today = date.today()
    subject = record.file_name or "SOP培训"
    trainer = record.trainer
    initiator = (record.initiator_department or "").strip()
    if not initiator:
        raise HTTPException(400, "该记录未配置发起部门（主办部门）")

    # ── 人员口径：主办部门全体在职员工 + 各涉及部门的一级培训师（去重）──
    emp_rows = (await session.execute(
        text(
            "SELECT name, status FROM hr.employees "
            "WHERE (department = :dept OR actual_department = :dept) "
            "AND is_deleted = false AND status NOT IN ('离职', '待审批', '病假', '产假') "
            "ORDER BY name"
        ),
        {"dept": initiator},
    )).fetchall()
    names = [r[0] for r in emp_rows]
    trainer_names: set[str] = set()
    for dept in departments:
        level1 = await _lookup_level1_trainer(session, str(dept).strip())
        if level1:
            trainer_names.add(level1)
    # 主办部门一级培训师作为授课讲师（如涉及部门有则其一；兜底登记表填写的培训师）
    initiator_trainer = await _lookup_level1_trainer(session, initiator)
    lead_trainer = initiator_trainer or (sorted(trainer_names)[0] if trainer_names else (trainer or ""))
    for t in sorted(trainer_names):
        if t and t not in names:
            names.append(t)
    sick, maternity = await _query_leave_counts(session, initiator)

    questions = await _lookup_exam_questions(session, [record.file_no] if record.file_no else [])
    buf = BytesIO()
    try:
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            notif_buf = generate_training_notification(TrainingNotificationInput(
                department=initiator, training_date=today, subject=subject,
                trainer=lead_trainer, content=record.change_note or "",
                trainee_names=names, issuer_department=initiator,
                issue_date=today, sick_count=sick, maternity_count=maternity,
                training_method=record.method or "",
            ))
            sign_buf = generate_training_sign_in_sheet(TrainingSignInSheetInput(
                training_date=today, department=initiator, training_subject=subject,
                topic=record.change_note or subject, instructor=lead_trainer,
                employee_names=names, employee_departments={n: initiator for n in names},
                sick_count=sick, maternity_count=maternity,
            ))
            zf.writestr("培训通知.docx", notif_buf.getvalue())
            zf.writestr("培训签到表.docx", sign_buf.getvalue())
            if questions:
                zf.writestr("培训试卷.docx", _generate_exam_docx(subject, questions).getvalue())
            zf.writestr(
                "培训登记表.docx",
                _generate_register_docx(
                    subject=subject, training_date=today, department=initiator,
                    trainer=lead_trainer, method=record.method or "",
                    trainee_names=names, sick=sick, maternity=maternity,
                ).getvalue(),
            )
    except FileNotFoundError as e:
        raise HTTPException(400, f"生成失败（模板缺失）: {e}")
    buf.seek(0)
    filename = f"SOP全套培训材料_{today.isoformat()}.zip"
    return StreamingResponse(
        buf,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename*=utf-8''{quote(filename)}"},
    )
@router.post("/sop-training-entries/batch-materials", summary="多条SOP生成一套培训材料")
async def batch_generate_materials(
    payload: dict,
    session: AsyncSession = Depends(get_db),
    hr_scope: HrAccessContext = Depends(get_hr_scope),
):
    """勾选二级表记录（1..n 条）生成**一套**培训材料。

    人员 = 各选中条目「分类人员」的并集（条目未选分类人员时直接 400 提示先选人）；
    选中多条时培训内容并列体现多个文件名称，共用同一套通知/签到表/试卷/登记表。
    """
    from app.modules.hr.models import SopTrainingEntry
    ids = payload.get("ids") or []
    if not ids:
        raise HTTPException(400, "请选择记录")
    try:
        entry_ids = [UUID(i) for i in ids]
    except (TypeError, ValueError):
        raise HTTPException(400, "记录 ID 格式不正确")
    rows = (await session.execute(
        select(SopTrainingEntry).where(
            SopTrainingEntry.is_deleted == False,  # noqa: E712
            SopTrainingEntry.id.in_(entry_ids),
        )
    )).scalars().all()
    if not rows:
        raise HTTPException(404, "记录不存在")
    # 数据范围：逐条校验（防止受限用户导出其他部门员工名单）
    for row in rows:
        _ensure_entry_in_scope(hr_scope, row)
    record_map = await _records_by_ids(session, [e.record_id for e in rows if e.record_id])

    # 每条记录必须先选好分类人员（材料人员口径 = 分类人员）
    file_names: list[str] = []
    file_nos: list[str] = []
    names: list[str] = []
    emp_nos: list[str] = []
    for e in rows:
        rec = record_map.get(e.record_id or "")
        fname = (rec.file_name if rec else None) or e.file_name or "SOP培训"
        if fname not in file_names:
            file_names.append(fname)
        if rec and rec.file_no and rec.file_no not in file_nos:
            file_nos.append(rec.file_no)
        personnel = _json_loads(e.personnel)
        if not personnel:
            raise HTTPException(400, f"「{fname}」尚未选择分类人员，请先在该条目选择分类人员后再生成材料")
        for p in personnel:
            nm = str(p.get("name", "")).strip()
            no = str(p.get("employee_number", "")).strip()
            if nm and nm not in names:
                names.append(nm)
                emp_nos.append(no)

    dept = str(rows[0].department or "").strip()
    subject = file_names[0] if len(file_names) == 1 else "；".join(file_names)
    content = "；".join(file_names)

    from app.modules.hr.notification_document_generator import (
        generate_training_notification,
    )
    from app.modules.hr.schemas import (
        TrainingNotificationInput,
        TrainingSignInSheetInput,
    )
    from app.modules.hr.signin_document_generator import generate_training_sign_in_sheet

    today = date.today()
    # 该部门培训师
    trainer = next((e.trainer for e in rows if e.trainer), None) or await _lookup_level1_trainer(session, dept)
    # 分类人员并集（过滤病假/产假/离职/待审批）
    if emp_nos:
        leave_rows = (await session.execute(
            text(
                "SELECT employee_number, status FROM hr.employees "
                "WHERE employee_number = ANY(:nos) AND is_deleted = false "
                "AND status IN ('病假', '产假', '离职', '待审批')"
            ),
            {"nos": emp_nos},
        )).fetchall()
        excluded = {r[0] for r in leave_rows}
        keep = [(no, nm) for no, nm in zip(emp_nos, names) if no not in excluded]
        names = [nm for _, nm in keep]
    sick, maternity = await _query_leave_counts(session, dept)

    buf = BytesIO()
    try:
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            notif_buf = generate_training_notification(TrainingNotificationInput(
                department=dept, training_date=today, subject=subject,
                trainer=trainer or "", content=content, trainee_names=names,
                issuer_department=dept, issue_date=today,
                sick_count=sick, maternity_count=maternity,
            ))
            sign_buf = generate_training_sign_in_sheet(TrainingSignInSheetInput(
                training_date=today, department=dept, training_subject=subject,
                topic=content, instructor=trainer or "",
                employee_names=names, employee_departments={n: dept for n in names},
                sick_count=sick, maternity_count=maternity,
            ))
            zf.writestr("培训通知.docx", notif_buf.getvalue())
            zf.writestr("培训签到表.docx", sign_buf.getvalue())
            # 试卷：按 SOP 编号匹配题库，有题才打包
            questions = await _lookup_exam_questions(session, file_nos)
            if questions:
                zf.writestr("培训试卷.docx", _generate_exam_docx(subject, questions).getvalue())
            # 培训登记表
            method = rows[0].complete_time or ""
            zf.writestr(
                "培训登记表.docx",
                _generate_register_docx(
                    subject=subject, training_date=today, department=dept,
                    trainer=trainer or "", method=method, trainee_names=names,
                    sick=sick, maternity=maternity,
                ).getvalue(),
            )
    except FileNotFoundError as e:
        raise HTTPException(400, f"生成失败（模板缺失）: {e}")
    buf.seek(0)
    filename = f"SOP培训材料_{today.isoformat()}.zip"
    return StreamingResponse(
        buf,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename*=utf-8''{quote(filename)}"},
    )


@router.get("/sop-training-entries/classifications", summary="部门自定义分类选项")
async def sop_entry_classifications(
    department: str = Query(..., description="部门"),
    session: AsyncSession = Depends(get_db),
    hr_scope: HrAccessContext = Depends(get_hr_scope),
):
    """该部门员工被打上的全部标签（所有创建人汇总），作为二级表分类选项。"""
    # 数据范围：受限用户只能查授权部门（员工标签/分类属于员工 PII）
    if not hr_scope.is_unrestricted:
        if not hr_scope.scoped_departments:
            raise HTTPException(403, "数据范围限制：仅可访问本人相关数据")
        if department not in hr_scope.scoped_departments:
            raise HTTPException(403, f"数据范围限制：仅可访问授权部门（{department}）")
    rows = (await session.execute(
        text("""
            SELECT t.tag_name, count(DISTINCT t.employee_number)
            FROM hr.employee_tags t
            JOIN hr.employees e ON e.employee_number = t.employee_number
            WHERE t.is_deleted = false AND e.is_deleted = false
              AND (e.department = :dept OR e.actual_department = :dept)
            GROUP BY t.tag_name ORDER BY t.tag_name
        """),
        {"dept": department},
    )).fetchall()
    return success_response(data=[{"tag_name": r[0], "count": r[1]} for r in rows])


@router.get("/sop-training-entries/personnel", summary="分类人员查询")
async def sop_entry_personnel(
    department: str = Query(..., description="部门"),
    classification: str = Query(..., description="自定义分类"),
    session: AsyncSession = Depends(get_db),
    hr_scope: HrAccessContext = Depends(get_hr_scope),
):
    """按部门 + 自定义分类拉取人员（标签可多打，人员可重复分类）。"""
    # 数据范围：受限用户只能查授权部门（员工姓名/工号/岗位/状态属于 PII）
    if not hr_scope.is_unrestricted:
        if not hr_scope.scoped_departments:
            raise HTTPException(403, "数据范围限制：仅可访问本人相关数据")
        if department not in hr_scope.scoped_departments:
            raise HTTPException(403, f"数据范围限制：仅可访问授权部门（{department}）")
    rows = (await session.execute(
        text("""
            SELECT e.employee_number, e.name, e.position, e.status
            FROM hr.employees e
            JOIN hr.employee_tags t ON t.employee_number = e.employee_number
            WHERE t.is_deleted = false AND e.is_deleted = false
              AND t.tag_name = :tag
              AND (e.department = :dept OR e.actual_department = :dept)
              AND e.status NOT IN ('离职', '待审批', '病假', '产假')
            ORDER BY e.employee_number
        """),
        {"dept": department, "tag": classification},
    )).fetchall()
    return success_response(data=[
        {"employee_number": r[0], "name": r[1], "position": r[2], "status": r[3]}
        for r in rows
    ])


@router.get("/sop-training-entries/export", summary="导出培训清单")
async def export_sop_entries(
    record_id: str | None = Query(None, description="登记记录筛选"),
    department: str | None = Query(None, description="部门筛选"),
    status: str | None = Query(None, description="状态筛选：待转训/已转训"),
    session: AsyncSession = Depends(get_db),
    hr_scope: HrAccessContext = Depends(get_hr_scope),
):
    """导出二级表培训清单（xlsx）。"""
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font

    from app.modules.hr.models import SopTrainingEntry
    stmt = select(SopTrainingEntry).where(SopTrainingEntry.is_deleted == False)  # noqa: E712
    if hr_scope.is_unrestricted:
        if department:
            stmt = stmt.where(SopTrainingEntry.department == department)
    else:
        scoped = hr_scope.scoped_departments
        if not scoped:
            raise HTTPException(403, "数据范围限制：无可见部门")
        stmt = stmt.where(SopTrainingEntry.department.in_(scoped))
        if department:
            if department not in scoped:
                raise HTTPException(403, "无权导出该部门清单")
            stmt = stmt.where(SopTrainingEntry.department == department)
    if record_id:
        stmt = stmt.where(SopTrainingEntry.record_id == record_id)
    if status:
        stmt = stmt.where(SopTrainingEntry.status == status)
    rows = (await session.execute(
        stmt.order_by(SopTrainingEntry.department, SopTrainingEntry.created_at)
    )).scalars().all()
    record_map = await _records_by_ids(session, [e.record_id for e in rows if e.record_id])

    wb = Workbook()
    ws = wb.active
    ws.title = "SOP培训清单"
    headers = ["部门", "文件名称", "文件编号", "培训方式", "培训师", "状态", "完成时间/课时", "自定义分类", "分类人员", "转培训人", "转培训时间"]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center")
    for e in rows:
        rec = record_map.get(e.record_id or "")
        names = "、".join(str(p.get("name", "")) for p in _json_loads(e.personnel) if p.get("name"))
        ws.append([
            e.department,
            rec.file_name if rec else "",
            rec.file_no if rec else "",
            rec.method if rec else "",
            e.trainer or "", e.status, e.complete_time or "",
            e.classification or "", names, e.transferred_by or "",
            str(e.transferred_at)[:19] if e.transferred_at else "",
        ])
    for col, width in zip("ABCDEFGHIJK", (18, 44, 20, 10, 14, 10, 20, 16, 40, 12, 20)):
        ws.column_dimensions[col].width = width
    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    filename = f"SOP培训清单_{date.today().isoformat()}.xlsx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename*=utf-8''{quote(filename)}"},
    )
