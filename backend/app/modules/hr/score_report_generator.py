"""Generate score report (成绩单) from template."""

from copy import deepcopy
from io import BytesIO
from pathlib import Path

from docx import Document


def _find_template() -> Path:
    candidates = [
        Path("assets/hr/成绩单模板.docx"),
        Path("../assets/hr/成绩单模板.docx"),
        Path(__file__).resolve().parent.parent.parent.parent / "assets/hr" / "成绩单模板.docx",
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError("模板文件未找到: 成绩单模板.docx")


def _set_cell_text(cell, text: str) -> None:
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    if cell.paragraphs:
        cell.paragraphs[0].add_run(str(text or ""))


def generate_score_report(
    *,
    training_content: str = "",
    training_date: str = "",
    department: str = "",
    scores: list[dict] | None = None,
) -> BytesIO:
    """生成成绩单，按人数自动追加行。scores: [{name, department, total_score}]"""
    doc = Document(str(_find_template()))
    table = doc.tables[0]
    scores = scores or []
    template_row_idx = 1  # 第一行数据行（跳过表头）作为样式模板

    if len(table.rows) <= template_row_idx:
        raise ValueError("成绩单模板至少需要一行数据行（不含表头）")

    # 人数超过模板行数时，自动追加行
    while len(table.rows) - 1 < len(scores):
        template_row = table.rows[template_row_idx]
        new_row = deepcopy(template_row._tr)
        table._tbl.append(new_row)

    for ri, s in enumerate(scores):
        row = table.rows[ri + 1]
        _set_cell_text(row.cells[0], str(ri + 1))
        _set_cell_text(row.cells[1], s.get("name", ""))
        _set_cell_text(row.cells[2], s.get("department", ""))
        _set_cell_text(row.cells[3], str(s.get("total_score", "")))

    # 清空多余行
    for ri in range(len(scores) + 1, len(table.rows)):
        for cell in table.rows[ri].cells:
            _set_cell_text(cell, "")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
