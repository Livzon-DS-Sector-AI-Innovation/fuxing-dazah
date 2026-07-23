"""个人培训登记表 Word 文档生成器

模板为 docxtemplater 风格的单页表格（{#台账}...{/台账} 包裹），
每名员工生成一页登记表，整个部门合并输出为一个 .docx 文件。
"""

import re
from copy import deepcopy
from io import BytesIO

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.modules.hr.template_utils import find_hr_template

# 简单文本占位符 -> 记录字段
_PLACEHOLDERS = ["姓名", "性别", "体现部门", "体现岗位", "学历", "毕业院校", "专业", "证书"]
# 带 customDate 过滤器的日期占位符，如 {入职日期| customDate:'YYYY-MM-DD'}
_DATE_PLACEHOLDERS = ["毕业时间", "入职日期"]




def _fill_paragraph(paragraph, record: dict) -> None:
    """合并 run 后整体替换占位符（占位符可能被 Word 拆分到多个 run）。"""
    if not paragraph.runs:
        return
    full = "".join(r.text for r in paragraph.runs)
    if "{" not in full:
        return
    for key in _PLACEHOLDERS:
        full = full.replace("{" + key + "}", record.get(key, ""))
    for key in _DATE_PLACEHOLDERS:
        full = re.sub(r"\{" + key + r"\|[^}]*\}", record.get(key, ""), full)
    paragraph.runs[0].text = full
    for r in paragraph.runs[1:]:
        r.text = ""


def _fill_table(table, record: dict) -> None:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _fill_paragraph(p, record)


def _page_break_paragraph():
    """构造仅含分页符的 w:p 元素。"""
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    return p


def generate_training_registration_sync(records: list[dict]) -> BytesIO:
    """生成个人培训登记表。records: 每名员工一个占位符字段 dict。"""
    doc = Document(str(find_hr_template("个人培训登记表.docx")))
    if not doc.tables or not records:
        return _save(doc)

    # 移除 {#台账} / {/台账} 循环标记段落
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if text in ("{#台账}", "{/台账}"):
            paragraph._element.getparent().remove(paragraph._element)

    # 先保留一份未填充的表格副本，再逐人克隆填充
    source_tbl = doc.tables[0]._tbl
    pristine = deepcopy(source_tbl)

    _fill_table(doc.tables[0], records[0])

    last_el = source_tbl
    for record in records[1:]:
        brk = _page_break_paragraph()
        last_el.addnext(brk)
        new_tbl = deepcopy(pristine)
        brk.addnext(new_tbl)
        last_el = new_tbl
        _fill_table(doc.tables[-1], record)

    return _save(doc)


def _save(doc) -> BytesIO:
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
