"""职称评审业务编排（v3：审批前置 + 内网投票）。

流程：员工飞书审批申报（部门负责人→HR 两道，飞书原生）→ 审批通过自动写入
申报表 → 系统同步落库 → 评委登录内网系统投票（匿名，7 维评价自动计算
综合等级与投票结果）→ 票数判定（同意÷(同意+不同意)≥2/3）。
小组评审结果为「评审合格/未通过」，最终名单经总经理确认后由 HR 公示，
系统不直接通知申报人。
"""

import asyncio
import base64
import json
import logging
import uuid
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any
from uuid import UUID

from fastapi import HTTPException
from sqlalchemy.exc import IntegrityError
from sqlalchemy.ext.asyncio import AsyncSession

from app.modules.hr.models import Employee
from app.modules.hr.title_review import models as m
from app.modules.hr.title_review.repository import (
    TitleReviewActivityRepository,
    TitleReviewApplicationRepository,
    TitleReviewDeptCommitteeRepository,
    TitleReviewDimensionRepository,
    TitleReviewJudgeRepository,
    TitleReviewLevelRepository,
    TitleReviewScoreRepository,
)
from app.modules.hr.title_review.schemas import (
    TitleReviewActivityCreate,
    TitleReviewActivityUpdate,
    TitleReviewApplicationOut,
    TitleReviewDeptCommitteeIn,
    TitleReviewJudgeAssignIn,
    TitleReviewJudgeOut,
    TitleReviewLevelIn,
    TitleReviewScoreOut,
)
from app.platform.audit.service import record_audit_log

logger = logging.getLogger(__name__)

# ─── 申报表字段映射（与原型一致） ───
APPLY_SELF_EVAL_FIELDS = [
    "岗位任务自我评价",
    "工作思想表现自我评价",
    "组织项目自我评价",
    "科技成果自我评价",
    "专利论文自我评价",
    "合理化建议自我评价",
    "培养指导自我评价",
]
APPLY_WORK_FIELDS = [
    "岗位规定的职责任务完成情况",
    "组织处理技术项目及效果",
    "专业工作合理化建议及采纳情况",
    "任现职以来主要专业技术/职业技能工作业绩",
    "专利/论文/著作/总结/报告及发表情况",
    "科技/技改/管理/研究成果及奖励情况",
    "培养指导专业技术人员学习工作情况",
    "工作思想表现及执行政策水平",
]
APPLY_ATTACHMENT_FIELDS = [
    "职称评审申报表",
    "业绩成果证明材料",
    "论文论著专利",
    "外部职称证书",
]
# 人工新建审批定义（大修版表单）字段 → 申报表列名；未列出的字段按原名写入。
# 多个表单字段映射到同一列时按表单顺序以换行合并。
APPROVAL_FORM_FIELD_MAP = {
    # 申报职级（技术/职业两个序列控件）
    "本次申报职级类型": "申报序列",
    "本次申报职级（技术）": "申报职级",
    "本次申报职级（技能）": "申报职级",
    "本次申报职级（职业）": "申报职级",
    "是否破格": "是否破格申报",
    # 7 项自评（附表4 维度列）
    "岗位规定职责自我评价": "岗位任务自我评价",
    "工作思想自我评价": "工作思想表现自我评价",
    "科技/管理/研究成果自我评价": "科技成果自我评价",
    "组织处理技术项目自我评价": "组织项目自我评价",
    "工作合理化建议自我评价": "合理化建议自我评价",
    "论文，著作等内容自我评价": "专利论文自我评价",
    "培养技术人员相关自我评价": "培养指导自我评价",
    # 8 项业绩陈述（同名列，名称微调过的在此对齐）
    "任现职以来撰写的专利、论文、著作、总结、报告": "专利/论文/著作/总结/报告及发表情况",
    "专利、论文、著作、总结、报告（本人负责部分）以及授权、刊出及交流情况 （时间、刊物或会议名称）": "专利/论文/著作/总结/报告及发表情况",
    "岗位规定的职责任务以及完成情况 （数量、质量、效益）": "岗位规定的职责任务完成情况",
    "科技、技改、管理、研究成果以及鉴定或奖励情况 （时间、等级、名次）": "科技/技改/管理/研究成果及奖励情况",
    "组织处理技术项目以及效果 （完成时间、结果）": "组织处理技术项目及效果",
    "对专业工作合理化建议内容以及建议采纳情况": "专业工作合理化建议及采纳情况",
    "培养、指导专业技术人员学习、工作情况": "培养指导专业技术人员学习工作情况",
    "工作思想表现及执行上级政策水平": "工作思想表现及执行政策水平",
}
# 申报表只读列类型：lookup(19)/formula(20)/系统自动列(1001-1005)
READONLY_FIELD_TYPES = {19, 20, 1001, 1002, 1003, 1004, 1005}
# 单选列类型（飞书 field type=3），写入前做选项校验
FEISHU_FIELD_SINGLE_SELECT = 3
# 审批表单图片证明控件（image 类型）：值为 URL 列表，落库为文本列（旧表单名保留兼容）
IMAGE_EVIDENCE_FIELDS = [
    "外部专业技术职称证书等证明材料上传",
    "证明材料上传（图片）",
    "参加过两项以上本专业或相关专业技术工作、技术管理，技术服务工作的业绩证明材料",
    "两项以上担任项目技术负责人的业绩证明材料",
]
# 员工信息表自动带出的个人档案字段（申报表/审批表单未体现的信息由此补充展示）
# 评定职级按年度滚动：近 5 个已评定年份 + 当前年度最高可申报
def _build_profile_fields() -> list[str]:
    year = datetime.now().year
    recent_years = [f"{y}年评定职级" for y in range(year - 5, year)]
    return [
        "学历",
        "司龄",
        "入职日期",
        "性别",
        "职务",
        "岗位职级",
        "毕业院校",
        "专业",
        "目前职级",
        *recent_years,
        "近5年年终绩效考评结果",
        f"{year}年最高可申报（根据年限）",
    ]


EMPLOYEE_PROFILE_FIELDS = _build_profile_fields()

# 员工信息表 table_id 缓存（app_token → table_id，绑定表格时失效）
_employee_table_cache: dict[str, str] = {}
# 审批镜像表名（飞书审批自动化把每次提交镜像进 Base 的这张表，含申请状态与 SourceID）
APPROVAL_TABLE_NAME = "福兴医药职称审批"
# 审批镜像表 table_id 缓存（app_token → table_id，绑定表格时失效）
_approval_table_cache: dict[str, str] = {}


def _as_str(value: Any) -> str | None:
    """飞书字段值转字符串（select/lookup 可能返回 list）。

    lookup 列值为 [{"text": "...", "type": "text"}] 结构，取 text 拼接；
    单选/文本等直接转字符串。
    """
    if value is None:
        return None
    if isinstance(value, str):
        return value
    if isinstance(value, list):
        parts: list[str] = []
        for v in value:
            if isinstance(v, dict):
                text = v.get("text")
                if text is not None:
                    parts.append(str(text))
            else:
                parts.append(str(v))
        return "".join(parts) if parts else None
    return str(value)


def _as_bool_text(value: Any) -> bool:
    return _as_str(value) == "是"


def _image_list_to_text(value: Any) -> str | None:
    """审批图片控件值（[{url,name,...}]）→ 「名称 链接」逐行文本。"""
    if not isinstance(value, list):
        return None
    lines: list[str] = []
    for item in value:
        if isinstance(item, dict):
            name = str(item.get("name") or "")
            url = str(item.get("url") or "")
            lines.append(f"{name} {url}".strip())
        else:
            lines.append(str(item))
    return "\n".join(line for line in lines if line) or None


def _cell_has_remote_url(text: str) -> bool:
    """图片单元格中是否还有飞书远程链接（未转存本地）。"""
    return "http" in text


async def _download_image_bytes(url: str) -> bytes:
    """下载图片（飞书签名链接有效期内可下载）——走平台集成层下载。"""
    from app.platform.integrations.feishu.http import download

    return await download(url)


async def _image_text_to_local(text: str) -> str:
    """图片证据文本：把飞书远程链接下载转存本地 /uploads，返回本地路径行。

    下载失败保留原链接（下次对账重试，签名有效期内仍可转存）。
    """
    from app.core.config import get_settings

    base_dir = Path(get_settings().UPLOAD_DIR) / "title_review"
    out: list[str] = []
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        idx = line.rfind("http")
        if idx < 0:
            out.append(line)
            continue
        url = line[idx:].strip()
        try:
            data = await _download_image_bytes(url)
        except Exception as exc:  # noqa: BLE001
            logger.warning("图片转存下载失败，保留原链接: %s", exc)
            out.append(line)
            continue
        ext = ".jpg"
        for candidate in (".png", ".jpg", ".jpeg", ".gif", ".webp"):
            if candidate in url.split("?")[0].lower():
                ext = candidate
                break
        base_dir.mkdir(parents=True, exist_ok=True)
        fname = f"{uuid.uuid4().hex}{ext}"
        await asyncio.to_thread((base_dir / fname).write_bytes, data)
        out.append(f"/uploads/title_review/{fname}")
    return "\n".join(out)


def _as_ts(value: Any) -> datetime | None:
    """datetime 字段（毫秒时间戳 int）→ datetime。"""
    if value is None:
        return None
    try:
        return datetime.fromtimestamp(int(value) / 1000).astimezone()
    except (TypeError, ValueError, OSError):
        return None


def _parse_attachments(raw: Any) -> list[dict[str, Any]] | None:
    """附件字段 → 元数据列表 [{file_token,name,size}]。"""
    if not isinstance(raw, list):
        return None
    result: list[dict[str, Any]] = []
    for item in raw:
        if isinstance(item, dict):
            result.append(
                {
                    "file_token": item.get("file_token", ""),
                    "name": item.get("name", ""),
                    "size": item.get("size", 0),
                }
            )
    return result or None




def compute_comprehensive_grade(dimension_grades: list[str | None]) -> str | None:
    """按《员工职级评定管理办法》附表5 自动计算综合等级（两档）。

    规则：合格=全维度中≥5项合格；不合格=超过2项不合格；
    维度未填齐或出现非法值则不计算（返回 None）。
    """
    total = len(m.DEFAULT_DIMENSION_NAMES)
    if len(dimension_grades) < total or any(g not in ("合格", "不合格") for g in dimension_grades):
        return None
    qualified = dimension_grades.count("合格")
    return "合格" if qualified >= 5 else "不合格"


def decide_by_votes(agree: int, oppose: int, pass_ratio: float) -> bool:
    """票数判定：同意÷(同意+不同意) ≥ pass_ratio（弃权不计分母；有同意且无反对即通过）。

    带 1e-9 容差：pass_ratio 允许 0.6667 这类四舍五入值，恰好 2/3 时判定通过。
    """
    voted = agree + oppose
    if voted == 0:
        return False
    return agree / voted >= pass_ratio - 1e-9


async def _audit(
    db: AsyncSession,
    *,
    action: str,
    user: Any = None,
    resource_type: str,
    resource_id: UUID | None = None,
    old_value: dict[str, Any] | None = None,
    new_value: dict[str, Any] | None = None,
    extra: dict[str, Any] | None = None,
) -> None:
    """审计落库（容错包装：审计失败不阻断业务）。"""
    try:
        await record_audit_log(
            db,
            action=action,
            user=user,
            resource_type=resource_type,
            resource_id=resource_id,
            old_value=old_value,
            new_value=new_value,
            extra=extra,
        )
    except Exception:
        logger.exception("审计日志记录失败: action=%s", action)


def _apply_fields_to_dict(fields: dict[str, Any]) -> dict[str, Any]:
    """原型申报表字段 → 内网结构化数据。"""
    self_evaluations: dict[str, str] = {}
    for k in APPLY_SELF_EVAL_FIELDS:
        val = _as_str(fields.get(k))
        if val is not None:
            self_evaluations[k] = val
    work_statements: dict[str, str] = {}
    for k in APPLY_WORK_FIELDS:
        text_val = _as_str(fields.get(k))
        if text_val is not None:
            work_statements[k] = text_val
    perf = _as_str(fields.get("近五年年终绩效考核结果"))
    if perf:
        work_statements["近五年年终绩效考核结果"] = perf
    resume = _as_str(fields.get("任职以来工作简历"))
    if resume:
        work_statements["任职以来工作简历"] = resume
    external_cert = _as_str(
        fields.get("是否具备外部专业技术职称或职业/执业技能证书")
    )
    if external_cert:
        work_statements["是否具备外部专业技术职称或职业/执业技能证书"] = external_cert
    for img_key in IMAGE_EVIDENCE_FIELDS:
        img_text = _as_str(fields.get(img_key))
        if img_text:
            work_statements[img_key] = img_text
    attachments: dict[str, list[dict[str, Any]]] = {}
    for k in APPLY_ATTACHMENT_FIELDS:
        attach_val = _parse_attachments(fields.get(k))
        if attach_val is not None:
            attachments[k] = attach_val
    return {
        "employee_no": (_as_str(fields.get("工号")) or "").strip(),
        "name": (_as_str(fields.get("姓名")) or "").strip(),
        "sequence": _as_str(fields.get("申报序列")),
        "tech_domain": _as_str(fields.get("申报领域")),
        "apply_level": _as_str(fields.get("申报职级")),
        "current_level": _as_str(fields.get("现任职级")),
        "is_exception": _as_bool_text(fields.get("是否破格申报")),
        "exception_reason": _as_str(fields.get("破格申报理由")),
        "tenure_start": _as_ts(fields.get("任现职开始时间")),
        "tenure_end": _as_ts(fields.get("任现职结束时间")),
        "self_evaluations": self_evaluations or None,
        "work_statements": work_statements or None,
        "attachments": attachments or None,
        "approval_instance_code": _as_str(fields.get("审批实例编号")),
    }


TERMINAL_APPLICATION_STATUSES = (
    m.APPLICATION_PASSED,
    m.APPLICATION_FAILED,
    m.APPLICATION_FINAL_PASSED,
    m.APPLICATION_FINAL_FAILED,
)


_SYNCED_FIELD_ATTRS: list[tuple[str, str]] = [
    ("employee_no", "employee_no"),
    ("name", "name"),
    ("sequence", "sequence"),
    ("tech_domain", "tech_domain"),
    ("apply_level", "apply_level"),
    ("current_level", "current_level"),
    ("self_evaluations", "self_evaluations"),
    ("work_statements", "work_statements"),
    ("attachments", "attachments"),
    ("tenure_start", "tenure_start"),
    ("tenure_end", "tenure_end"),
]


def _apply_synced_fields(
    application: m.TitleReviewApplication, parsed: dict[str, Any], fields: dict[str, Any]
) -> bool:
    """把飞书解析值同步到申报对象（事件/对账共用），返回是否有实际变化。

    空值保留旧值（Base 清空不覆盖内网）；「是否破格申报」例外：该列一旦
    出现在变更字段里即整体覆盖（改「否」为「是」或清空都需如实体现）。
    """
    changed = False
    for attr, key in _SYNCED_FIELD_ATTRS:
        new_value = parsed[key] or getattr(application, attr)
        if new_value != getattr(application, attr):
            setattr(application, attr, new_value)
            changed = True
    if "是否破格申报" in fields:
        if parsed["is_exception"] != application.is_exception:
            application.is_exception = parsed["is_exception"]
            changed = True
        if parsed["exception_reason"] != application.exception_reason:
            application.exception_reason = parsed["exception_reason"]
            changed = True
    code = parsed["approval_instance_code"]
    if code and code != application.approval_instance_code:
        # 重复提交时飞书行编号累积，同步到内网
        application.approval_instance_code = code
        changed = True
    return changed


def _filter_writable_row(
    row: dict[str, Any], table_fields: list[dict[str, Any]]
) -> dict[str, Any]:
    """按申报表实际列过滤写入行：仅保留可写列，单选列值不在选项中则丢弃。

    lookup/formula/系统列只读不写（如工号/学历/司龄/现任职级/近五年绩效考核
    均为 lookup，由「姓名」写入后自动从员工信息表带出）。
    """
    writable: dict[str, set[str]] = {}
    for field in table_fields:
        name = field.get("field_name") or ""
        if not name or field.get("type") in READONLY_FIELD_TYPES:
            continue
        options: set[str] = set()
        if field.get("type") == FEISHU_FIELD_SINGLE_SELECT:  # 单选列做选项校验
            options = {
                str(opt.get("name", ""))
                for opt in (field.get("property") or {}).get("options") or []
            }
        writable[name] = options
    filtered: dict[str, Any] = {}
    for key, value in row.items():
        if key not in writable:
            continue
        if writable[key]:
            text = _as_str(value)
            if text is None or text not in writable[key]:
                logger.warning(
                    "申报表单选列值不在选项中，跳过写入: col=%s value=%s", key, value
                )
                continue
        filtered[key] = value
    return filtered


def _extract_approval_code_from_source_id(source_id: str) -> str | None:
    """审批镜像表 SourceID 解码出审批实例编号（与审批 API 实例编号格式对齐）。

    SourceID 形如 base64("{instance_id}:{实例编号}-{版本}:{suffix}")，
    解码后取第二段并去掉末尾「-N」版本后缀。
    """
    try:
        decoded = base64.b64decode(source_id).decode("utf-8")
    except Exception:  # noqa: BLE001
        return None
    parts = decoded.split(":")
    if len(parts) < 2:
        return None
    code = parts[1].strip()
    head, _, tail = code.rpartition("-")
    if tail.isdigit():
        code = head
    return code or None


def _resume_date_text(value: Any) -> str:
    """镜像表日期值（毫秒时间戳/字符串）→ YYYY-MM-DD 文本。"""
    text = _as_str(value) or ""
    if text.isdigit() and len(text) >= 10:
        try:
            return datetime.fromtimestamp(int(text) / 1000).strftime("%Y-%m-%d")
        except (ValueError, OSError):
            return text
    return text


def _merge_resume_columns(fields: dict[str, Any]) -> None:
    """审批镜像表把「任职以来工作简历」fieldList 拆成独立列，还原为原文本格式。"""
    start = _resume_date_text(fields.get("任职以来工作简历_开始时间"))
    end = _resume_date_text(fields.get("任职以来工作简历_结束时间"))
    dept = _as_str(fields.get("任职以来工作简历_部门/担任职位"))
    lines: list[str] = []
    if start or end:
        lines.append(f"任职时间：{start} ~ {end}".strip(" ~"))
    if dept:
        lines.append(f"部门/担任职位：{dept}")
    if lines:
        fields["任职以来工作简历"] = "\n".join(lines)


class TitleReviewService:
    """职称评审编排（活动/职级组/部门评审组/申报/投票/流程）。"""

    def __init__(self, session: AsyncSession) -> None:
        self.session = session
        self.activity_repo = TitleReviewActivityRepository(session)
        self.level_repo = TitleReviewLevelRepository(session)
        self.dimension_repo = TitleReviewDimensionRepository(session)
        self.application_repo = TitleReviewApplicationRepository(session)
        self.judge_repo = TitleReviewJudgeRepository(session)
        self.score_repo = TitleReviewScoreRepository(session)
        self.committee_repo = TitleReviewDeptCommitteeRepository(session)
        # 数据范围（由路由层按 HrAccessContext 注入）：部门集合 或 本人工号
        self._scope_departments: set[str] | None = None
        self._scope_employee_no: str | None = None

    def set_scope(
        self, departments: set[str] | None, employee_no: str | None
    ) -> None:
        """设置数据范围：departments 非 None 时按部门过滤；否则按 employee_no 过滤本人。"""
        self._scope_departments = departments
        self._scope_employee_no = employee_no

    def _scope_applications(
        self, applications: list[m.TitleReviewApplication]
    ) -> list[m.TitleReviewApplication]:
        if self._scope_departments is not None:
            return [
                a for a in applications if a.department in self._scope_departments
            ]
        if self._scope_employee_no is not None:
            return [
                a for a in applications if a.employee_no == self._scope_employee_no
            ]
        return applications

    # ═══ 活动 CRUD ═══

    async def create_activity(
        self, data: TitleReviewActivityCreate, user: Any = None
    ) -> m.TitleReviewActivity:
        activity = m.TitleReviewActivity(
            name=data.name,
            status=m.ACTIVITY_DRAFT,
            apply_deadline=data.apply_deadline,
            review_deadline=data.review_deadline,
            pass_ratio=data.pass_ratio,
            feishu_app_token=data.feishu_app_token,
            apply_table_id=data.apply_table_id,
            vote_table_id=data.vote_table_id,
            approval_code=data.approval_code,
        )
        await self.activity_repo.create(activity)
        levels = data.levels or [
            TitleReviewLevelIn(**tpl) for tpl in m.DEFAULT_LEVEL_TEMPLATE
        ]
        for i, lv in enumerate(levels, start=1):
            await self.level_repo.create(
                m.TitleReviewLevel(
                    activity_id=activity.id,
                    sequence=lv.sequence,
                    level_name=lv.level_name,
                    sort_order=i,
                )
            )
        # 固定 7 个评价项（投票表列名与评价项同名）
        for i, name in enumerate(m.DEFAULT_DIMENSION_NAMES, start=1):
            await self.dimension_repo.create(
                m.TitleReviewDimension(
                    activity_id=activity.id,
                    name=name,
                    feishu_field_name=name,
                    sort_order=i,
                )
            )
        await _audit(
            self.session,
            action="hr.title.activity_create",
            user=user,
            resource_type="title_review_activity",
            resource_id=activity.id,
            new_value={"name": activity.name, "pass_ratio": activity.pass_ratio},
        )
        return activity

    async def get_activity(self, activity_id: UUID) -> m.TitleReviewActivity:
        activity = await self.activity_repo.get_by_id(activity_id)
        if not activity:
            raise HTTPException(404, "评定活动不存在")
        return activity

    async def list_activities_with_progress(
        self, *, status: str | None, keyword: str | None, page: int, page_size: int
    ) -> tuple[list[m.TitleReviewActivity], int, dict[UUID, int], dict[UUID, int], dict[UUID, int]]:
        """活动分页 + 申报数/评委总数/已投票评委数（分组查询避免 N+1）。"""
        from sqlalchemy import func
        from sqlalchemy import select as sa_select

        activities, total = await self.activity_repo.list_all(
            status=status, keyword=keyword, page=page, page_size=page_size
        )
        ids = [a.id for a in activities]
        app_counts: dict[UUID, int] = {}
        voted_counts: dict[UUID, int] = {}
        total_judge_counts: dict[UUID, int] = {}
        if ids:
            app_rows = (
                await self.session.execute(
                    sa_select(m.TitleReviewApplication.activity_id, func.count())
                    .where(
                        m.TitleReviewApplication.activity_id.in_(ids),
                        m.TitleReviewApplication.is_deleted == False,  # noqa: E712
                    )
                    .group_by(m.TitleReviewApplication.activity_id)
                )
            ).all()
            app_counts = {row[0]: row[1] for row in app_rows}
            judge_rows = (
                await self.session.execute(
                    sa_select(
                        m.TitleReviewJudge.activity_id,
                        func.count(),
                        func.count().filter(m.TitleReviewJudge.vote_result.isnot(None)),
                    )
                    .where(
                        m.TitleReviewJudge.activity_id.in_(ids),
                        m.TitleReviewJudge.is_deleted == False,  # noqa: E712
                    )
                    .group_by(m.TitleReviewJudge.activity_id)
                )
            ).all()
            total_judge_counts = {row[0]: row[1] for row in judge_rows}
            voted_counts = {row[0]: row[2] for row in judge_rows}
        return activities, total, app_counts, voted_counts, total_judge_counts

    async def update_activity(
        self, activity_id: UUID, data: TitleReviewActivityUpdate, user: Any = None
    ) -> m.TitleReviewActivity:
        activity = await self.get_activity(activity_id)
        old = {"name": activity.name, "pass_ratio": activity.pass_ratio}

        if data.name is not None and data.name != activity.name:
            self._ensure_draft(activity, "活动名称")
            activity.name = data.name
        if activity.status == m.ACTIVITY_CLOSED:
            raise HTTPException(400, "活动已结束，不可修改")
        if "apply_deadline" in data.model_fields_set:
            activity.apply_deadline = data.apply_deadline
        if "review_deadline" in data.model_fields_set:
            activity.review_deadline = data.review_deadline
        if "pass_ratio" in data.model_fields_set and data.pass_ratio is not None:
            activity.pass_ratio = data.pass_ratio
        for field_name, value in (
            ("feishu_app_token", data.feishu_app_token),
            ("apply_table_id", data.apply_table_id),
            ("vote_table_id", data.vote_table_id),
            ("approval_code", data.approval_code),
        ):
            if value is not None and getattr(activity, field_name) != value:
                self._ensure_draft(activity, "飞书表格绑定")
                setattr(activity, field_name, value)

        if data.levels is not None:
            self._ensure_draft(activity, "职级组")
            existing = await self.level_repo.list_by_activity(activity.id)
            for old_lv in existing:
                old_lv.is_deleted = True
            for i, lv in enumerate(data.levels, start=1):
                await self.level_repo.create(
                    m.TitleReviewLevel(
                        activity_id=activity.id,
                        sequence=lv.sequence,
                        level_name=lv.level_name,
                        sort_order=i,
                    )
                )

        activity = await self.activity_repo.update(activity)
        await _audit(
            self.session,
            action="hr.title.activity_update",
            user=user,
            resource_type="title_review_activity",
            resource_id=activity.id,
            old_value=old,
            new_value={"name": activity.name, "pass_ratio": activity.pass_ratio},
        )
        return activity

    async def delete_activity(self, activity_id: UUID, user: Any = None) -> None:
        activity = await self.get_activity(activity_id)
        if activity.status != m.ACTIVITY_DRAFT:
            raise HTTPException(400, "仅配置中（draft）的活动可删除")
        count = await self.application_repo.count_by_activity(activity.id)
        if count > 0:
            raise HTTPException(400, f"活动已有 {count} 条申报记录，不可删除")
        for lv in await self.level_repo.list_by_activity(activity.id):
            lv.is_deleted = True
        for dim in await self.dimension_repo.list_by_activity(activity.id):
            dim.is_deleted = True
        activity.is_deleted = True
        await self.activity_repo.update(activity)
        await _audit(
            self.session,
            action="hr.title.activity_delete",
            user=user,
            resource_type="title_review_activity",
            resource_id=activity.id,
            old_value={"name": activity.name},
        )

    async def bind_tables(self, activity_id: UUID, user: Any = None) -> m.TitleReviewActivity:
        """校验绑定表格可访问 + 订阅事件 + 回写投票表评价项 field_id。"""
        from app.modules.hr.title_review import bitable_client as bc

        activity = await self.get_activity(activity_id)
        if not (activity.feishu_app_token and activity.apply_table_id):
            raise HTTPException(400, "请先在活动编辑中粘贴 app_token、申报表 table_id")
        # 同一申报表禁止绑定多个活动（审批实例编号全局唯一，重复绑定会在对账时冲突）
        duplicate = await self.activity_repo.find_by_feishu_binding(
            activity.feishu_app_token, activity.apply_table_id, exclude_id=activity.id
        )
        if duplicate:
            raise HTTPException(
                400,
                f"该申报表已绑定活动「{duplicate.name}」，同一申报表不能绑定多个活动；"
                "如原活动绑定信息有误，请修复原活动后重新绑定",
            )
        # 重新绑定表格后，员工信息表/审批镜像表 table_id 缓存失效（Base 表结构可能变化）
        _employee_table_cache.pop(activity.feishu_app_token, None)
        _approval_table_cache.pop(activity.feishu_app_token, None)
        try:
            apply_fields = await bc.list_fields(
                activity.feishu_app_token, activity.apply_table_id
            )
            vote_fields: list[Any] = []
            if activity.vote_table_id:
                vote_fields = await bc.list_fields(
                    activity.feishu_app_token, activity.vote_table_id
                )
        except bc.TitleReviewBitableError as exc:
            raise HTTPException(
                502, f"绑定飞书表格失败（所用飞书应用 {bc.active_app_id()}）：{exc}"
            ) from exc
        # 事件订阅为增强能力：应用无 drive:drive 权限时降级为 5 分钟对账模式，不阻断绑定
        try:
            await bc.subscribe_bitable(activity.feishu_app_token)
        except bc.TitleReviewBitableError as exc:
            logger.warning("事件订阅失败（降级为 5 分钟对账模式）: %s", exc)
        apply_names = {f.get("field_name") for f in apply_fields}
        missing_apply = {"姓名", "工号", "附件4评审结果", "同意票数", "不同意票数", "弃权票数"} - apply_names
        if missing_apply:
            raise HTTPException(400, f"申报表缺少必需列：{'、'.join(sorted(missing_apply))}")
        vote_name_to_id: dict[Any, Any] = {}
        if activity.vote_table_id:
            vote_name_to_id = {f.get("field_name"): f.get("field_id") for f in vote_fields}
            missing_vote = set(m.DEFAULT_DIMENSION_NAMES) | {"评审人编号", "评审人角色", "综合等级", "投票结果", "投票状态", "评审意见"} - set(vote_name_to_id)
            if missing_vote:
                raise HTTPException(400, f"投票表缺少必需列：{'、'.join(sorted(missing_vote))}")
        dims = await self.dimension_repo.list_by_activity(activity.id)
        for dim in dims:
            fid = vote_name_to_id.get(dim.feishu_field_name)
            if fid:
                dim.feishu_field_id = fid
        await self.session.flush()
        await _audit(
            self.session,
            action="hr.title.activity_bind_tables",
            user=user,
            resource_type="title_review_activity",
            resource_id=activity.id,
            new_value={
                "feishu_app_token": activity.feishu_app_token,
                "apply_table_id": activity.apply_table_id,
                "vote_table_id": activity.vote_table_id,
            },
        )
        return activity

    # ─── 状态流转 ───

    async def open_activity(self, activity_id: UUID, user: Any = None) -> m.TitleReviewActivity:
        activity = await self.get_activity(activity_id)
        if activity.status != m.ACTIVITY_DRAFT:
            raise HTTPException(400, f"仅配置中（draft）的活动可开启，当前：{activity.status}")
        if not (activity.feishu_app_token and activity.apply_table_id):
            raise HTTPException(400, "请先绑定飞书表格（app_token/申报表）")
        levels = await self.level_repo.list_by_activity(activity.id)
        if not levels:
            raise HTTPException(400, "请先配置职级组")
        activity.status = m.ACTIVITY_OPEN
        activity = await self.activity_repo.update(activity)
        await _audit(
            self.session,
            action="hr.title.activity_status",
            user=user,
            resource_type="title_review_activity",
            resource_id=activity.id,
            old_value={"status": m.ACTIVITY_DRAFT},
            new_value={"status": m.ACTIVITY_OPEN},
        )
        return activity

    async def start_review(self, activity_id: UUID, user: Any = None) -> m.TitleReviewActivity:
        activity = await self.get_activity(activity_id)
        if activity.status != m.ACTIVITY_OPEN:
            raise HTTPException(400, f"仅申报中（open）的活动可开启评审，当前：{activity.status}")
        activity.status = m.ACTIVITY_REVIEWING
        activity = await self.activity_repo.update(activity)
        # 批量流转：submitted → voting，并按部门自动分配评委
        assigned = 0
        applications = await self.application_repo.list_all_by_activity(activity.id)
        committees = {c.department: c for c in await self.committee_repo.list_all()}
        for application in applications:
            assigned += await self._promote_and_assign_if_reviewing(
                activity, application, committees
            )
        if assigned:
            logger.info("开始评审批量分配评委: activity=%s assigned=%s", activity.id, assigned)
        await _audit(
            self.session,
            action="hr.title.activity_status",
            user=user,
            resource_type="title_review_activity",
            resource_id=activity.id,
            old_value={"status": m.ACTIVITY_OPEN},
            new_value={"status": m.ACTIVITY_REVIEWING},
        )
        return activity

    async def close_activity(self, activity_id: UUID, user: Any = None) -> m.TitleReviewActivity:
        activity = await self.get_activity(activity_id)
        if activity.status not in (m.ACTIVITY_OPEN, m.ACTIVITY_REVIEWING):
            raise HTTPException(400, f"当前状态不可关闭：{activity.status}")
        old_status = activity.status
        activity.status = m.ACTIVITY_CLOSED
        activity = await self.activity_repo.update(activity)
        await _audit(
            self.session,
            action="hr.title.activity_status",
            user=user,
            resource_type="title_review_activity",
            resource_id=activity.id,
            old_value={"status": old_status},
            new_value={"status": m.ACTIVITY_CLOSED},
        )
        return activity

    async def get_levels(self, activity_id: UUID) -> list[m.TitleReviewLevel]:
        await self.get_activity(activity_id)
        return await self.level_repo.list_by_activity(activity_id)

    async def get_dimensions(self, activity_id: UUID) -> list[m.TitleReviewDimension]:
        await self.get_activity(activity_id)
        return await self.dimension_repo.list_by_activity(activity_id)

    # ═══ 部门评审组 ═══

    async def list_departments(self) -> list[str]:
        """员工档案中去重的实际部门（缺省回落体现部门；与申报落库、自动分配口径一致）。"""
        from sqlalchemy import select as sa_select

        result = await self.session.execute(
            sa_select(Employee.actual_department, Employee.department)
            .where(Employee.is_deleted == False)  # noqa: E712
            .distinct()
        )
        departments = {
            actual or dept for actual, dept in result.all() if actual or dept
        }
        return sorted(str(d) for d in departments if d)



    async def list_committees(self) -> list[m.TitleReviewDeptCommittee]:
        return await self.committee_repo.list_all()

    async def upsert_committee(
        self, data: TitleReviewDeptCommitteeIn, user: Any = None
    ) -> m.TitleReviewDeptCommittee:
        committee = await self.committee_repo.get_by_department(data.department)
        members = [item.model_dump(mode="json") for item in data.committee_members]
        if committee:
            committee.manager_employee_id = data.manager_employee_id
            committee.manager_name = data.manager_name
            committee.leader_employee_id = data.leader_employee_id
            committee.leader_name = data.leader_name
            committee.committee_members = members or None
            committee = await self.committee_repo.update(committee)
        else:
            committee = await self.committee_repo.create(
                m.TitleReviewDeptCommittee(
                    department=data.department,
                    manager_employee_id=data.manager_employee_id,
                    manager_name=data.manager_name,
                    leader_employee_id=data.leader_employee_id,
                    leader_name=data.leader_name,
                    committee_members=members or None,
                )
            )
        await _audit(
            self.session,
            action="hr.title.committee_upsert",
            user=user,
            resource_type="title_review_dept_committee",
            resource_id=committee.id,
            new_value={"department": committee.department},
        )
        # 评审组保存后：立即为评审期该部门投票中的申报补齐评委（幂等，只补缺）
        if members:
            added = await self._backfill_judges_for_department(committee.department)
            if added:
                logger.info(
                    "评审组 %s 补分配评委 %d 人", committee.department, added
                )
        return committee

    async def _backfill_judges_for_department(self, department: str) -> int:
        """评审组成员变更后，为该部门评审期投票中的申报补齐评委（幂等）。"""
        total = 0
        committees = {c.department: c for c in await self.committee_repo.list_all()}
        for activity in await self.activity_repo.list_active():
            if activity.status != m.ACTIVITY_REVIEWING:
                continue
            applications = await self.application_repo.list_all_by_activity(activity.id)
            for application in applications:
                if application.department != department:
                    continue
                if application.status != m.APPLICATION_VOTING:
                    continue
                total += await self._auto_assign_for_application(
                    activity, application, committees
                )
        return total

    async def delete_committee(self, committee_id: UUID, user: Any = None) -> None:
        committee = await self.committee_repo.get_by_id(committee_id)
        if not committee:
            raise HTTPException(404, "部门评审组不存在")
        committee.is_deleted = True
        await self.committee_repo.update(committee)
        await _audit(
            self.session,
            action="hr.title.committee_delete",
            user=user,
            resource_type="title_review_dept_committee",
            resource_id=committee.id,
            old_value={"department": committee.department},
        )

    # ═══ 申报同步（飞书申报表事件 → 落库） ═══

    async def sync_apply_record_added(
        self, activity_id: UUID, record_id: str, fields: dict[str, Any]
    ) -> m.TitleReviewApplication | None:
        activity = await self.get_activity(activity_id)
        existing = await self.application_repo.get_by_feishu_record(activity.id, record_id)
        if existing:
            return None
        parsed = _apply_fields_to_dict(fields)
        employee = await self._match_employee(parsed["employee_no"], parsed["name"])
        profile = None
        if employee and activity.feishu_app_token:
            # 员工信息表自动带出个人档案（学历/司龄/入职日期/职务/岗位职级等）
            profile = await self._load_employee_profile(
                activity, parsed["name"], parsed["employee_no"]
            )
        application = m.TitleReviewApplication(
            activity_id=activity.id,
            employee_id=employee.id if employee else None,
            employee_no=parsed["employee_no"],
            name=parsed["name"],
            department=employee.actual_department or employee.department if employee else None,
            sequence=parsed["sequence"],
            tech_domain=parsed["tech_domain"],
            apply_level=parsed["apply_level"],
            # 现任职级为空即留空：无职称人员如实体现，不做档案兜底
            current_level=parsed["current_level"],
            is_exception=parsed["is_exception"],
            exception_reason=parsed["exception_reason"],
            tenure_start=parsed["tenure_start"],
            tenure_end=parsed["tenure_end"],
            self_evaluations=parsed["self_evaluations"],
            work_statements=parsed["work_statements"],
            attachments=parsed["attachments"],
            profile=profile,
            profile_refreshed_at=datetime.now().astimezone() if profile else None,
            feishu_record_id=record_id,
            approval_instance_code=parsed["approval_instance_code"],
            status=m.APPLICATION_SUBMITTED if employee else m.APPLICATION_INVALID,
        )
        try:
            # SAVEPOINT 隔离本次插入：竞态失败只回滚自身，
            # 不丢弃同一会话中其他申报已 flush 的更新（对账循环/事件事务）
            async with self.session.begin_nested():
                await self.application_repo.create(application)
        except IntegrityError:
            # WS 事件与 5 分钟对账并发时唯一约束竞态：对方已建，按已存在处理
            existing = await self.application_repo.get_by_feishu_record(activity.id, record_id)
            if not existing:
                raise
            logger.warning("申报并发创建冲突，按已存在处理: record_id=%s", record_id)
            return existing
        # 活动已进入评审期 → 自动流转到投票并按部门分配评委
        await self._promote_and_assign_if_reviewing(activity, application)
        await _audit(
            self.session,
            action="hr.title.application_sync",
            resource_type="title_review_application",
            resource_id=application.id,
            extra={"record_id": record_id, "valid": employee is not None},
        )
        return application

    async def sync_apply_record_edited(
        self, activity_id: UUID, record_id: str, fields: dict[str, Any]
    ) -> m.TitleReviewApplication | None:
        activity = await self.get_activity(activity_id)
        application = await self.application_repo.get_by_feishu_record(activity.id, record_id)
        if not application:
            return await self.sync_apply_record_added(activity.id, record_id, fields)
        parsed = _apply_fields_to_dict(fields)
        changed = _apply_synced_fields(application, parsed, fields)
        if changed and application.status in TERMINAL_APPLICATION_STATUSES:
            # 终态申报放开申报信息更新（票数/判定结果仍冻结），记审计便于追溯
            await _audit(
                self.session,
                action="hr.title.application_sync",
                resource_type="title_review_application",
                resource_id=application.id,
                extra={"record_id": record_id, "terminal_override": True},
            )
        employee = await self._match_employee(parsed["employee_no"], parsed["name"])
        application.employee_id = employee.id if employee else application.employee_id
        if employee and not application.department:
            # INVALID 期间部门为空，员工匹配成功后补上，否则自动分配匹配不到评审组
            application.department = employee.actual_department or employee.department
        if application.status == m.APPLICATION_INVALID and employee:
            application.status = m.APPLICATION_SUBMITTED
            await self.application_repo.update(application)
            await self._promote_and_assign_if_reviewing(activity, application)
            return application
        await self.application_repo.update(application)
        return application

    async def sync_apply_record_deleted(self, activity_id: UUID, record_id: str) -> None:
        activity = await self.get_activity(activity_id)
        application = await self.application_repo.get_by_feishu_record(activity.id, record_id)
        if not application:
            return
        if application.status in TERMINAL_APPLICATION_STATUSES:
            logger.warning("已判定申报不可撤回: record_id=%s", record_id)
            return
        application.is_deleted = True
        await self.application_repo.update(application)

    async def _match_employee(
        self, employee_no: str, name: str
    ) -> Employee | None:
        """姓名+工号匹配在职员工（排除离职）。"""
        from app.modules.hr.repository import EmployeeRepository

        if not employee_no or not name:
            return None
        employee = await EmployeeRepository(self.session).get_by_employee_number(employee_no)
        if not employee or employee.status == "离职":
            return None
        if employee.name != name:
            return None
        return employee

    async def _load_employee_profile(
        self, activity: m.TitleReviewActivity, name: str, employee_no: str
    ) -> dict[str, Any] | None:
        """从 Base 员工信息表按姓名+工号取个人档案（容错：失败不阻断申报落库）。

        审批表单/申报表未体现的信息（学历/司龄/入职日期/职务/岗位职级/
        目前职级/近5年绩效/2026最高可申报等）由此自动补充。
        """
        from app.modules.hr.title_review import bitable_client as bc

        app_token = activity.feishu_app_token
        if not app_token:
            return None
        try:
            emp_table_id = _employee_table_cache.get(app_token)
            if emp_table_id is None:
                tables = await bc.list_tables(app_token)
                emp_table = next(
                    (t for t in tables if t.get("name") == "员工信息表"), None
                )
                if not emp_table:
                    logger.warning("Base 内未找到员工信息表: %s", app_token)
                    return None
                emp_table_id = str(emp_table.get("table_id") or "")
                _employee_table_cache[app_token] = emp_table_id
            records = await bc.list_all_records(
                app_token,
                emp_table_id,
                filter_expr=f'CurrentValue.[姓名]="{name}"',
            )
            for item in records:
                fields = item.get("fields") or {}
                if _as_str(fields.get("工号")) != employee_no:
                    continue
                profile = {
                    key: _as_str(fields.get(key))
                    for key in EMPLOYEE_PROFILE_FIELDS
                    if _as_str(fields.get(key)) is not None
                }
                return profile or None
            logger.warning(
                "员工信息表未匹配到姓名+工号: %s %s", name, employee_no
            )
        except Exception as exc:  # noqa: BLE001
            logger.warning("员工信息表档案补充失败: %s", exc)
        return None

    async def default_committee_members(
        self, application_id: UUID
    ) -> list[dict[str, Any]]:
        """申报人部门评审组的默认评委（评定小组成员）。"""
        application = await self.application_repo.get_by_id(application_id)
        if not application or not application.department:
            return []
        committee = await self.committee_repo.get_by_department(application.department)
        if not committee or not committee.committee_members:
            return []
        return list(committee.committee_members)

    async def _next_judge_code(self, application_id: UUID) -> str:
        """下一个匿名评审人编号：包含软删除行取最大值+1，编号永不复用。"""
        from sqlalchemy import select as sa_select

        rows = (await self.session.execute(
            sa_select(m.TitleReviewJudge.judge_code).where(
                m.TitleReviewJudge.application_id == application_id
            )
        )).scalars().all()
        max_index = 0
        for code in rows:
            if code and code.startswith("P"):
                try:
                    max_index = max(max_index, int(code[1:]))
                except ValueError:
                    continue
        return f"P{max_index + 1}"

    async def _auto_assign_for_application(
        self,
        activity: m.TitleReviewActivity,
        application: m.TitleReviewApplication,
        committees: dict[str, m.TitleReviewDeptCommittee] | None = None,
    ) -> int:
        """按申报人部门自动分配评委（部门评审组评定小组）；无配置返回 0。

        committees: 批量场景预加载的 {部门: 评审组} 映射（避免逐申报查询）。
        """
        if not application.department:
            return 0
        if committees is not None:
            committee = committees.get(application.department)
        else:
            committee = await self.committee_repo.get_by_department(application.department)
        if not committee or not committee.committee_members:
            logger.info("部门 %s 未配置评审组，跳过自动分配", application.department)
            return 0
        from sqlalchemy import select as sa_select

        existing = await self.judge_repo.list_by_application(application.id)
        existing_ids = {j.judge_employee_id for j in existing}
        # 批量加载评审组成员员工档案（一条 IN 查询替代逐成员查询）
        member_ids: list[UUID] = []
        for member in committee.committee_members:
            member_id = member.get("employee_id")
            if not member_id:
                continue
            try:
                member_ids.append(UUID(str(member_id)))
            except (TypeError, ValueError):
                continue
        emps: dict[UUID, Employee] = {}
        if member_ids:
            emp_rows = (
                await self.session.execute(
                    sa_select(Employee).where(
                        Employee.id.in_(member_ids),
                        Employee.is_deleted == False,  # noqa: E712
                    )
                )
            ).scalars().all()
            emps = {e.id: e for e in emp_rows}
        added = 0
        new_judges: list[m.TitleReviewJudge] = []
        for emp_id in member_ids:
            if emp_id in existing_ids:
                continue
            emp = emps.get(emp_id)
            if not emp or emp.status == "离职":
                continue
            judge = await self.judge_repo.create(
                m.TitleReviewJudge(
                    activity_id=activity.id,
                    application_id=application.id,
                    judge_employee_id=emp.id,
                    judge_name=emp.name,
                    judge_employee_no=emp.employee_number,
                    judge_code=await self._next_judge_code(application.id),
                    judge_role="评委",
                )
            )
            new_judges.append(judge)
            added += 1
        if new_judges:
            await _audit(
                self.session,
                action="hr.title.judge_auto_assign",
                resource_type="title_review_application",
                resource_id=application.id,
                new_value={
                    "department": application.department,
                    "judges": [j.judge_code for j in new_judges],
                },
            )
            await self._notify_new_judges(activity, application, new_judges)
        return added

    async def _notify_new_judges(
        self,
        activity: m.TitleReviewActivity,
        application: m.TitleReviewApplication,
        new_judges: list[m.TitleReviewJudge],
    ) -> None:
        """提醒新评委登录内网投票（容错：提醒失败不阻断业务）。"""
        if not new_judges:
            return
        from app.modules.hr.title_review.notify import send_judge_reminder

        for j in new_judges:
            try:
                await send_judge_reminder(
                    judge_name=j.judge_name,
                    activity_name=activity.name,
                    applicant_name=application.name,
                    judge_code=j.judge_code,
                )
            except Exception:  # noqa: BLE001
                logger.warning("评委提醒发送失败: %s", j.judge_name)

    async def _promote_and_assign_if_reviewing(
        self,
        activity: m.TitleReviewActivity,
        application: m.TitleReviewApplication,
        committees: dict[str, m.TitleReviewDeptCommittee] | None = None,
    ) -> int:
        """评审期新申报自动流转：submitted → voting 并分配评委（幂等），返回分配数。"""
        if application.status != m.APPLICATION_SUBMITTED:
            return 0
        if activity.status != m.ACTIVITY_REVIEWING:
            return 0
        application.status = m.APPLICATION_VOTING
        await self.application_repo.update(application)
        return await self._auto_assign_for_application(
            activity, application, committees
        )

    async def assign_judges(
        self, application_id: UUID, data: TitleReviewJudgeAssignIn, user: Any = None
    ) -> list[m.TitleReviewJudge]:
        """指定/调整评委（幂等 upsert，匿名编号）→ 提醒评委登录内网投票。"""
        application = await self.application_repo.get_by_id(application_id)
        if not application:
            raise HTTPException(404, "申报记录不存在")
        activity = await self.get_activity(application.activity_id)
        if activity.status not in (m.ACTIVITY_OPEN, m.ACTIVITY_REVIEWING):
            raise HTTPException(400, f"活动状态不可指定评委：{activity.status}")
        if application.status == m.APPLICATION_INVALID:
            raise HTTPException(400, "员工档案未匹配，无法指定评委（请补全员工档案后重新同步）")
        if application.status not in (m.APPLICATION_VOTING,):
            raise HTTPException(400, f"当前状态不可指定评委：{application.status}")

        from app.modules.hr.repository import EmployeeRepository

        emp_repo = EmployeeRepository(self.session)
        employees = []
        for item in data.judges:
            emp = await emp_repo.get_by_id(item.employee_id)
            if not emp or emp.status == "离职":
                raise HTTPException(400, f"评委不存在或已离职: {item.employee_id}")
            employees.append((emp, item.role))

        existing_judges = await self.judge_repo.list_by_application(application.id)
        existing_map = {j.judge_employee_id: j for j in existing_judges}
        target_ids = {item.employee_id for item in data.judges}

        for j in existing_judges:
            if j.judge_employee_id not in target_ids:
                if j.vote_result is not None:
                    raise HTTPException(400, f"评委 {j.judge_name} 已投票，不可撤换")
                j.is_deleted = True
                await self.judge_repo.update(j)

        # 编号统一分配器（含软删除行，永不复用）
        new_judges: list[m.TitleReviewJudge] = []
        for emp, role in employees:
            if emp.id in existing_map and not existing_map[emp.id].is_deleted:
                continue
            judge = await self.judge_repo.create(
                m.TitleReviewJudge(
                    activity_id=activity.id,
                    application_id=application.id,
                    judge_employee_id=emp.id,
                    judge_name=emp.name,
                    judge_employee_no=emp.employee_number,
                    judge_code=await self._next_judge_code(application.id),
                    judge_role=role,
                )
            )
            new_judges.append(judge)

        await self.session.flush()

        all_judges = await self.judge_repo.list_by_application(application.id)
        await _audit(
            self.session,
            action="hr.title.judge_assign",
            user=user,
            resource_type="title_review_application",
            resource_id=application.id,
            new_value={
                "judges": [{"code": j.judge_code, "name": j.judge_name, "role": j.judge_role} for j in all_judges],
            },
        )
        # 提醒新评委登录内网系统投票
        await self._notify_new_judges(activity, application, new_judges)
        return all_judges

    async def sync_vote_record(
        self, activity_id: UUID, record_id: str, fields: dict[str, Any]
    ) -> None:
        """投票表记录变更：投票结果/综合等级/7 维评价落库，全投完自动判定。"""
        activity = await self.get_activity(activity_id)
        judge = None
        if record_id:
            judge = await self.judge_repo.get_by_feishu_record(record_id)
        if not judge:
            judge_code = _as_str(fields.get("评审人编号"))
            if judge_code:
                from sqlalchemy import select as sa_select

                # 评审人编号仅在单个申报内唯一（每个申报都有 P1），跨申报会出现重号：
                # 命中多个时不做猜测，跳过本次同步（避免投错申报）
                result = await self.session.execute(
                    sa_select(m.TitleReviewJudge).where(
                        m.TitleReviewJudge.activity_id == activity_id,
                        m.TitleReviewJudge.judge_code == judge_code,
                        m.TitleReviewJudge.is_deleted == False,  # noqa: E712
                    )
                )
                matches = list(result.scalars().all())
                if len(matches) == 1:
                    judge = matches[0]
                elif len(matches) > 1:
                    logger.warning(
                        "评审人编号 %s 在活动内命中 %d 个评委，无法唯一定位，跳过: record_id=%s",
                        judge_code, len(matches), record_id,
                    )
                    return
        if not judge:
            logger.warning("无法匹配投票表行: record_id=%s", record_id)
            return
        application = await self.application_repo.get_by_id(judge.application_id)
        if not application:
            return

        vote_result = _as_str(fields.get("投票结果"))
        judge.vote_result = vote_result
        judge.comprehensive_grade = _as_str(fields.get("综合等级"))
        judge.review_comment = _as_str(fields.get("评审意见"))
        if vote_result:
            judge.voted_at = _as_ts(fields.get("投票时间")) or datetime.now().astimezone()
        elif _as_str(fields.get("投票状态")) == "已作废":
            judge.vote_result = None
            judge.voted_at = None
        judge = await self.judge_repo.update(judge)

        dims = await self.dimension_repo.list_by_activity(activity.id)
        grades = {
            dim.name: _as_str(fields.get(dim.feishu_field_name)) for dim in dims
        }
        await self._upsert_dimension_scores(
            activity.id, application.id, judge.id, dims, grades, judge.voted_at
        )
        await _audit(
            self.session,
            action="hr.title.vote_sync",
            resource_type="title_review_judge",
            resource_id=judge.id,
            extra={"judge_code": judge.judge_code, "vote_result": vote_result},
        )
        # 回写投票表“投票状态”列
        if activity.feishu_app_token and activity.vote_table_id:
            from app.modules.hr.title_review import bitable_client as bc

            try:
                await bc.update_record(
                    activity.feishu_app_token,
                    activity.vote_table_id,
                    record_id,
                    {"投票状态": "已投票" if vote_result else "未投票"},
                )
            except bc.TitleReviewBitableError as exc:
                logger.warning("回写投票状态失败: record_id=%s error=%s", record_id, exc)
        # 全部评委已投票 → 自动判定
        await self._maybe_finalize(application.id)

    async def _upsert_dimension_scores(
        self,
        activity_id: UUID,
        application_id: UUID,
        judge_id: UUID,
        dims: list[m.TitleReviewDimension],
        grades: dict[str, str | None],
        voted_at: datetime | None,
    ) -> None:
        """维度分数幂等 upsert（单条查询 + 内存匹配，替代逐维度查询）。"""
        existing_scores = await self.score_repo.list_by_judge(judge_id)
        existing_by_dim: dict[UUID, m.TitleReviewScore] = {
            s.dimension_id: s for s in existing_scores
        }
        for dim in dims:
            grade = grades.get(dim.name)
            if grade is None:
                continue
            existing = existing_by_dim.get(dim.id)
            if existing:
                existing.grade = grade
                existing.voted_at = voted_at
                await self.score_repo.update(existing)
            else:
                await self.score_repo.create(
                    m.TitleReviewScore(
                        activity_id=activity_id,
                        application_id=application_id,
                        judge_id=judge_id,
                        dimension_id=dim.id,
                        dimension_name=dim.name,
                        grade=grade,
                        voted_at=voted_at,
                    )
                )

    async def _maybe_finalize(self, application_id: UUID) -> None:
        """全部评委已投票 → 自动票数判定（幂等）。"""
        application = await self.application_repo.get_by_id(application_id)
        if not application or application.status not in (
            m.APPLICATION_VOTING, m.APPLICATION_PASSED, m.APPLICATION_FAILED
        ):
            return
        judges = await self.judge_repo.list_by_application(application_id)
        if judges and all(j.vote_result is not None for j in judges):
            await self.finalize_by_votes(application_id)

    async def finalize_by_votes(
        self, application_id: UUID, user: Any = None, force: bool = False
    ) -> m.TitleReviewApplication:
        """票数判定：统计票数 → 回写申报表 → 终审卡或结果通知。"""
        application = await self.application_repo.get_by_id(application_id)
        if not application:
            raise HTTPException(404, "申报记录不存在")
        if application.status not in (m.APPLICATION_VOTING, m.APPLICATION_PASSED, m.APPLICATION_FAILED):
            raise HTTPException(400, f"当前状态不可判定：{application.status}")
        judges = await self.judge_repo.list_by_application(application.id)
        if not judges:
            raise HTTPException(400, "尚未指定评委，无法判定")
        if not force and not all(j.vote_result is not None for j in judges):
            raise HTTPException(400, "仍有评委未投票，如需提前判定请确认")

        agree = sum(1 for j in judges if j.vote_result == m.VOTE_AGREE)
        oppose = sum(1 for j in judges if j.vote_result == m.VOTE_OPPOSE)
        abstain = sum(1 for j in judges if j.vote_result == m.VOTE_ABSTAIN)
        activity = await self.get_activity(application.activity_id)
        passed = decide_by_votes(agree, oppose, activity.pass_ratio)

        old = {
            "status": application.status,
            "agree_votes": application.agree_votes,
            "oppose_votes": application.oppose_votes,
            "abstain_votes": application.abstain_votes,
            "final_result": application.final_result,
        }
        application.agree_votes = agree
        application.oppose_votes = oppose
        application.abstain_votes = abstain
        application.final_result = m.APPLICATION_PASSED if passed else m.APPLICATION_FAILED
        application.status = m.APPLICATION_PASSED if passed else m.APPLICATION_FAILED
        application = await self.application_repo.update(application)

        # 回写申报表：票数 3 列 + 附件4评审结果
        await self._writeback_votes(application, passed)
        await _audit(
            self.session,
            action="hr.title.vote_finalize",
            user=user,
            resource_type="title_review_application",
            resource_id=application.id,
            old_value=old,
            new_value={
                "status": application.status,
                "agree_votes": agree,
                "oppose_votes": oppose,
                "abstain_votes": abstain,
            },
        )

        # 小组评审结果并非最终结果（最终名单须经总经理确认），不直接通知申报人；
        # 结果由 HR 在公示环节线下发布。
        return application

    async def _writeback_votes(
        self, application: m.TitleReviewApplication, passed: bool
    ) -> None:
        """回写申报表：同意/不同意/弃权票数 + 附件4评审结果。"""
        from app.modules.hr.title_review import bitable_client as bc

        activity = await self.get_activity(application.activity_id)
        if not (activity.feishu_app_token and activity.apply_table_id and application.feishu_record_id):
            return
        try:
            await bc.update_record(
                activity.feishu_app_token,
                activity.apply_table_id,
                application.feishu_record_id,
                {
                    "同意票数": application.agree_votes,
                    "不同意票数": application.oppose_votes,
                    "弃权票数": application.abstain_votes,
                    "附件4评审结果": "通过(三分之二以上同意)" if passed else "不通过",
                },
            )
        except bc.TitleReviewBitableError as exc:
            logger.warning("回写申报表票数失败: record_id=%s error=%s", application.feishu_record_id, exc)

    # ═══ 查询 ═══

    async def list_applications(
        self,
        activity_id: UUID,
        *,
        status: str | None,
        keyword: str | None,
        page: int,
        page_size: int,
    ) -> tuple[list[m.TitleReviewApplication], int]:
        await self.get_activity(activity_id)
        return await self.application_repo.list_by_activity(
            activity_id,
            status=status,
            keyword=keyword,
            page=page,
            page_size=page_size,
            departments=self._scope_departments,
            employee_no=self._scope_employee_no,
        )

    async def get_application_detail(
        self, application_id: UUID
    ) -> tuple[m.TitleReviewApplication, list[m.TitleReviewJudge]]:
        application = await self.application_repo.get_by_id(application_id)
        if not application:
            raise HTTPException(404, "申报记录不存在")
        if not self._scope_applications([application]):
            raise HTTPException(403, "数据范围限制：无权访问该申报")
        judges = await self.judge_repo.list_by_application(application.id)
        return application, judges

    async def get_results(self, activity_id: UUID) -> list[dict[str, Any]]:
        """活动评审结果（票数、各评委投票明细）——hr:title:scores:read 数据。"""
        await self.get_activity(activity_id)
        applications = self._scope_applications(
            await self.application_repo.list_all_by_activity(activity_id)
        )
        # 批量加载评委与维度分数（两条查询替代 2×N 条）
        judges_by_app: dict[UUID, list[m.TitleReviewJudge]] = {}
        for j in await self.judge_repo.list_by_activity(activity_id):
            judges_by_app.setdefault(j.application_id, []).append(j)
        scores_by_judge: dict[UUID, list[m.TitleReviewScore]] = {}
        for s in await self.score_repo.list_by_activity(activity_id):
            scores_by_judge.setdefault(s.judge_id, []).append(s)

        results: list[dict[str, Any]] = []
        for application in applications:
            judges = judges_by_app.get(application.id, [])
            judge_out = []
            for j in judges:
                item = TitleReviewJudgeOut.model_validate(j).model_dump(mode="json")
                item["scores"] = [
                    TitleReviewScoreOut.model_validate(s).model_dump(mode="json")
                    for s in scores_by_judge.get(j.id, [])
                ]
                judge_out.append(item)
            voted = application.agree_votes + application.oppose_votes
            results.append(
                {
                    "application": TitleReviewApplicationOut.model_validate(
                        application
                    ).model_dump(mode="json"),
                    "judges": judge_out,
                    "vote_ratio": round(application.agree_votes / voted, 4) if voted else None,
                }
            )
        return results

    # ═══ 评委内网投票（v3：投票在内网系统完成） ═══

    async def _resolve_judge_for_user(
        self, judge_id: UUID, employee_no: str | None
    ) -> tuple[m.TitleReviewJudge, m.TitleReviewApplication]:
        """校验评委任务归属：仅本人（工号匹配）可操作自己的投票任务。"""
        from app.modules.hr.repository import EmployeeRepository

        judge = await self.judge_repo.get_by_id(judge_id)
        if not judge:
            raise HTTPException(404, "投票任务不存在")
        if not employee_no:
            raise HTTPException(403, "无法识别您的员工身份，请联系管理员")
        employee = await EmployeeRepository(self.session).get_by_employee_number(employee_no)
        if not employee or employee.id != judge.judge_employee_id:
            raise HTTPException(403, "该投票任务不属于您")
        application = await self.application_repo.get_by_id(judge.application_id)
        if not application:
            raise HTTPException(404, "申报记录不存在")
        return judge, application

    async def list_my_judge_tasks(
        self, employee_no: str | None
    ) -> list[dict[str, Any]]:
        """评委视角：我的投票任务列表（不含其他评委信息）。"""
        if not employee_no:
            return []
        from app.modules.hr.repository import EmployeeRepository

        employee = await EmployeeRepository(self.session).get_by_employee_number(employee_no)
        if not employee:
            return []
        from sqlalchemy import select as sa_select

        result = await self.session.execute(
            sa_select(m.TitleReviewJudge)
            .where(
                m.TitleReviewJudge.judge_employee_id == employee.id,
                m.TitleReviewJudge.is_deleted == False,  # noqa: E712
            )
            .order_by(m.TitleReviewJudge.vote_result.is_(None).desc(), m.TitleReviewJudge.created_at)
        )
        judges = list(result.scalars().all())
        if not judges:
            return []
        # 批量加载申报/活动/分数（避免逐条 N+1 查询）
        app_ids = {j.application_id for j in judges}
        app_rows = (
            await self.session.execute(
                sa_select(m.TitleReviewApplication).where(
                    m.TitleReviewApplication.id.in_(app_ids),
                    m.TitleReviewApplication.is_deleted == False,  # noqa: E712
                )
            )
        ).scalars().all()
        applications = {a.id: a for a in app_rows}
        act_ids = {a.activity_id for a in app_rows}
        act_rows = (
            await self.session.execute(
                sa_select(m.TitleReviewActivity).where(m.TitleReviewActivity.id.in_(act_ids))
            )
        ).scalars().all()
        activities = {a.id: a for a in act_rows}
        judge_ids = {j.id for j in judges}
        score_rows = (
            await self.session.execute(
                sa_select(m.TitleReviewScore).where(
                    m.TitleReviewScore.judge_id.in_(judge_ids),
                    m.TitleReviewScore.is_deleted == False,  # noqa: E712
                )
            )
        ).scalars().all()
        scores_by_judge: dict[UUID, list[m.TitleReviewScore]] = {}
        for s in score_rows:
            scores_by_judge.setdefault(s.judge_id, []).append(s)
        dim_rows = (
            await self.session.execute(
                sa_select(m.TitleReviewDimension).where(
                    m.TitleReviewDimension.activity_id.in_(act_ids),
                    m.TitleReviewDimension.is_deleted == False,  # noqa: E712
                )
            )
        ).scalars().all()
        dims_by_activity: dict[UUID, list[m.TitleReviewDimension]] = {}
        for d in sorted(dim_rows, key=lambda x: x.sort_order):
            dims_by_activity.setdefault(d.activity_id, []).append(d)

        tasks: list[dict[str, Any]] = []
        for judge in judges:
            application = applications.get(judge.application_id)
            if not application:
                continue
            activity = activities.get(application.activity_id)
            if not activity or activity.status == m.ACTIVITY_CLOSED:
                continue
            scores = scores_by_judge.get(judge.id, [])
            # 分级合格标准（附表5 表11）：工程师/技师及以上=高档，其余=低档
            tier = (
                "high"
                if application.apply_level in m.HIGH_TIER_LEVELS
                else "low"
            )
            standards = m.DIMENSION_GRADE_STANDARDS[tier]
            dims = dims_by_activity.get(activity.id, [])
            tasks.append(
                {
                    "judge_id": judge.id,
                    "judge_code": judge.judge_code,
                    "status": "voted" if judge.vote_result else "pending",
                    "vote_result": judge.vote_result,
                    "comprehensive_grade": judge.comprehensive_grade,
                    "review_comment": judge.review_comment,
                    "voted_at": judge.voted_at,
                    "dimension_grades": {s.dimension_name: s.grade for s in scores},
                    "grade_tier": tier,
                    "dimensions": [
                        {"name": d.name, "standard": standards.get(d.name, "")}
                        for d in dims
                    ],
                    "activity_name": activity.name,
                    "application": TitleReviewApplicationOut.model_validate(
                        application
                    ).model_dump(mode="json"),
                }
            )
        return tasks

    async def submit_vote(
        self,
        judge_id: UUID,
        employee_no: str | None,
        data: Any,
    ) -> m.TitleReviewJudge:
        """评委提交投票：7 维评价 → 自动计算综合等级与投票结果；全投完自动判定。

        投票结果按《员工职级评定管理办法》附表5自动计算：
        综合等级 合格 → 同意；不合格 → 不同意（维度未填齐不可提交）。
        """
        judge, application = await self._resolve_judge_for_user(judge_id, employee_no)
        if application.status not in (m.APPLICATION_VOTING, m.APPLICATION_PASSED, m.APPLICATION_FAILED):
            raise HTTPException(400, f"当前状态不可投票：{application.status}")
        activity = await self.activity_repo.get_by_id(application.activity_id)
        if activity and activity.review_deadline:
            deadline = activity.review_deadline
            if deadline.tzinfo is None:
                # 客户端传入无时区的截止时间：按本地时区解释，避免 naive/aware 比较 TypeError
                deadline = deadline.astimezone()
            if datetime.now().astimezone() > deadline:
                raise HTTPException(400, "评审已截止，不可再投票")

        dims = await self.dimension_repo.list_by_activity(application.activity_id)
        grades = [(data.dimension_grades or {}).get(d.name) for d in dims]
        if len(dims) < len(m.DEFAULT_DIMENSION_NAMES) or any(
            g not in ("合格", "不合格") for g in grades
        ):
            raise HTTPException(
                400,
                f"请完成全部 {len(m.DEFAULT_DIMENSION_NAMES)} 项维度评价，投票结果将自动计算",
            )
        auto_grade = compute_comprehensive_grade(grades)
        judge.vote_result = (
            m.VOTE_OPPOSE if auto_grade == "不合格" else m.VOTE_AGREE
        )
        judge.comprehensive_grade = auto_grade
        judge.review_comment = data.review_comment
        judge.voted_at = datetime.now().astimezone()
        judge = await self.judge_repo.update(judge)

        await self._upsert_dimension_scores(
            application.activity_id,
            application.id,
            judge.id,
            dims,
            dict(data.dimension_grades or {}),
            judge.voted_at,
        )
        await _audit(
            self.session,
            action="hr.title.judge_vote",
            resource_type="title_review_judge",
            resource_id=judge.id,
            extra={"judge_code": judge.judge_code, "vote_result": judge.vote_result},
        )
        # 全部评委已投票 → 自动判定
        await self._maybe_finalize(application.id)
        return judge

    # ═══ 审批先行同步（飞书审批通过 → 写申报表） ═══

    async def sync_approval_instances(self, activity_id: UUID) -> dict[str, Any]:
        """同步已通过的审批实例到申报表（事件自然落库）。

        数据源优先 Base 审批镜像表「福兴医药职称审批」（含申请状态与实例编号，
        不受审批 API 限流影响）；镜像表不存在时回退飞书审批 API。
        """
        from app.modules.hr.title_review import bitable_client as bc

        activity = await self.get_activity(activity_id)
        stats: dict[str, Any] = {
            "approval_instances_fetched": 0,
            "approval_synced": 0,
            "approval_updated": 0,
            "approval_skipped": 0,
        }
        if not activity.approval_code:
            return stats
        if not (activity.feishu_app_token and activity.apply_table_id):
            return stats

        # 申报表现状：审批实例编号（支持逗号累积）→ 行；姓名+工号 → 行（同员工覆盖更新）
        existing_map: dict[str, str] = {}
        employee_row_map: dict[tuple[str, str], dict[str, Any]] = {}
        try:
            records = await bc.list_all_records(
                activity.feishu_app_token, activity.apply_table_id
            )
        except bc.TitleReviewBitableError as exc:
            logger.warning("拉取申报表失败: %s", exc)
            return stats
        for item in records:
            fields = item.get("fields") or {}
            raw_codes = _as_str(fields.get("审批实例编号")) or ""
            for code_part in raw_codes.replace("，", ",").split(","):
                code_part = code_part.strip()
                if code_part:
                    existing_map[code_part] = str(item.get("record_id") or "")
            name = _as_str(fields.get("姓名"))
            emp_no = _as_str(fields.get("工号"))
            if name and emp_no:
                employee_row_map.setdefault((name, emp_no), item)

        # 申报表可写列（跳过 lookup/formula/系统只读列，单选列做选项校验）
        try:
            table_fields = await bc.list_fields(
                activity.feishu_app_token, activity.apply_table_id
            )
        except bc.TitleReviewBitableError as exc:
            logger.warning("拉取申报表字段失败: %s", exc)
            return stats

        rows_to_write: list[dict[str, Any]] = []
        # 数据源：优先审批镜像表，不存在/拉取失败时回退审批 API
        fetched = await self._pending_approval_rows_from_table(activity, existing_map)
        if fetched is None:
            fetched = await self._pending_approval_rows_from_api(activity, existing_map)
        if fetched is None:
            return stats
        pending, skipped = fetched
        stats["approval_skipped"] += skipped
        stats["approval_instances_fetched"] = len(pending) + skipped

        for code, fields in pending:
            if not _as_str(fields.get("姓名")) or not _as_str(fields.get("工号")):
                logger.warning("审批实例缺少姓名/工号，跳过: %s", code)
                continue
            # 表单字段映射到申报表列名（未列出的按原名写入；多字段同列时换行合并）
            row: dict[str, Any] = {}
            for form_name, value in fields.items():
                col = APPROVAL_FORM_FIELD_MAP.get(form_name, form_name)
                if col == "任职以来工作简历":
                    # fieldList 控件：值为 [[{name, value}...]]（嵌套列表）或 JSON 字符串 → 「name：value」逐条换行
                    items = value
                    if isinstance(items, str):
                        try:
                            items = json.loads(items)
                        except (TypeError, ValueError):
                            items = None
                    if isinstance(items, list) and items and isinstance(items[0], list):
                        items = items[0]  # 展开一层嵌套
                    if isinstance(items, list):
                        lines: list[str] = []
                        for item in items:
                            if isinstance(item, dict):
                                name = str(item.get("name") or "")
                                raw = item.get("value")
                                # 时间区间控件：去掉技术名 DateInterval，保留起止日期
                                if (
                                    item.get("type") == "dateInterval"
                                    or (isinstance(raw, dict) and ("start" in raw or "end" in raw))
                                ) and isinstance(raw, dict):
                                    val = f'{str(raw.get("start") or "")[:10]} ~ {str(raw.get("end") or "")[:10]}'.strip(" ~")
                                    lines.append(f"任职时间：{val}" if val else "")
                                    continue
                                val = str(raw or "")
                                lines.append(f"{name}：{val}" if name and val else (name or val))
                            else:
                                lines.append(str(item))
                        value = "\n".join(line for line in lines if line)
                if col in IMAGE_EVIDENCE_FIELDS and isinstance(value, list):
                    # 图片控件：下载转存本地（飞书签名链接约 1 天过期），本地路径永久可看
                    text = _image_list_to_text(value) or ""
                    value = await _image_text_to_local(text)
                if col in row:
                    existing = _as_str(row[col]) or ""
                    incoming = _as_str(value) or ""
                    if existing and incoming:
                        row[col] = f"{existing}\n{incoming}"
                    elif incoming:
                        row[col] = value
                    continue
                row[col] = value
            row["审批实例编号"] = code
            # 审批附件：有 file_token 才写申报表附件列，否则元数据并入申报说明
            for attach_key in APPLY_ATTACHMENT_FIELDS:
                raw = row.get(attach_key)
                if not isinstance(raw, list):
                    continue
                valid = [a for a in raw if isinstance(a, dict) and a.get("file_token")]
                if valid:
                    row[attach_key] = [{"file_token": a["file_token"]} for a in valid]
                else:
                    row.pop(attach_key, None)
                    meta = "、".join(
                        str(a.get("title") or a.get("name") or a.get("url") or "")
                        for a in raw
                        if isinstance(a, dict)
                    )
                    if meta:
                        note = _as_str(row.get("申报说明")) or ""
                        row["申报说明"] = f"{note}\n【{attach_key}】{meta}".strip()
            # 只写申报表存在的可写列；lookup 列由「姓名」自动从员工信息表带出
            row = _filter_writable_row(row, table_fields)
            if not row.get("姓名"):
                logger.warning("审批实例过滤后无姓名列，跳过: %s", code)
                continue
            # 同员工已有申报行 → 覆盖更新该行，审批实例编号累积（防旧实例重新写回）
            emp_key = (
                _as_str(fields.get("姓名")) or "",
                _as_str(fields.get("工号")) or "",
            )
            target = employee_row_map.get(emp_key)
            if target:
                record_id = str(target.get("record_id") or "")
                old_codes = (
                    _as_str((target.get("fields") or {}).get("审批实例编号")) or ""
                )
                row["审批实例编号"] = f"{old_codes},{code}" if old_codes else code
                try:
                    await bc.update_record(
                        activity.feishu_app_token,
                        activity.apply_table_id,
                        record_id,
                        row,
                    )
                    stats["approval_updated"] += 1
                    logger.info(
                        "同员工申报覆盖更新: %s record=%s", emp_key[0], record_id
                    )
                except bc.TitleReviewBitableError as exc:
                    logger.warning("覆盖更新申报行失败: %s", exc)
                continue
            rows_to_write.append(row)

        if rows_to_write:
            try:
                await bc.batch_create_records(
                    activity.feishu_app_token, activity.apply_table_id, rows_to_write
                )
                stats["approval_synced"] = len(rows_to_write)
            except bc.TitleReviewBitableError as exc:
                logger.warning("写入申报表失败: %s", exc)
                stats["approval_synced"] = 0
        # 图片证据转存兜底：飞书签名链接约 1 天过期，对账时把残留远程链接下载转存本地
        try:
            stats["image_urls_refreshed"] = await self._refresh_image_urls(
                activity, records
            )
        except Exception as exc:  # noqa: BLE001
            logger.warning("图片转存兜底失败: %s", exc)
        await _audit(
            self.session,
            action="hr.title.approval_sync",
            resource_type="title_review_activity",
            resource_id=activity.id,
            extra=stats,
        )
        return stats

    async def _pending_approval_rows_from_table(
        self, activity: m.TitleReviewActivity, existing_map: dict[str, str]
    ) -> tuple[list[tuple[str, dict[str, Any]]], int] | None:
        """从 Base 审批镜像表读取已通过且未写入的实例。

        返回 [(实例编号, 表单字段), ...] 与重复跳过数；镜像表不存在返回 None
        （调用方回退审批 API）。
        """
        from app.modules.hr.title_review import bitable_client as bc

        app_token = activity.feishu_app_token
        if not app_token:
            return None
        table_id = _approval_table_cache.get(app_token)
        if table_id is None:
            tables = await bc.list_tables(app_token)
            approval_table = next(
                (t for t in tables if t.get("name") == APPROVAL_TABLE_NAME), None
            )
            if not approval_table:
                logger.info(
                    "Base 内未找到审批镜像表「%s」，回退审批 API: %s",
                    APPROVAL_TABLE_NAME, app_token,
                )
                return None
            table_id = str(approval_table.get("table_id") or "")
            _approval_table_cache[app_token] = table_id
        rows = await bc.list_all_records(app_token, table_id)
        pending: list[tuple[str, dict[str, Any]]] = []
        skipped = 0
        for item in rows:
            fields = item.get("fields") or {}
            if _as_str(fields.get("申请状态")) != "已通过":
                continue
            source_id = _as_str(fields.get("SourceID")) or ""
            code = _extract_approval_code_from_source_id(source_id)
            if not code:
                logger.warning(
                    "审批镜像行 SourceID 无法解析，跳过: record_id=%s",
                    item.get("record_id"),
                )
                continue
            if code in existing_map:
                skipped += 1
                continue
            _merge_resume_columns(fields)
            pending.append((code, fields))
        return pending, skipped

    async def _pending_approval_rows_from_api(
        self, activity: m.TitleReviewActivity, existing_map: dict[str, str]
    ) -> tuple[list[tuple[str, dict[str, Any]]], int] | None:
        """审批 API 兜底：拉取近 30 天已通过且未同步的实例。"""
        from app.modules.hr.title_review import approval_client as ac

        if not activity.approval_code:
            return None
        now = datetime.now().astimezone()
        end_ms = int(now.timestamp() * 1000)
        # 窗口起点取 max(活动创建时间, 30 天前)：避免活动配置晚于审批通过时漏同步
        window_start = now - timedelta(days=30)
        if activity.created_at and activity.created_at < window_start:
            window_start = activity.created_at
        start_ms = int(window_start.timestamp() * 1000)
        try:
            codes = await ac.list_instance_codes(
                activity.approval_code, start_ms, end_ms
            )
        except Exception as exc:  # noqa: BLE001
            logger.warning("拉取审批实例失败: activity=%s error=%s", activity.id, exc)
            return None
        pending_codes = [code for code in codes if code not in existing_map]
        skipped = len(codes) - len(pending_codes)
        # 并发拉取审批实例详情（限并发 5，独立请求无需串行等待）
        sem = asyncio.Semaphore(5)

        async def _fetch_instance(code: str) -> dict[str, Any] | None:
            try:
                async with sem:
                    return await ac.get_instance(code)
            except ac.TitleReviewBitableError as exc:
                logger.warning("拉取审批实例详情失败: %s error=%s", code, exc)
                return None

        instances = await asyncio.gather(
            *(_fetch_instance(code) for code in pending_codes)
        )
        pending: list[tuple[str, dict[str, Any]]] = []
        for code, instance in zip(pending_codes, instances):
            if instance is None or instance.get("status") != "APPROVED":
                continue
            fields = ac.form_widgets_to_fields(instance.get("form") or [])
            pending.append((code, fields))
        return pending, skipped

    async def _refresh_image_urls(
        self,
        activity: m.TitleReviewActivity,
        records: list[dict[str, Any]],
    ) -> int:
        """图片转存兜底：把单元格中残留的飞书远程链接下载转存为本地路径（幂等）。"""
        from app.modules.hr.title_review import bitable_client as bc

        app_token = activity.feishu_app_token
        apply_table_id = activity.apply_table_id
        if not (app_token and apply_table_id):
            return 0
        refreshed = 0
        for item in records:
            fields = item.get("fields") or {}
            update: dict[str, Any] = {}
            for col in IMAGE_EVIDENCE_FIELDS:
                text = _as_str(fields.get(col)) or ""
                if not _cell_has_remote_url(text):
                    continue
                local = await _image_text_to_local(text)
                if local != text:
                    update[col] = local
            if not update:
                continue
            try:
                await bc.update_record(
                    app_token,
                    apply_table_id,
                    str(item.get("record_id") or ""),
                    update,
                )
                refreshed += 1
                logger.info(
                    "图片已转存本地: record=%s cols=%s",
                    item.get("record_id"),
                    list(update),
                )
            except bc.TitleReviewBitableError as exc:
                logger.warning("图片转存写表失败: %s", exc)
        return refreshed

    # ═══ 兜底对账 ═══

    async def reconcile_activity(self, activity_id: UUID) -> dict[str, Any]:
        """拉飞书全量记录与本地双向对账（幂等）。"""
        from app.modules.hr.title_review import bitable_client as bc

        activity = await self.get_activity(activity_id)
        stats: dict[str, Any] = {
            "applications_created": 0,
            "applications_updated": 0,
            "terminal_applications_updated": 0,
            "applications_removed": 0,
            "votes_updated": 0,
            "errors": [],
        }
        # 审批先行：先同步审批通过的实例 → 写入申报表（随后本对账立即拉取落库）
        if activity.approval_code:
            approval_stats = await self.sync_approval_instances(activity.id)
            stats.update(approval_stats)
        if not (activity.feishu_app_token and activity.apply_table_id):
            return stats

        try:
            feishu_records = await bc.list_all_records(
                activity.feishu_app_token, activity.apply_table_id
            )
        except bc.TitleReviewBitableError as exc:
            stats["errors"].append(f"拉取申报表失败: {exc}")
            return stats
        feishu_ids = {item.get("record_id") for item in feishu_records}
        # 预加载本地申报 → 内存映射（替代逐行 get_by_feishu_record 的 N 次查询）
        local_apps = await self.application_repo.list_all_by_activity(activity.id)
        local_by_record = {
            a.feishu_record_id: a for a in local_apps if a.feishu_record_id
        }
        for item in feishu_records:
            record_id = str(item.get("record_id") or "")
            fields = item.get("fields") or {}
            existing = local_by_record.get(record_id)
            if existing:
                parsed = _apply_fields_to_dict(fields)
                changed = _apply_synced_fields(existing, parsed, fields)
                if changed:
                    stats["applications_updated"] += 1
                    if existing.status in TERMINAL_APPLICATION_STATUSES:
                        # 终态申报仅申报信息放开更新，票数/判定结果仍冻结
                        stats["terminal_applications_updated"] += 1
                profile_stale = (
                    existing.profile is None
                    or existing.profile_refreshed_at is None
                    or existing.profile_refreshed_at
                    < datetime.now().astimezone() - timedelta(hours=6)
                )
                refreshed = None
                if profile_stale and activity.feishu_app_token:
                    # 档案缺失或超过 6 小时自动重拉（员工信息表补录数据后生效）
                    refreshed = await self._load_employee_profile(
                        activity,
                        existing.name or parsed["name"],
                        existing.employee_no or parsed["employee_no"],
                    )
                    if refreshed is not None:
                        existing.profile = refreshed
                        existing.profile_refreshed_at = datetime.now().astimezone()
                if changed or refreshed is not None:
                    await self.application_repo.update(existing)
            else:
                try:
                    await self.sync_apply_record_added(activity.id, record_id, fields)
                    stats["applications_created"] += 1
                except IntegrityError:
                    # 审批实例编号全局唯一：申报行已归属其他活动时记入 errors 而非 500
                    logger.warning(
                        "申报记录重复/已归属其他活动，跳过: record_id=%s", record_id
                    )
                    stats["errors"].append(
                        f"申报记录 {record_id} 与已有记录重复（可能已绑定其他活动），已跳过"
                    )
        for application in local_apps:
            if (
                application.feishu_record_id
                and application.feishu_record_id not in feishu_ids
                and application.status not in TERMINAL_APPLICATION_STATUSES
            ):
                application.is_deleted = True
                await self.application_repo.update(application)
                stats["applications_removed"] += 1

        if activity.vote_table_id:
            try:
                vote_records = await bc.list_all_records(
                    activity.feishu_app_token, activity.vote_table_id
                )
            except bc.TitleReviewBitableError as exc:
                stats["errors"].append(f"拉取投票表失败: {exc}")
            else:
                for item in vote_records:
                    record_id = str(item.get("record_id") or "")
                    fields = item.get("fields") or {}
                    if not fields:
                        continue
                    try:
                        await self.sync_vote_record(activity.id, record_id, fields)
                        stats["votes_updated"] += 1
                    except Exception as exc:  # noqa: BLE001
                        logger.warning("对账同步投票行失败: record_id=%s error=%s", record_id, exc)
        await _audit(
            self.session,
            action="hr.title.reconcile",
            resource_type="title_review_activity",
            resource_id=activity.id,
            extra=stats,
        )
        return stats

    # ─── 校验辅助 ───

    @staticmethod
    def _ensure_draft(activity: m.TitleReviewActivity, field: str) -> None:
        if activity.status != m.ACTIVITY_DRAFT:
            raise HTTPException(400, f"{field}仅配置中（draft）可修改，当前：{activity.status}")


    async def get_summary(self, activity_id: UUID) -> dict[str, Any]:
        """结果汇总统计：总量 + 按职级分组 + 按部门分组。"""
        applications = self._scope_applications(
            await self.application_repo.list_all_by_activity(activity_id)
        )

        def _rate(passed: int, total: int) -> float | None:
            return round(passed / total * 100, 1) if total else None

        def _stats(apps: list[m.TitleReviewApplication]) -> dict[str, Any]:
            passed = sum(1 for a in apps if a.status in (m.APPLICATION_PASSED, m.APPLICATION_FINAL_PASSED))
            failed = sum(1 for a in apps if a.status in (m.APPLICATION_FAILED, m.APPLICATION_FINAL_FAILED))
            pending = sum(1 for a in apps if a.status not in (
                m.APPLICATION_PASSED, m.APPLICATION_FINAL_PASSED,
                m.APPLICATION_FAILED, m.APPLICATION_FINAL_FAILED,
            ))
            return {"applications": len(apps), "passed": passed, "failed": failed,
                    "pending": pending, "pass_rate": _rate(passed, len(apps))}

        by_level: dict[tuple[str, str], list[m.TitleReviewApplication]] = {}
        by_department: dict[str, list[m.TitleReviewApplication]] = {}
        for a in applications:
            key = (a.sequence or "未分类", a.apply_level or "未分类")
            by_level.setdefault(key, []).append(a)
            by_department.setdefault(a.department or "未分类", []).append(a)

        return {
            "total": _stats(applications),
            "by_level": [
                {"sequence": seq, "level_name": lv, **_stats(apps)}
                for (seq, lv), apps in sorted(by_level.items())
            ],
            "by_department": [
                {"department": dept, **_stats(apps)}
                for dept, apps in sorted(by_department.items())
            ],
        }

    async def export_results_xlsx(self, activity_id: UUID) -> tuple[bytes, str]:
        """导出评审结果汇总 xlsx（三个工作表：汇总统计/申报明细/评委明细）。"""
        import io

        from openpyxl import Workbook

        activity = await self.get_activity(activity_id)
        results = await self.get_results(activity_id)
        summary = await self.get_summary(activity_id)

        wb = Workbook()
        ws1 = wb.active
        assert ws1 is not None  # 新建工作簿必有默认工作表
        # Sheet1 汇总
        ws1.title = "汇总统计"
        ws1.append(["职称评审结果汇总", activity.name])
        ws1.append(["总体", "申报数", "通过", "未通过", "评审中", "通过率"])
        t = summary["total"]
        ws1.append(["合计", t["applications"], t["passed"], t["failed"], t["pending"], f'{t["pass_rate"]}%' if t["pass_rate"] is not None else "-"])
        ws1.append([])
        ws1.append(["按职级分组"])
        ws1.append(["序列", "职级", "申报数", "通过", "未通过", "评审中", "通过率"])
        for r in summary["by_level"]:
            ws1.append([r["sequence"], r["level_name"], r["applications"], r["passed"], r["failed"], r["pending"], f'{r["pass_rate"]}%' if r["pass_rate"] is not None else "-"])
        ws1.append([])
        ws1.append(["按部门分组"])
        ws1.append(["部门", "申报数", "通过", "未通过", "评审中", "通过率"])
        for r in summary["by_department"]:
            ws1.append([r["department"], r["applications"], r["passed"], r["failed"], r["pending"], f'{r["pass_rate"]}%' if r["pass_rate"] is not None else "-"])

        # Sheet2 申报明细
        ws2 = wb.create_sheet("申报明细")
        headers = ["序号", "姓名", "工号", "部门", "序列", "技术领域", "申报职级", "现任职级",
                   "是否破格", "同意票", "不同意票", "弃权票", "通过比例", "综合等级分布", "结果", "附件4综合意见"]
        ws2.append(headers)
        result_text = {m.APPLICATION_PASSED: "通过", m.APPLICATION_FINAL_PASSED: "通过",
                       m.APPLICATION_FAILED: "未通过", m.APPLICATION_FINAL_FAILED: "未通过"}
        for i, row in enumerate(results, start=1):
            a = row["application"]
            grades = {"合格": 0, "不合格": 0}
            for j in row["judges"]:
                g = j.get("comprehensive_grade")
                if g in grades:
                    grades[g] += 1
            ws2.append([
                i, a["name"], a["employee_no"], a["department"], a["sequence"], a["tech_domain"],
                a["apply_level"], a["current_level"], "是" if a["is_exception"] else "否",
                a["agree_votes"], a["oppose_votes"], a["abstain_votes"],
                f'{row["vote_ratio"] * 100:.1f}%' if row["vote_ratio"] is not None else "-",
                f"合格{grades['合格']}/不合格{grades['不合格']}",
                result_text.get(a["status"], a["status"]),
                a["final_opinion"] or "",
            ])

        # Sheet3 评委明细
        ws3 = wb.create_sheet("评委明细")
        ws3.append(["申报人", "工号", "申报职级", "评委编号", "评委姓名", "角色", "投票结果",
                    "综合等级", "维度评价", "评审意见", "投票时间"])
        for row in results:
            a = row["application"]
            for j in row["judges"]:
                dims = "；".join(f"{s['dimension_name']}:{s.get('grade') or '-'}" for s in j.get("scores", []))
                ws3.append([
                    a["name"], a["employee_no"], a["apply_level"], j["judge_code"], j["judge_name"],
                    j["judge_role"], j["vote_result"] or "未投", j["comprehensive_grade"] or "-",
                    dims, j["review_comment"] or "",
                    str(j["voted_at"] or ""),
                ])

        buf = io.BytesIO()
        await asyncio.to_thread(wb.save, buf)
        buf.seek(0)
        return buf.getvalue(), f"职称评审结果汇总_{activity.name}.xlsx"

    async def generate_roster_docx(
        self, activity_id: UUID, application_ids: list[UUID]
    ) -> tuple[bytes, str]:
        """生成最终名单 docx（表格：序号/职务/姓名/职级认定结果）。

        由 HR 勾选小组评审合格人员后导出，作为总经理确认与公示的依据。
        """
        import io

        from docx import Document
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Pt

        activity = await self.get_activity(activity_id)
        applications = await self.application_repo.list_all_by_activity(activity_id)
        applications = self._scope_applications(applications)
        by_id = {a.id: a for a in applications}
        selected = [by_id[i] for i in application_ids if i in by_id]

        def _apply_song_font(run: Any, size: float = 10.5) -> None:
            """统一宋体（含中文字体 eastAsia）。"""
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(size)

        doc = Document()
        # 文档默认样式统一宋体（五号）
        normal = doc.styles["Normal"]
        normal.font.name = "宋体"
        normal.font.size = Pt(10.5)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中
        title_run = title.add_run(f"{activity.name} 职级认定结果名单")
        title_run.bold = True
        _apply_song_font(title_run, size=16)

        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ["序号", "部门", "职务", "姓名", "本年度认定职称"]
        for i, text in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = text
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in cell.paragraphs[0].runs:
                _apply_song_font(run)
                run.bold = True
        for idx, a in enumerate(selected, start=1):
            row = table.add_row().cells
            # 职务取自员工信息表档案（岗位信息）
            position = ((a.profile or {}).get("职务")) or "-"
            values = [str(idx), a.department or "-", position, a.name, a.apply_level or "-"]
            for col, text in enumerate(values):
                row[col].text = text
                for run in row[col].paragraphs[0].runs:
                    _apply_song_font(run)

        buf = io.BytesIO()
        await asyncio.to_thread(doc.save, buf)
        buf.seek(0)
        return buf.getvalue(), f"{activity.name}_职级认定结果名单.docx"
