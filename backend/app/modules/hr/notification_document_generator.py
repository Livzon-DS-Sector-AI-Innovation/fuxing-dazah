"""培训通知 Word 文档生成器 — 基于表格模板."""

from datetime import date, datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from pydantic import BaseModel


class TrainingNotificationInput(BaseModel):
    department: str
    training_date: date
    subject: str
    training_time_start: str | None = None
    training_time_end: str | None = None
    face_to_face_time_start: str | None = None
    face_to_face_time_end: str | None = None
    self_study_time_start: str | None = None
    self_study_time_end: str | None = None
    face_date: date | None = None
    self_study_date: date | None = None
    time_slots: list[dict] | None = None
    location: str | None = None
    trainer: str | None = None
    training_method: str | None = None
    assessment_method: str | None = None
    content: str | None = None
    trainee_names: list[str] = []
    issuer_department: str | None = None
    issue_date: date | None = None


def _find_template() -> Path:
    candidates = [
        Path("assets/hr/7.4培训通知书.docx"),
        Path("../assets/hr/7.4培训通知书.docx"),
        Path(__file__).resolve().parent.parent.parent.parent
        / "assets/hr"
        / "7.4培训通知书.docx",
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError("模板文件未找到: 7.4培训通知书.docx")


def _set_cell(cell, text: str) -> None:
    """Set cell text, preserving first run's formatting."""
    first = None
    for p in cell.paragraphs:
        for r in p.runs:
            if first is None:
                first = r
            r.text = ""
    if first is not None:
        first.text = str(text or "")
    elif cell.paragraphs:
        cell.paragraphs[0].add_run(str(text or ""))


def _compute_hours(start: str | None, end: str | None) -> str:
    if not start or not end:
        return ""
    try:
        s = datetime.strptime(start, "%H:%M")
        e = datetime.strptime(end, "%H:%M")
        diff = (e - s).total_seconds() / 3600
        if diff <= 0:
            return ""
        rounded = round(diff * 2) / 2
        if rounded == int(rounded):
            return f"{int(rounded)}小时"
        return f"{rounded}小时"
    except ValueError:
        return ""


def generate_training_notification(data: TrainingNotificationInput) -> BytesIO:
    """Generate training notification docx from template."""
    template_path = _find_template()
    doc = Document(str(template_path))
    table = doc.tables[0]

    # ── Row 0: 培训内容 (label col0, value cols 1-3 merged) ──
    topic_parts = [data.subject]
    if data.content:
        topic_parts.append(data.content)
    _set_cell(table.rows[0].cells[1], " — ".join(topic_parts))

    # ── Row 1: 培训日期 | value (含时间) | 课时 | value ──
    # 优先使用 time_slots（多时段），否则回退到旧字段
    if data.time_slots:
        slot_parts = []
        total_minutes = 0
        for s in data.time_slots:
            d = s.get("date", "")
            t_start = s.get("start", "")
            t_end = s.get("end", "")
            if d and t_start:
                slot_parts.append(f"{d} {t_start}—{t_end}")
            if t_start and t_end:
                try:
                    s_dt = datetime.strptime(t_start, "%H:%M")
                    e_dt = datetime.strptime(t_end, "%H:%M")
                    diff = (e_dt - s_dt).total_seconds() / 60
                    if diff > 0:
                        total_minutes += diff
                except ValueError:
                    pass
        date_text = "；".join(slot_parts)
        total_hours = round(total_minutes / 30) / 2
        hours_text = f"{int(total_hours)}小时" if total_hours == int(total_hours) else f"{total_hours}小时" if total_hours > 0 else ""
        _set_cell(table.rows[1].cells[1], date_text)
        _set_cell(table.rows[1].cells[3], hours_text)
    else:
        face_start = getattr(data, "face_to_face_time_start", None) or data.training_time_start
        face_end = getattr(data, "face_to_face_time_end", None) or data.training_time_end
        self_start = getattr(data, "self_study_time_start", None)
        self_end = getattr(data, "self_study_time_end", None)
        face_date = getattr(data, "face_date", None) or data.training_date
        self_date = getattr(data, "self_study_date", None)
        # 格式化中文日期
        def _fmt_date(d, t_start, t_end, label):
            if not d or not t_start:
                return ""
            ds = f"{d.year}年{d.month}月{d.day}日"
            if t_start and t_end:
                return f"{label}：{ds} {t_start}—{t_end}"
            return f"{label}：{ds} {t_start}"
        parts = []
        face_part = _fmt_date(face_date, face_start, face_end, "面授")
        if face_part:
            parts.append(face_part)
        if self_date and self_start:
            parts.append(_fmt_date(self_date, self_start, self_end, "自学"))
        date_text = "；".join(parts) if parts else (str(data.training_date) if data.training_date else "")
        _set_cell(table.rows[1].cells[1], date_text)
        # 课时仅按面授时间计算
        _set_cell(table.rows[1].cells[3], _compute_hours(face_start, face_end))

    # ── Row 2: 培训方式 | value | 授课人 | value ──
    _set_cell(table.rows[2].cells[1], data.training_method or "")
    _set_cell(table.rows[2].cells[3], data.trainer or "")

    # ── Row 3: 培训对象 (merged cols 1-3) ──
    # 培训对象统一以「x部门 全体成员」展示，不再列人名；附病假/产假人数
    if data.trainee_names:
        people = f"「{data.department}」全体成员（{len(data.trainee_names)}人）"
    else:
        people = f"「{data.department}」全体成员"
    sick = data.sick_count if data.sick_count is not None else 0
    maternity = data.maternity_count if data.maternity_count is not None else 0
    people += f"（病假{sick}人、产假{maternity}人除外）"
    _set_cell(table.rows[3].cells[1], people)

    # ── Row 4: 培训地点 (merged cols 1-3) ──
    _set_cell(table.rows[4].cells[1], data.location or "")

    # ── Row 5: 考核方式 (merged cols 1-3) ──
    _set_cell(table.rows[5].cells[1], data.assessment_method or "")

    # ── Row 6: 注意事项 (merged cols 1-3) ──
    _set_cell(table.rows[6].cells[1],
        "1. 请培训人员自带笔记本、笔，做好笔记。\n"
        "2. 请部门安排好参训人员的工作时间，做到培训工作两不误。\n"
        "3. 不得无故缺席、迟到，到场签到，有特殊情况须提前请假。"
    )

    # ── Para 0: department line ──
    for p_idx in (0, 1):
        if p_idx < len(doc.paragraphs):
            p = doc.paragraphs[p_idx]
            # Clear all runs first
            for r in p.runs:
                r.text = ""
    if len(doc.paragraphs) > 0:
        dept_name = data.issuer_department or data.department or ""
        doc.paragraphs[0].runs[0].text = f"部门/Dept：{dept_name}          签发人/ Issued by："
    if len(doc.paragraphs) > 1:
        # 优先从 time_slots 取首个日期，否则使用 training_date
        issue_d = data.training_date
        if data.time_slots:
            first_slot_date = data.time_slots[0].get("date", "")
            if first_slot_date:
                try:
                    issue_d = datetime.strptime(first_slot_date, "%Y-%m-%d").date()
                except ValueError:
                    pass
        d = issue_d
        doc.paragraphs[1].runs[0].text = f"{d.year}年{d.month:02d}月{d.day:02d}日"

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
