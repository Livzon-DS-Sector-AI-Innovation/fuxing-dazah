"""花名册 Word 文档生成器"""

import re
from copy import deepcopy
from io import BytesIO
from pathlib import Path

from docx import Document

TEMPLATE = Path(__file__).parent.parent.parent.parent / "assets" / "hr" / "花名册-模板.docx"


def generate_roster_sync(employees: list[tuple], department: str | None = None) -> BytesIO:
    """同步生成花名册。employees: list of (name, department, gender, education, hire_date, status)"""
    if not TEMPLATE.exists():
        raise FileNotFoundError(f"模板未找到: {TEMPLATE}")

    doc = Document(str(TEMPLATE))

    # 修正页眉：保持 run 结构，仅加下划线
    for section in doc.sections:
        for p in section.header.paragraphs:
            for r in p.runs:
                t = r.text or ''
                if t.strip() == '第':
                    r.text = '第____'
                elif t.strip() == '共':
                    r.text = '共____'

    table = doc.tables[0]
    if len(table.rows) < 2 or not employees:
        return _save(doc)

    template_row = table.rows[1]
    tr_el = template_row._tr

    # Clone rows
    last_tr = tr_el
    for _ in range(1, len(employees)):
        new_tr = deepcopy(tr_el)
        last_tr.addnext(new_tr)
        last_tr = new_tr

    # Fill rows
    for i, emp in enumerate(employees):
        row_idx = 1 + i
        if row_idx >= len(table.rows):
            break
        name, dept, gender, edu, hire_date, status = emp
        hire_str = str(hire_date) if hire_date else ""

        mapping = {
            "{$index+1}": str(i + 1),
            "{体现部门}": dept or "",
            "{姓名}": name or "",
            "{性别}": gender or "",
            "{学历}": edu or "",
            "{人员状态}": status or "在职",
        }

        for cell in table.rows[row_idx].cells:
            if not cell.paragraphs or not cell.paragraphs[0].runs:
                continue
            full = "".join(r.text for r in cell.paragraphs[0].runs)
            for ph, val in mapping.items():
                full = full.replace(ph, val)
            full = re.sub(r'\{入职日期\|[^}]*\}', hire_str, full)
            full = re.sub(r'\{\#[\w一-鿿/]+\}', '', full)
            full = re.sub(r'\{\/[\w一-鿿/]+\}', '', full)

            cell.paragraphs[0].runs[0].text = full
            for r in cell.paragraphs[0].runs[1:]:
                r.text = ""

    return _save(doc)


def _save(doc: Document) -> BytesIO:
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
