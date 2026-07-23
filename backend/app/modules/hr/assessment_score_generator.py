"""实操考核成绩单 Word 文档生成器."""

from datetime import date
from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def generate_assessment_score_sheet(
    training_content: str,
    training_date: str,
    department: str,
    scores: list[dict],  # [{"name": "张三", "department": "仓储部", "score": 95}, ...]
) -> BytesIO:
    """生成实操考核成绩单 Word 文档。"""
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)

    # ── 标题 ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('丽珠集团福州福兴医药有限公司')
    run.bold = True
    run.font.size = Pt(14)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('员工培训考核成绩单')
    run.bold = True
    run.font.size = Pt(12)

    # ── 信息行 ──
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    info_text = f'培训内容：{training_content}\n培训日期：{training_date}\n培训部门：{department}'
    run = info.add_run(info_text)
    run.font.size = Pt(10.5)

    doc.add_paragraph()  # 空行

    # ── 表格 ──
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    hdr = table.rows[0]
    headers = ['序号', '姓名', '部门', '成绩']
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10.5)

    # 数据行
    for idx, s in enumerate(scores, 1):
        row = table.add_row()
        values = [str(idx), s.get('name', ''), s.get('department', ''), str(s.get('score', ''))]
        for i, text in enumerate(values):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.size = Pt(10.5)

    doc.add_paragraph()  # 空行

    # ── 页脚信息 ──
    footer_text = (
        'TEM.02.1206.002\n'
        'LIVZON GROUP FUZHOU FUXING PHARMACEUTICAL CO., LTD.\n'
        '丽珠集团福州福兴医药有限公司\n'
        'ADD: No. 8 Nangang Road, Jiangyin Industrial Concentration Zone, '
        'Fuqing, Fuzhou City, Fujian Province, P. R. China. 350309\n'
        'TEL: +86-591-85966932; FAX:+86-591-85966925\n'
        'E-MAIL: regulation@fxpharm.com; fxqa@fxpharm.com\n'
        'Website: www.fxpharm.com'
    )
    footer = doc.add_paragraph()
    run = footer.add_run(footer_text)
    run.font.size = Pt(9)
    run.font.color.rgb = None  # default color

    # ── 签名行 ──
    sign = doc.add_paragraph()
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = sign.add_run(f'考核人：{"_" * 10}    日期：{date.today().strftime("%Y年%m月%d日")}')
    run.font.size = Pt(10.5)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
