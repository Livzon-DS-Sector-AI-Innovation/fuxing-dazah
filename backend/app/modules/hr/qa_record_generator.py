"""现场问答/实操记录表 Word 文档生成器 — 基于表格模板."""

import copy
from datetime import date
from io import BytesIO

from docx import Document

from app.modules.hr.template_utils import find_hr_template

NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# 无题目导出时，考题区保留的空白行数（供打印后手写）
_BLANK_QUESTION_ROWS = 5


def _replace_all_in_cell(cell, mapping: dict[str, str]) -> None:
    """替换单元格中所有占位符（处理跨 run 拆分的情况）."""
    for p in cell.paragraphs:
        # 先拼接整段文本
        full_text = "".join(r.text or "" for r in p.runs)
        # 替换所有占位符
        for old, new in mapping.items():
            full_text = full_text.replace(old, new)
        # 写回第一个 run，清空其余
        if p.runs:
            p.runs[0].text = full_text
            for r in p.runs[1:]:
                r.text = ""


def _clear_cell(cell) -> None:
    """清空单元格文本."""
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""


def _set_cell_text(cell, text: str) -> None:
    """设置单元格文本."""
    _clear_cell(cell)
    if cell.paragraphs:
        runs = cell.paragraphs[0].runs
        if runs:
            runs[0].text = text
        else:
            cell.paragraphs[0].add_run(text)


def _insert_row_after(table, after_idx: int) -> int:
    """在指定行之后插入一行（深拷贝样式），返回新行索引."""
    ref_row = table.rows[after_idx]
    new_tr = copy.deepcopy(ref_row._tr)
    # 清空新行所有单元格
    for tc in new_tr.findall(f".//{{{NAMESPACE}}}t"):
        tc.text = ""
        tc.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    ref_row._tr.addnext(new_tr)
    return after_idx + 1


def generate_qa_record(
    *,
    training_content: str = "",
    training_purpose: str = "",
    training_date: date | None = None,
    training_method: str = "",
    training_department: str = "",
    questions: list[dict] | None = None,
    trainee_names: list[str] | None = None,
    scores: list[dict] | None = None,
    trainer_name: str = "",
) -> BytesIO:
    """生成现场问答/实操记录表."""
    template_path = find_hr_template("现场问答，实操记录表-模板1.docx")
    doc = Document(str(template_path))
    table = doc.tables[0]

    date_str = str(training_date) if training_date else ""
    questions = questions or []
    trainee_names = trainee_names or []

    # ── 第 1 遍：替换所有占位符 ──
    header_mapping = {
        "{培训内容}": training_content or "",
        "{实施日期}": date_str,
        "{培训方式}": training_method or "",
        "{培训对象}": training_department or "",
    }

    # 遍历所有单元格，替换头部占位符和清除模板标记
    for row in table.rows:
        for cell in row.cells:
            _replace_all_in_cell(cell, header_mapping)
            # 清除模板标记
            _replace_all_in_cell(cell, {
                "{#题库}": "",
                "{/题库}": "",
                "{$index+1}": "",
                "{#记录_签到表}": "",
                "{/记录_签到表}": "",
            })

    # 培训目的（行1，cols 3-8）—— 模板无占位符，直接写入
    if training_purpose:
        _set_cell_text(table.rows[1].cells[3], training_purpose or "")

    # ── 考题区域：row 5 是模板，row 6 是空行 ──
    question_tpl_row = 5

    # 先清除模板行遗留标记
    for cell in table.rows[question_tpl_row].cells:
        _replace_all_in_cell(cell, {
            "{文件编号}": "",
            "{题目}": "",
            "{正确答案}": "",
            "{分值}": "",
            "{姓名}": "",
            "{答题结果}": "",
            "{总分}": "",
        })

    if questions:
        # 第一道题填入模板行
        q0 = questions[0]
        _set_cell_text(table.rows[question_tpl_row].cells[0], "1")
        _set_cell_text(table.rows[question_tpl_row].cells[1], q0.get("file_no", ""))
        _set_cell_text(table.rows[question_tpl_row].cells[4], q0.get("question", ""))
        _set_cell_text(table.rows[question_tpl_row].cells[7], q0.get("answer", ""))
        _set_cell_text(table.rows[question_tpl_row].cells[8], str(q0.get("score", "")))

        # 其余题目：在模板行之后插入
        insert_pos = question_tpl_row
        for qi, q in enumerate(questions[1:], start=2):
            insert_pos = _insert_row_after(table, insert_pos)
            _set_cell_text(table.rows[insert_pos].cells[0], str(qi))
            _set_cell_text(table.rows[insert_pos].cells[1], q.get("file_no", ""))
            _set_cell_text(table.rows[insert_pos].cells[4], q.get("question", ""))
            _set_cell_text(table.rows[insert_pos].cells[7], q.get("answer", ""))
            _set_cell_text(table.rows[insert_pos].cells[8], str(q.get("score", "")))
    else:
        # 没有题目：清空模板行占位符，并补足空白行供打印后手写
        for cell in table.rows[question_tpl_row].cells:
            _clear_cell(cell)
        insert_pos = question_tpl_row
        for _ in range(_BLANK_QUESTION_ROWS - 1):
            insert_pos = _insert_row_after(table, insert_pos)

    # 删除模板自带的空行
    empty_row_idx = question_tpl_row + 1 + (max(0, len(questions) - 1) if questions else _BLANK_QUESTION_ROWS - 1)
    if empty_row_idx < len(table.rows):
        table._tbl.remove(table.rows[empty_row_idx]._tr)

    # ── 受训人员区域：找到「姓名」表头行 ──
    trainee_header_row = None
    for ri in range(question_tpl_row + 1, len(table.rows)):
        cell_text = table.rows[ri].cells[0].text.strip() if table.rows[ri].cells else ""
        if "姓名" in cell_text and "得分情况" in (table.rows[ri].cells[2].text if len(table.rows[ri].cells) > 2 else ""):
            trainee_header_row = ri
            break
    if trainee_header_row is None:
        trainee_header_row = question_tpl_row + 2  # fallback
    trainee_tpl_row = trainee_header_row + 1

    # 清除受训人员模板行遗留标记
    for cell in table.rows[trainee_tpl_row].cells:
        _replace_all_in_cell(cell, {"{姓名}": "", "{答题结果}": "", "{总分}": ""})

    # 成绩行优先（含得分情况/考核日期/总分），否则退回仅姓名
    # 行内 cell 布局：0=姓名 2=得分情况 6=考核日期 8=总分（合并单元格）
    def _fill_trainee_row(row_idx: int, entry: dict) -> None:
        cells = table.rows[row_idx].cells
        _set_cell_text(cells[0], entry.get("name", ""))
        if len(cells) > 2:
            _set_cell_text(cells[2], entry.get("result_text", ""))
        if len(cells) > 6:
            _set_cell_text(cells[6], entry.get("assessed_date", ""))
        if len(cells) > 8:
            _set_cell_text(cells[8], entry.get("total_score", ""))

    entries = (
        [dict(s) for s in scores]
        if scores
        else [{"name": n} for n in trainee_names]
    )
    if entries:
        _fill_trainee_row(trainee_tpl_row, entries[0])
        insert_pos = trainee_tpl_row
        for entry in entries[1:]:
            insert_pos = _insert_row_after(table, insert_pos)
            _fill_trainee_row(insert_pos, entry)
    else:
        for cell in table.rows[trainee_tpl_row].cells:
            _clear_cell(cell)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
